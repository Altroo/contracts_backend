"""
contract/st_doc.py
~~~~~~~~~~~~~~~~~~
Generates a Sous-Traitance DOCX using python-docx with CDL branding
(dark #0F0F1A + gold #B8973A, Cormorant Garamond + Inter).
"""
# i18n: skip-file — bilingual document generator; FR+EN content is intentional

from datetime import datetime
from io import BytesIO
from typing import Any, cast

from django.http import HttpResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from core.models import CompanyConfig
from .pdf import _fmt_date, _fmt_amt
from .st_i18n import (
    LOT_LABELS,
    LOT_NORMES,
    LOT_ASSURANCES,
    LOT_DESC_DEFAULT,
    LOT_OBLIGATIONS,
    LOT_RECEPTION,
    FORME_LABELS,
    TYPE_PRIX_LABELS,
    DELAI_UNIT_LABELS,
    st_t,
)

# ── helpers ──────────────────────────────────────────────────────────────────


def _resolve_lot_keys(lot_type) -> list:
    """Handle JSONField (list) or legacy string for st_lot_type."""
    if isinstance(lot_type, list):
        return [k for k in lot_type if k]
    return [lot_type] if lot_type else []


def _resolve_type_prix_keys(type_prix) -> list:
    """Handle JSONField (list) or legacy string for st_type_prix. Defaults to forfaitaire."""
    if isinstance(type_prix, list):
        keys = [k for k in type_prix if k]
        return keys if keys else ["forfaitaire"]
    return [type_prix] if type_prix else ["forfaitaire"]


# ── colour palette ───────────────────────────────────────────────────────────

DARK = RGBColor(0x0F, 0x0F, 0x1A)
GOLD = RGBColor(0xB8, 0x97, 0x3A)
INK = RGBColor(0x1C, 0x1C, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x6B, 0x6B, 0x80)
DARK_HEX = "0F0F1A"
GOLD_HEX = "B8973A"
GOLD_PALE_HEX = "F7F0E0"
CREAM_HEX = "FAF6EC"
BORDER_HEX = "E2D9C8"


# ── OxmlElement helpers ──────────────────────────────────────────────────────


def _num_to_words_mad(n: int, lang: str) -> str:
    """Convert integer n to words for the contract 'arrêté à la somme de' line."""
    try:
        from num2words import num2words  # type: ignore[import-untyped]

        locale = "fr" if lang == "fr" else "en"
        return num2words(n, lang=locale).capitalize()
    except Exception:  # ImportError or any num2words error
        return f"{n:,}".replace(",", "\u202f")


def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _para_bg(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _para_border_left(para, color: str = "B8973A", size: str = "12"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _para_border_bottom(para, color: str = "E2D9C8", size: str = "4"):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(val[1]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val[0])
            tcBorders.append(el)
        else:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_width(cell, cm_val):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm_val * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_row_height(row, pt_val):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(pt_val * 20)))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


def _run_shd(run, hex_color: str):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


# ── Generator class ──────────────────────────────────────────────────────────


class SousTraitanceDOCGenerator:
    """Generate a python-docx DOCX for a Sous-Traitance contract."""

    def __init__(self, contract, language: str = "fr"):
        self.c = contract
        self.lang = language
        self.fr = language == "fr"
        self.doc = Document()
        self._ep = None
        self._art_num = 0
        self._configure_page()

    # ── setup ────────────────────────────────────────────────────────────────

    def _configure_page(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        normal = cast(Any, self.doc.styles["Normal"])
        normal.font.name = "Inter"
        normal.font.size = Pt(9)
        normal.font.color.rgb = INK
        normal.paragraph_format.space_after = Pt(2)
        normal.paragraph_format.line_spacing = 1.75

    # ── helpers ──────────────────────────────────────────────────────────────

    def _t(self, key: str) -> str:
        return st_t(key, self.lang)

    @property
    def ep(self):
        if self._ep is None:
            try:
                self._ep = CompanyConfig.objects.get(company=self.c.company)
            except CompanyConfig.DoesNotExist:
                self._ep = None
        return self._ep

    @property
    def _ht(self) -> float:
        try:
            return float(self.c.montant_ht or 0)
        except (TypeError, ValueError):
            return 0.0

    @property
    def _tva_pct(self) -> float:
        try:
            return float(self.c.tva if self.c.tva is not None else 20)
        except (TypeError, ValueError):
            return 20.0

    @property
    def _tva_amt(self) -> float:
        return self._ht * self._tva_pct / 100

    @property
    def _ttc(self) -> float:
        return self._ht + self._tva_amt

    def _dev(self) -> str:
        return self.c.devise or "MAD"

    def _lot(self) -> str:
        labels = LOT_LABELS.get(self.lang, LOT_LABELS["fr"])
        lot_keys = _resolve_lot_keys(self.c.st_lot_type)
        if not lot_keys:
            return ""
        return " / ".join(labels.get(k, k) for k in lot_keys)

    def _next_art(self) -> int:
        self._art_num += 1
        return self._art_num

    def _add_empty(self, pt=2):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(pt)
        p.paragraph_format.space_after = Pt(0)
        return p

    def _add_text(self, text, size=Pt(9), color=INK, bold=False, align=None):
        p = self.doc.add_paragraph()
        if align:
            p.alignment = align
        r = p.add_run(text)
        r.font.size = size
        r.font.color.rgb = color
        r.bold = bold
        return p, r

    def _add_bullet(self, text, size=Pt(8.5), color=INK, indent_cm=0.6):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"• {text}")
        r.font.size = size
        r.font.color.rgb = color
        return p, r

    # ── art title ────────────────────────────────────────────────────────────

    def _add_art_title(self, title: str):
        n = self._next_art()
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _para_bg(p, GOLD_PALE_HEX)
        _para_border_left(p, GOLD_HEX, "12")

        # Article number badge
        num_run = p.add_run(f" {n:02d} ")
        num_run.font.size = Pt(8)
        num_run.font.color.rgb = GOLD
        num_run.bold = True
        _run_shd(num_run, DARK_HEX)

        # Title text
        tr = p.add_run(f"  {title}")
        tr.font.size = Pt(8)
        tr.font.color.rgb = DARK
        tr.bold = True

    def _add_sub_title(self, text: str):
        p, r = self._add_text(text, size=Pt(8.5), color=DARK, bold=True)
        p.paragraph_format.space_before = Pt(6)
        return p

    # ── section builders ─────────────────────────────────────────────────────

    def _build_topstrip(self):
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)
        cell = table.cell(0, 0)
        _cell_bg(cell, GOLD_HEX)
        _set_row_height(table.rows[0], 3)
        cell.text = ""

    def _build_header(self):
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        ep = self.ep
        ep_name = ep.name if ep else "CASA DI LUSSO"
        ep_rc = ep.rc if ep else ""
        ep_ice = ep.ice if ep else ""
        ep_addr = ep.adresse if ep else ""

        # Left — logo
        left = table.cell(0, 0)
        _set_cell_width(left, 9.5)
        pl = left.paragraphs[0]
        lr = pl.add_run("CASA ")
        lr.font.name = "Cormorant Garamond"
        lr.font.size = Pt(20)
        lr.font.color.rgb = DARK
        lr.bold = True
        gr = pl.add_run("DI LUSSO")
        gr.font.name = "Cormorant Garamond"
        gr.font.size = Pt(20)
        gr.font.color.rgb = GOLD
        gr.bold = True

        pt = left.add_paragraph()
        tr = pt.add_run("ENTREPRENEUR PRINCIPAL")
        tr.font.size = Pt(7.5)
        tr.font.color.rgb = MUTED
        tr.font.all_caps = True

        pi = left.add_paragraph()
        ir = pi.add_run(f"{ep_name}\nRC {ep_rc} · ICE {ep_ice}\n{ep_addr}")
        ir.font.size = Pt(7)
        ir.font.color.rgb = MUTED

        # Right — ref & date
        right = table.cell(0, 1)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref = self.c.numero_contrat or ""
        pr = right.paragraphs[0]
        rr = pr.add_run(f"{self._t('ref')} {ref}")
        rr.font.size = Pt(9)
        rr.font.color.rgb = GOLD
        rr.bold = True
        _run_shd(rr, DARK_HEX)

        pd = right.add_paragraph()
        pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dr = pd.add_run(f"{self._t('date_label')} {_fmt_date(self.c.date_contrat)}")
        dr.font.size = Pt(8)
        dr.font.color.rgb = MUTED

        if self.c.st_projet:
            pp = right.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pr2 = pp.add_run(f"{self._t('projet')} : {self.c.st_projet.name}")
            pr2.font.size = Pt(8)
            pr2.font.color.rgb = MUTED

        # separator
        sep = self.doc.add_paragraph()
        _para_border_bottom(sep, BORDER_HEX, "2")
        sep.paragraph_format.space_after = Pt(8)

    def _build_title(self):
        sep = self.doc.add_paragraph()
        _para_border_bottom(sep, "EEEEEE", "2")
        sep.paragraph_format.space_after = Pt(6)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self._t("contrat_title"))
        r.font.name = "Cormorant Garamond"
        r.font.size = Pt(14)
        r.font.color.rgb = DARK
        r.bold = True
        r.font.all_caps = True

        ps = self.doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = ps.add_run(self._t("preamble"))
        rs.font.size = Pt(8)
        rs.font.color.rgb = MUTED

        lot_label = self._lot()
        if lot_label:
            pb = self.doc.add_paragraph()
            pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            br = pb.add_run(f"  {lot_label}  ")
            br.font.size = Pt(7.5)
            br.font.color.rgb = GOLD
            br.bold = True
            _run_shd(br, GOLD_PALE_HEX)

        sep2 = self.doc.add_paragraph()
        _para_border_bottom(sep2, "EEEEEE", "2")
        sep2.paragraph_format.space_after = Pt(8)

    def _build_parties(self):
        t = self._t
        ep = self.ep
        c = self.c
        lang = self.lang

        # "ENTRE LES SOUSSIGNÉS"
        pe, re_ = self._add_text(t("entre"), size=Pt(9), color=MUTED, bold=True)
        pe.paragraph_format.space_after = Pt(4)

        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        # ── EP cell ──
        ep_cell = table.cell(0, 0)
        _set_cell_width(ep_cell, 8.5)
        _cell_bg(ep_cell, DARK_HEX)
        _cell_borders(
            ep_cell,
            top=(GOLD_HEX, 2),
            bottom=(GOLD_HEX, 2),
            left=(GOLD_HEX, 2),
            right=(GOLD_HEX, 2),
        )

        lbl_p = ep_cell.paragraphs[0]
        lr = lbl_p.add_run(t("lbl_ep"))
        lr.font.size = Pt(7)
        lr.font.color.rgb = GOLD
        lr.bold = True
        lr.font.all_caps = True

        ep_name = ep.name if ep else "CASA DI LUSSO SARL"
        ep_forme = ep.forme_juridique if ep else "SARL"
        ep_capital = _fmt_amt(ep.capital, self._dev()) if ep and ep.capital else "—"
        ep_rc = ep.rc if ep else "—"
        ep_ice = ep.ice if ep else "—"
        ep_if = ep.identifiant_fiscal if ep else "—"
        ep_addr = ep.adresse if ep else "—"
        ep_rep = ep.representant if ep else "—"
        ep_qualite = ep.qualite_representant if ep else "Gérant"

        ep_lines = [
            (t("raison_sociale"), ep_name, True),
            (t("forme_juridique"), ep_forme, False),
            (t("capital_social"), ep_capital, False),
            (t("rc"), ep_rc, False),
            (t("ice"), ep_ice, False),
            (t("if_label"), ep_if, False),
            (t("siege"), ep_addr, False),
            (t("representant"), ep_rep, False),
            (t("qualite"), ep_qualite, False),
        ]
        for lbl, val, bold in ep_lines:
            pp = ep_cell.add_paragraph()
            rr = pp.add_run(f"{lbl} : ")
            rr.font.size = Pt(8.5)
            rr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            vr = pp.add_run(val)
            vr.font.size = Pt(8.5)
            vr.font.color.rgb = WHITE
            vr.bold = bold

        ci_p = ep_cell.add_paragraph()
        ci_r = ci_p.add_run(t("ci_ep"))
        ci_r.font.size = Pt(8)
        ci_r.italic = True
        ci_r.font.color.rgb = MUTED

        # ── ST cell ──
        st_cell = table.cell(0, 1)
        _set_cell_width(st_cell, 8.5)
        _cell_bg(st_cell, CREAM_HEX)
        _cell_borders(
            st_cell,
            top=(GOLD_HEX, 2),
            bottom=(GOLD_HEX, 2),
            left=(GOLD_HEX, 2),
            right=(GOLD_HEX, 2),
        )

        sl = st_cell.paragraphs[0]
        slr = sl.add_run(t("lbl_st"))
        slr.font.size = Pt(7)
        slr.font.color.rgb = GOLD
        slr.bold = True
        slr.font.all_caps = True

        st_forme_key = c.st_forme or ""
        st_forme_lbl = FORME_LABELS.get(lang, FORME_LABELS["fr"]).get(
            st_forme_key, st_forme_key
        )

        st_lines = [
            (t("raison_sociale"), c.st_name or "—", True),
            (t("forme_juridique"), st_forme_lbl, False),
            (
                t("capital_social"),
                _fmt_amt(c.st_capital, self._dev()) if c.st_capital else "—",
                False,
            ),
            (t("rc"), c.st_rc or "—", False),
            (t("ice"), c.st_ice or "—", False),
            (t("if_label"), c.st_if or "—", False),
            (t("cnss"), c.st_cnss or "—", False),
            (t("siege"), c.st_addr or "—", False),
            (t("representant"), c.st_rep or "—", False),
            (t("cin"), c.st_cin or "—", False),
            (t("qualite"), c.st_qualite or "—", False),
            (t("telephone"), c.st_tel or "—", False),
            (t("email"), c.st_email or "—", False),
            (t("rib"), c.st_rib or "—", False),
            (t("banque"), c.st_banque or "—", False),
        ]
        for lbl, val, bold in st_lines:
            pp = st_cell.add_paragraph()
            rr = pp.add_run(f"{lbl} : ")
            rr.font.size = Pt(8.5)
            rr.font.color.rgb = MUTED
            vr = pp.add_run(str(val))
            vr.font.size = Pt(8.5)
            vr.font.color.rgb = DARK
            vr.bold = bold

        ci_s = st_cell.add_paragraph()
        ci_sr = ci_s.add_run(t("ci_st"))
        ci_sr.font.size = Pt(8)
        ci_sr.italic = True
        ci_sr.font.color.rgb = MUTED

        # "Ensemble désignées" closing sentence
        ens_p = self.doc.add_paragraph()
        ens_r = ens_p.add_run(t("parties_ensemble"))
        ens_r.font.size = Pt(8.5)
        ens_r.italic = True
        ens_r.font.color.rgb = MUTED
        ens_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ens_p.paragraph_format.space_before = Pt(4)
        ens_p.paragraph_format.space_after = Pt(8)

    # ── articles ─────────────────────────────────────────────────────────────

    def _build_art_parties(self):
        t = self._t
        self._add_art_title(t("art_parties"))
        self._add_sub_title(t("partie_ep_title"))
        self._add_text(t("ci_ep"), size=Pt(8.5), color=MUTED)
        self._add_sub_title(t("partie_st_title"))
        self._add_text(t("ci_st"), size=Pt(8.5), color=MUTED)

    def _build_art_objet(self):
        t = self._t
        c = self.c
        lang = self.lang
        self._add_art_title(t("art_objet"))

        self._add_text(t("objet_intro"), size=Pt(8.5))

        lot_keys = _resolve_lot_keys(c.st_lot_type)
        lot_key = lot_keys[0] if lot_keys else ""
        labels_map = LOT_LABELS.get(lang, LOT_LABELS["fr"])
        lot_label = " / ".join(labels_map.get(k, k) for k in lot_keys) if lot_keys else ""
        lot_desc = c.st_lot_description or LOT_DESC_DEFAULT.get(
            lang, LOT_DESC_DEFAULT["fr"]
        ).get(lot_key, "")
        normes = LOT_NORMES.get(lot_key, "")

        self._add_text(lot_label, size=Pt(9), bold=True)
        if lot_desc:
            self._add_text(lot_desc, size=Pt(8.5))

        projet = c.st_projet
        info_data = [
            (t("objet_projet"), projet.name if projet else "—"),
            (t("objet_adresse"), projet.adresse if projet and projet.adresse else "—"),
            (
                t("objet_mo"),
                projet.maitre_ouvrage if projet and projet.maitre_ouvrage else "—",
            ),
            (t("objet_permis"), projet.permis if projet and projet.permis else "—"),
            (t("objet_normes"), normes),
        ]
        table = self.doc.add_table(rows=len(info_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)
        for i, (lbl, val) in enumerate(info_data):
            _set_cell_width(table.cell(i, 0), 5)
            lr = table.cell(i, 0).paragraphs[0].add_run(lbl)
            lr.font.size = Pt(8.5)
            lr.font.color.rgb = DARK
            lr.bold = True
            vr = table.cell(i, 1).paragraphs[0].add_run(str(val))
            vr.font.size = Pt(8.5)
            vr.font.color.rgb = INK
        self._add_text(t("objet_declaration"), size=Pt(8.5))

    def _build_art_docs(self):
        t = self._t
        self._add_art_title(t("art_docs"))
        self._add_text(t("docs_intro"), size=Pt(8.5))
        lot_keys = _resolve_lot_keys(self.c.st_lot_type)
        lot_key = lot_keys[0] if lot_keys else ""
        normes = LOT_NORMES.get(lot_key, "")
        docs = st_t("docs_list", self.lang)
        if isinstance(docs, list):
            for i, d in enumerate(docs, 1):
                text = f"{d} : {normes}" if i == 7 and normes else d
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.4)
                r = p.add_run(f"{i}. {text}")
                r.font.size = Pt(8.5)
                r.font.color.rgb = INK
        self._add_text(t("docs_contradiction"), size=Pt(8.5))

    def _build_art_prix(self):
        t = self._t
        c = self.c
        lang = self.lang
        dev = self._dev()
        self._add_art_title(t("art_prix"))

        # 4.1 Montant
        self._add_sub_title(t("prix_montant"))

        type_prix_keys = _resolve_type_prix_keys(c.st_type_prix)
        type_prix_lbl = " / ".join(
            TYPE_PRIX_LABELS.get(lang, TYPE_PRIX_LABELS["fr"]).get(k, k)
            for k in type_prix_keys
        )
        self._add_text(f"{t('prix_type')} {type_prix_lbl}", size=Pt(8.5), bold=True)

        # Financial box as table
        fin_rows = [
            (t("prix_ht"), _fmt_amt(self._ht, dev), False),
            (
                t("prix_tva").format(tva=f"{self._tva_pct:g}"),
                _fmt_amt(self._tva_amt, dev),
                False,
            ),
            (t("prix_ttc"), _fmt_amt(self._ttc, dev), True),
        ]
        ftable = self.doc.add_table(rows=len(fin_rows), cols=2)
        ftable.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(ftable)
        for i, (lbl, val, is_grand) in enumerate(fin_rows):
            _cell_bg(ftable.cell(i, 0), CREAM_HEX)
            _cell_bg(ftable.cell(i, 1), CREAM_HEX)
            lr = ftable.cell(i, 0).paragraphs[0].add_run(lbl)
            lr.font.size = Pt(9 if is_grand else 8.5)
            lr.font.color.rgb = DARK
            lr.bold = is_grand
            vr = ftable.cell(i, 1).paragraphs[0].add_run(val)
            vr.font.size = Pt(11 if is_grand else 8.5)
            vr.font.color.rgb = GOLD if is_grand else DARK
            vr.bold = True
            ftable.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # "Arrêté à la somme de" line
        arrete_p = self.doc.add_paragraph()
        arrete_r = arrete_p.add_run(
            f"{t('prix_arrete')} {_num_to_words_mad(round(self._ttc), lang)} {dev}."
        )
        arrete_r.font.size = Pt(8)
        arrete_r.italic = True
        arrete_r.font.color.rgb = MUTED
        self._add_text(t("prix_comprend"), size=Pt(8.5))

        # 4.2 Modalités
        self._add_sub_title(t("prix_modalites"))
        delai_pay = c.st_delai_paiement or 30
        rib_info = ""
        if c.st_rib or c.st_banque:
            if self.lang == "fr":
                rib_info = f" sur le compte n\u00b0\u00a0{c.st_rib or '\u2014'} ouvert aupr\u00e8s de {c.st_banque or '\u2014'}"
            else:
                rib_info = f" on account n\u00b0\u00a0{c.st_rib or '\u2014'} held at {c.st_banque or '\u2014'}"
        self._add_text(t("prix_virement_intro").format(rib_info=rib_info), size=Pt(8.5))
        self._add_text(t("prix_delai_paiement").format(days=delai_pay), size=Pt(8.5))

        # Tranches
        tranches = c.st_tranches or []
        if tranches:
            tt = self.doc.add_table(rows=1, cols=len(tranches))
            tt.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(tt)
            for i, tr in enumerate(tranches):
                cell = tt.cell(0, i)
                _cell_bg(cell, CREAM_HEX)
                _cell_borders(
                    cell,
                    top=(BORDER_HEX, 2),
                    bottom=(BORDER_HEX, 2),
                    left=(BORDER_HEX, 2),
                    right=(BORDER_HEX, 2),
                )
                lbl = tr.get("label", f"Tranche {i+1}")
                pct = float(tr.get("pourcentage", 0))
                amt = self._ht * pct / 100
                lp = cell.paragraphs[0]
                lr2 = lp.add_run(lbl)
                lr2.font.size = Pt(7)
                lr2.font.color.rgb = MUTED
                lr2.bold = True
                lr2.font.all_caps = True
                ap = cell.add_paragraph()
                ar = ap.add_run(_fmt_amt(amt, dev))
                ar.font.size = Pt(12)
                ar.font.color.rgb = DARK
                ar.bold = True
                dp = cell.add_paragraph()
                dr = dp.add_run(f"{pct:g}%")
                dr.font.size = Pt(7)
                dr.font.color.rgb = MUTED

        # 4.3 Avance
        avance_pct = float(c.st_avance or 0)
        if avance_pct > 0:
            avance_amt = self._ht * avance_pct / 100
            self._add_sub_title(t("prix_avance"))
            self._add_text(
                t("prix_avance_text").format(
                    pct=f"{avance_pct:g}", amount=_fmt_amt(avance_amt, dev)
                ),
                size=Pt(8.5),
            )

        # 4.4 Retenue
        ret_pct = float(c.st_retenue_garantie or 0)
        garantie_ret_mois = c.st_garantie_mois or 12
        if ret_pct > 0:
            self._add_sub_title(t("prix_retenue"))
            self._add_text(
                t("prix_retenue_text").format(
                    pct=f"{ret_pct:g}", months=garantie_ret_mois
                ),
                size=Pt(8.5),
            )

        # 4.5 Supplémentaires
        self._add_sub_title(t("prix_supplementaires"))
        self._add_text(t("prix_supplementaires_text"), size=Pt(8.5))

    def _build_art_delais(self):
        t = self._t
        c = self.c
        lang = self.lang
        self._add_art_title(t("art_delais"))

        val = c.st_delai_val or "—"
        unit_key = c.st_delai_unit or "mois"
        unit_lbl = DELAI_UNIT_LABELS.get(lang, DELAI_UNIT_LABELS["fr"]).get(
            unit_key, unit_key
        )
        pen_taux = c.st_penalite_taux or 1
        plafond = c.st_plafond_penalite or 10

        self._add_sub_title(t("delais_global"))
        self._add_text(
            t("delais_global_text").format(val=val, unit=unit_lbl), size=Pt(8.5)
        )
        self._add_sub_title(t("delais_planning"))
        self._add_text(t("delais_planning_text"), size=Pt(8.5))
        self._add_sub_title(t("delais_penalites"))
        self._add_text(
            t("delais_penalites_text").format(
                taux=f"{float(pen_taux):g}", plafond=f"{float(plafond):g}"
            ),
            size=Pt(8.5),
        )
        self._add_sub_title(t("delais_resiliation"))
        self._add_text(t("delais_resiliation_text"), size=Pt(8.5))

    def _build_art_obligations_st(self):
        t = self._t
        lang = self.lang
        lot_keys = _resolve_lot_keys(self.c.st_lot_type)
        lot_key = lot_keys[0] if lot_keys else ""
        self._add_art_title(t("art_obligations_st"))

        self._add_sub_title(t("oblig_generales"))
        gen_items = st_t("oblig_generales_list", lang)
        if isinstance(gen_items, list):
            for item in gen_items:
                self._add_bullet(item)

        specific = LOT_OBLIGATIONS.get(lang, LOT_OBLIGATIONS["fr"]).get(lot_key, [])
        if specific:
            self._add_sub_title(t("oblig_specifiques"))
            for item in specific:
                self._add_bullet(item)

    def _build_art_obligations_ep(self):
        t = self._t
        self._add_art_title(t("art_obligations_ep"))
        items = st_t("oblig_ep_list", self.lang)
        if isinstance(items, list):
            for item in items:
                self._add_bullet(item)

    def _build_art_assurances(self):
        t = self._t
        lang = self.lang
        lot_keys = _resolve_lot_keys(self.c.st_lot_type)
        lot_key = lot_keys[0] if lot_keys else ""
        self._add_art_title(t("art_assurances"))

        self._add_text(t("assurances_intro"), size=Pt(8.5))
        lot_assur = LOT_ASSURANCES.get(lang, LOT_ASSURANCES["fr"]).get(lot_key, "")
        if lot_assur:
            self._add_text(lot_assur, size=Pt(8.5), bold=True)

        rules = st_t("assurances_rules", lang)
        if isinstance(rules, list):
            for rule in rules:
                self._add_bullet(rule)

        actives = self.c.st_clauses_actives or []
        if "tTRC" in actives:
            self._add_text(t("assurances_trc"), size=Pt(8.5))

    def _build_art_reception(self):
        t = self._t
        lang = self.lang
        c = self.c
        lot_keys = _resolve_lot_keys(c.st_lot_type)
        lot_key = lot_keys[0] if lot_keys else ""
        self._add_art_title(t("art_reception"))

        reception_text = LOT_RECEPTION.get(lang, LOT_RECEPTION["fr"]).get(lot_key, "")
        delai_res = c.st_delai_reserves or 30
        garantie_mois = c.st_garantie_mois or 12

        self._add_sub_title(t("reception_provisoire"))
        self._add_text(t("reception_provisoire_text"), size=Pt(8.5))
        if reception_text:
            self._add_text(reception_text, size=Pt(8.5))
        self._add_text(t("reception_pv_text"), size=Pt(8.5))
        self._add_sub_title(t("reception_reserves"))
        self._add_text(
            t("reception_reserves_text").format(days=delai_res), size=Pt(8.5)
        )
        self._add_sub_title(t("reception_definitive"))
        self._add_text(
            t("reception_definitive_text").format(months=garantie_mois), size=Pt(8.5)
        )
        self._add_sub_title(t("reception_decennale"))
        self._add_text(t("reception_decennale_text"), size=Pt(8.5))

    def _build_art_responsabilite(self):
        t = self._t
        self._add_art_title(t("art_responsabilite"))
        self._add_text(t("resp_resultat"), size=Pt(8.5))
        self._add_text(t("resp_recours"), size=Pt(8.5))
        self._add_text(t("resp_personnel"), size=Pt(8.5))
        self._add_text(
            t("resp_defaillance").format(days=self.c.st_delai_med or 8), size=Pt(8.5)
        )

    def _build_art_resiliation(self):
        t = self._t
        c = self.c
        delai_med = c.st_delai_med or 8
        self._add_art_title(t("art_resiliation"))
        self._add_sub_title(t("resil_faute"))
        self._add_text(t("resil_faute_intro"), size=Pt(8.5))
        items = st_t("resil_faute_list", self.lang)
        if isinstance(items, list):
            for item in items:
                self._add_bullet(item)
        self._add_text(t("resil_faute_effet").format(days=delai_med), size=Pt(8.5))
        self._add_sub_title(t("resil_convenance"))
        self._add_text(t("resil_convenance_text"), size=Pt(8.5))
        self._add_sub_title(t("resil_consequences"))
        self._add_text(t("resil_consequences_text"), size=Pt(8.5))

    def _build_art_hse(self):
        t = self._t
        self._add_art_title(t("art_hse"))
        items = st_t("hse_list", self.lang)
        if isinstance(items, list):
            for item in items:
                self._add_text(item, size=Pt(8.5))

    def _build_optional_clauses(self):
        actives = self.c.st_clauses_actives or []
        projet = self.c.st_projet
        mo = (
            projet.maitre_ouvrage
            if projet and projet.maitre_ouvrage
            else ("le Maître d'Ouvrage" if self.fr else "the Employer")
        )
        clause_confid_amount = _fmt_amt(self._ht * 0.2, self._dev())

        if "tConfid" in actives:
            self._add_art_title(self._t("clause_confid"))
            n = self._art_num
            if self.fr:
                self._add_text(
                    f"{n}.1. Le Sous-Traitant s'engage à traiter comme strictement confidentielles toutes les informations techniques, commerciales et financières dont il pourrait avoir connaissance dans le cadre de l'exécution du présent contrat.",
                    size=Pt(8.5),
                )
                self._add_text(
                    f"{n}.2. Cette obligation de confidentialité s'étend au personnel du Sous-Traitant et à tout tiers auquel il pourrait faire appel. Elle survit à l'extinction du contrat pour une durée de 3 ans.",
                    size=Pt(8.5),
                )
                self._add_text(
                    f"{n}.3. Toute violation de cette clause expose le Sous-Traitant au paiement d'une indemnité forfaitaire de {clause_confid_amount}, sans préjudice du droit de l'Entrepreneur Principal à demander réparation du préjudice réel subi.",
                    size=Pt(8.5),
                )
            else:
                self._add_text(
                    f"{n}.1. The Subcontractor undertakes to treat as strictly confidential all technical, commercial and financial information that may come to its knowledge in the performance of this contract.",
                    size=Pt(8.5),
                )
                self._add_text(
                    f"{n}.2. This confidentiality obligation extends to the Subcontractor's personnel and to any third party it may engage. It survives termination of the contract for a period of 3 years.",
                    size=Pt(8.5),
                )
                self._add_text(
                    f"{n}.3. Any breach of this clause exposes the Subcontractor to a lump-sum indemnity of {clause_confid_amount}, without prejudice to the Principal Contractor's right to claim compensation for the actual loss suffered.",
                    size=Pt(8.5),
                )

        if "tNonConc" in actives:
            self._add_art_title(self._t("clause_non_conc"))
            if self.fr:
                self._add_text(
                    f"Le Sous-Traitant s'interdit, pendant toute la durée du contrat et pendant une période de 12 mois suivant son terme, de contracter directement avec le Maître d'Ouvrage {mo} ou avec tout acquéreur des lots du projet, pour des prestations similaires à celles objet du présent contrat, dans un rayon de 20 km autour du projet.",
                    size=Pt(8.5),
                )
                self._add_text(
                    "Toute violation expose le Sous-Traitant au paiement d'une indemnité forfaitaire de 30% du montant HT du présent contrat.",
                    size=Pt(8.5),
                )
            else:
                self._add_text(
                    f"The Subcontractor undertakes, throughout the term of the contract and for a period of 12 months after its expiry, not to contract directly with the Employer {mo} or with any purchaser of lots in the project for services similar to those covered by this contract, within a radius of 20 km around the project.",
                    size=Pt(8.5),
                )
                self._add_text(
                    "Any breach exposes the Subcontractor to a lump-sum indemnity equal to 30% of the pre-tax amount of this contract.",
                    size=Pt(8.5),
                )

        if "tNonDeb" in actives:
            self._add_art_title(self._t("clause_non_deb"))
            if self.fr:
                self._add_text(
                    "Chacune des Parties s'interdit de recruter ou de tenter de recruter, directement ou indirectement, tout salarié ou collaborateur de l'autre Partie ayant participé à l'exécution du présent contrat, pendant toute la durée du contrat et pendant une période de 12 mois suivant son terme.",
                    size=Pt(8.5),
                )
                self._add_text(
                    "Toute violation entraînera le paiement d'une indemnité forfaitaire équivalente à 12 mois de rémunération brute du salarié concerné.",
                    size=Pt(8.5),
                )
            else:
                self._add_text(
                    "Each Party undertakes not to recruit or attempt to recruit, directly or indirectly, any employee or collaborator of the other Party who has participated in the performance of this contract, throughout the term of the contract and for a period of 12 months after its expiry.",
                    size=Pt(8.5),
                )
                self._add_text(
                    "Any breach shall result in a lump-sum indemnity equal to 12 months of the gross remuneration of the employee concerned.",
                    size=Pt(8.5),
                )

        if "tCascade" in actives:
            self._add_art_title(self._t("clause_cascade"))
            if self.fr:
                self._add_text(
                    "Le Sous-Traitant s'interdit formellement de sous-traiter tout ou partie des travaux objets du présent contrat à un tiers, sauf accord écrit et préalable de l'Entrepreneur Principal.",
                    size=Pt(8.5),
                )
                self._add_text(
                    "En cas de violation, l'Entrepreneur Principal pourra résilier immédiatement le contrat aux torts exclusifs du Sous-Traitant, sans mise en demeure préalable, et sans préjudice de dommages et intérêts.",
                    size=Pt(8.5),
                )
            else:
                self._add_text(
                    "The Subcontractor is formally prohibited from subcontracting all or part of the works covered by this contract to a third party unless it has obtained the Principal Contractor's prior written consent.",
                    size=Pt(8.5),
                )
                self._add_text(
                    "In the event of breach, the Principal Contractor may terminate the contract immediately at the Subcontractor's sole fault, without prior notice and without prejudice to damages.",
                    size=Pt(8.5),
                )

        if "tEnviro" in actives:
            self._add_art_title(self._t("clause_enviro"))
            if self.fr:
                self._add_text(
                    "Le Sous-Traitant s'engage à respecter la législation environnementale en vigueur au Maroc, notamment la loi n° 11-03 relative à la protection et à la mise en valeur de l'environnement. Il s'engage à :",
                    size=Pt(8.5),
                )
                for item in [
                    "Gérer ses déchets de chantier de manière responsable et les évacuer vers les décharges autorisées",
                    "Limiter les nuisances sonores et les émissions de poussière",
                    "Utiliser des matériaux respectueux de l'environnement dans la mesure du possible",
                    "Respecter les horaires de chantier pour limiter les nuisances au voisinage",
                ]:
                    self._add_bullet(item)
            else:
                self._add_text(
                    "The Subcontractor undertakes to comply with environmental legislation in force in Morocco, in particular Law No. 11-03 on environmental protection and enhancement. It undertakes to:",
                    size=Pt(8.5),
                )
                for item in [
                    "Manage site waste responsibly and remove it to authorised dumps",
                    "Limit noise nuisance and dust emissions",
                    "Use environmentally friendly materials wherever possible",
                    "Respect site working hours in order to limit disturbance to neighbouring properties",
                ]:
                    self._add_bullet(item)

        for toggle_key, title_key, text_key in [
            ("tPI", "clause_pi", "clause_pi_text"),
            ("tExclus", "clause_exclus", "clause_exclus_text"),
            ("tRevision", "clause_revision", "clause_revision_text"),
        ]:
            if toggle_key in actives:
                self._add_art_title(self._t(title_key))
                self._add_text(self._t(text_key), size=Pt(8.5))

    def _build_trailing_articles(self):
        t = self._t
        c = self.c
        actives = c.st_clauses_actives or []

        # Litiges
        self._add_art_title(t("art_litiges"))
        if "tMediat" in actives:
            self._add_text(t("litiges_mediation"), size=Pt(8.5))
            self._add_text(t("litiges_tribunal_mediation"), size=Pt(8.5))
        else:
            self._add_text(t("litiges_tribunal"), size=Pt(8.5))

        # Force Majeure
        self._add_art_title(t("art_force_majeure"))
        self._add_text(t("force_majeure_text"), size=Pt(8.5))

        # Dispositions
        self._add_art_title(t("art_dispositions"))
        items = st_t("dispositions_items", self.lang)
        if isinstance(items, list):
            for item in items:
                self._add_bullet(item)

    def _build_observations(self):
        obs = self.c.st_observations
        if not obs:
            return
        self._add_art_title(self._t("special_clauses_title"))
        self._add_text(obs, size=Pt(8.5))

    def _build_signatures(self):
        t = self._t
        c = self.c
        ep = self.ep
        date_str = _fmt_date(c.date_contrat)
        ville = c.ville_signature or "Tanger"
        ep_name = ep.name if ep else "CASA DI LUSSO SARL"
        st_name = c.st_name or "Sous-Traitant"

        self._add_empty(8)
        sep = self.doc.add_paragraph()
        _para_border_bottom(sep, BORDER_HEX, "4")
        sep.paragraph_format.space_after = Pt(6)

        info_p, info_r = self._add_text(
            f"{t('fait_a')} {ville}, {t('le')} {date_str}, {t('en_exemplaires')}",
            size=Pt(8.5),
            color=INK,
        )
        _para_bg(info_p, CREAM_HEX)

        self._add_empty(4)

        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        for i, (label, name, role) in enumerate(
            [
                (t("lbl_ep"), ep_name, ep.qualite_representant if ep else "Gérant"),
                (t("lbl_st"), st_name, c.st_qualite or ""),
            ]
        ):
            cell = table.cell(0, i)
            _cell_borders(
                cell,
                top=(BORDER_HEX, 2),
                bottom=(BORDER_HEX, 2),
                left=(BORDER_HEX, 2),
                right=(BORDER_HEX, 2),
            )

            lp = cell.paragraphs[0]
            lr = lp.add_run(label)
            lr.font.size = Pt(7)
            lr.font.color.rgb = GOLD
            lr.bold = True
            lr.font.all_caps = True

            np_ = cell.add_paragraph()
            nr = np_.add_run(t("lu_approuve"))
            nr.font.size = Pt(8)
            nr.font.color.rgb = MUTED
            nr.italic = True

            # signature line
            sp = cell.add_paragraph()
            sp.paragraph_format.space_before = Pt(30)
            _para_border_bottom(sp, "CCCCCC", "2")

            snp = cell.add_paragraph()
            snr = snp.add_run(name)
            snr.font.size = Pt(8.5)
            snr.font.color.rgb = DARK
            snr.bold = True

            srp = cell.add_paragraph()
            srr = srp.add_run(role)
            srr.font.size = Pt(7.5)
            srr.font.color.rgb = MUTED

        # Paraphes note
        par_p, par_r = self._add_text(
            t("paraphes"), size=Pt(7.5), color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER
        )
        par_r.italic = True

    def _build_annexe(self):
        actives = self.c.st_clauses_actives or []
        if "tAnnexe" not in actives:
            return
        t = self._t
        items = st_t("annexe_items", self.lang)
        if not isinstance(items, list):
            return

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _para_bg(p, GOLD_PALE_HEX)
        _para_border_left(p, GOLD_HEX, "12")
        tr = p.add_run(f"  {t('annexe_title')}")
        tr.font.size = Pt(8)
        tr.font.color.rgb = DARK
        tr.bold = True

        table = self.doc.add_table(rows=len(items) + 1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for j, hdr in enumerate(
            [t("annexe_col_no"), t("annexe_col_doc"), t("annexe_col_status")]
        ):
            cell = table.cell(0, j)
            _cell_bg(cell, DARK_HEX)
            r = cell.paragraphs[0].add_run(hdr)
            r.font.size = Pt(7)
            r.font.color.rgb = GOLD
            r.bold = True
            if j != 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_width(cell, 2)

        for i, item in enumerate(items, start=1):
            row_idx = i
            no_cell = table.cell(row_idx, 0)
            no_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            no_cell.paragraphs[0].add_run(str(i)).font.size = Pt(8.5)
            table.cell(row_idx, 1).paragraphs[0].add_run(item).font.size = Pt(8.5)
            status_cell = table.cell(row_idx, 2)
            status_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = status_cell.paragraphs[0].add_run(t("annexe_status_blank"))
            r.font.size = Pt(8.5)
            if row_idx % 2 == 0:
                for j in range(3):
                    _cell_bg(table.cell(row_idx, j), CREAM_HEX)

    def _build_footer(self):
        ep = self.ep
        ep_name = ep.name if ep else "CASA DI LUSSO SARL"
        rc = ep.rc if ep else ""
        now = datetime.now()
        gen_date = now.strftime("%d/%m/%Y")
        gen_time = now.strftime("%H:%M")

        self._add_empty(6)
        sep = self.doc.add_paragraph()
        _para_border_bottom(sep, BORDER_HEX, "2")
        sep.paragraph_format.space_after = Pt(4)

        self._add_text(
            f"{ep_name} — Entrepreneur Principal\n"
            f"RC {rc} · Tanger, {'Maroc' if self.fr else 'Morocco'}\n"
            f"Document généré le {gen_date} à {gen_time}",
            size=Pt(7),
            color=MUTED,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # ── main assembly ────────────────────────────────────────────────────────

    def _build(self):
        self._art_num = 0
        self._build_topstrip()
        self._build_header()
        self._build_title()
        self._build_parties()
        self._build_art_parties()
        self._build_art_objet()
        self._build_art_docs()
        self._build_art_prix()
        self._build_art_delais()
        self._build_art_obligations_st()
        self._build_art_obligations_ep()
        self._build_art_assurances()
        self._build_art_reception()
        self._build_art_responsabilite()
        self._build_art_resiliation()
        self._build_art_hse()
        self._build_optional_clauses()
        self._build_trailing_articles()
        self._build_observations()
        self._build_signatures()
        self._build_annexe()
        self._build_footer()

    def generate_response(self) -> HttpResponse:
        """Return an HttpResponse with the generated DOCX."""
        self._build()
        buf = BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        ref = (self.c.numero_contrat or "contract").replace("/", "-")
        filename = f"contrat_st_{self.c.id}_{ref}.docx"
        response = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response
