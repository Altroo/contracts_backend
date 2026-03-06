import re
from datetime import datetime, timedelta
from io import BytesIO

from django.http import HttpResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from .pdf import _fmt_date, _fmt_amt
from .i18n import QUALITE_LABELS
from .bl_i18n import (
    BL_COMPANY,
    FOURNITURES_LABELS,
    EAU_ELEC_LABELS,
    GARANTIE_UNITE_LABELS,
    GARANTIE_TYPE_LABELS,
    CLAUSE_RESILIATION_LABELS,
    bl_t,
)

NAVY = RGBColor(0x0A, 0x16, 0x28)
BLUE = RGBColor(0x2A, 0x7F, 0xFF)
GOLD = RGBColor(0xC8, 0xA9, 0x6E)
TEXT = RGBColor(0x1E, 0x2D, 0x4A)
GRAY = RGBColor(0x8A, 0x99, 0xB3)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
RED = RGBColor(0xEF, 0x44, 0x44)
NAVY_HEX = "0A1628"
BLUE_HEX = "2A7FFF"
GOLD_HEX = "C8A96E"
LIGHT_HEX = "F0F4FA"
BORDER_HEX = "DCE4F0"
BODY_HEX = "3A4E6E"


def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _para_bg(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _para_border_left(para, color: str = "2A7FFF", size: str = "12"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _para_border_bottom(para, color: str = "DCE4F0", size: str = "4"):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(val[1]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val[0])
            tcBorders.append(el)
        else:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_width(cell, cm_val):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm_val * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_row_height(row, pt_val):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(pt_val * 20)))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


def _run_shd(run, hex_color: str):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _garantie_text(contract, lang: str = "fr") -> str:
    nb = int(contract.garantie_nb or 0)
    if nb == 0:
        return bl_t("sans_garantie", lang)
    unite = contract.garantie_unite or "ans"
    labels = GARANTIE_UNITE_LABELS[lang]
    if unite == "mois":
        return f"{nb} {labels['mois']}"
    return f"{nb} {labels['ans_p'] if nb > 1 else labels['ans_s']}"


def _solde_pct(contract) -> float:
    return 100 - float(contract.acompte or 0) - float(contract.tranche2 or 0)


def _parse_strong(html_text: str) -> list[tuple[str, bool]]:
    """Parse an HTML string with <strong> tags into (text, bold) tuples."""
    parts: list[tuple[str, bool]] = []
    for i, segment in enumerate(re.split(r"<strong>(.*?)</strong>", html_text)):
        if not segment:
            continue
        parts.append((segment, i % 2 == 1))
    return parts


class BluelineDOCGenerator:
    """Generate a python-docx DOCX for a Blueline Works contract."""

    def __init__(self, contract, language: str = "fr"):
        self.c = contract
        self.lang = language
        self.fr = language == "fr"
        self.doc = Document()
        self._configure_page()

    def _t(self, key: str) -> str:
        return bl_t(key, self.lang)

    def _configure_page(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        normal = self.doc.styles["Normal"]
        normal.font.name = "DM Sans"
        normal.font.size = Pt(9)
        normal.font.color.rgb = TEXT
        normal.paragraph_format.space_after = Pt(2)
        normal.paragraph_format.line_spacing = 1.65

    def _add_empty(self, pt=2):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(pt)
        return p

    def _add_text(
        self, text, size=Pt(9), bold=False, color=TEXT, align=None, space_after=Pt(3)
    ):
        p = self.doc.add_paragraph()
        if align:
            p.alignment = align
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.font.size = size
        run.bold = bold
        run.font.color.rgb = color
        return p, run

    def _add_section_title(self, title: str):
        """Add Blueline section title with left blue border."""
        p, run = self._add_text(
            title, size=Pt(9), bold=True, color=NAVY, space_after=Pt(6)
        )
        _para_bg(p, LIGHT_HEX)
        _para_border_left(p, BLUE_HEX, "12")
        p.paragraph_format.space_before = Pt(10)
        return p

    def _add_article_title(self, title: str):
        """Add article title bar (gray background)."""
        p, run = self._add_text(
            title, size=Pt(8), bold=True, color=NAVY, space_after=Pt(0)
        )
        _para_bg(p, LIGHT_HEX)
        _para_border_bottom(p, BORDER_HEX, "4")
        return p

    def _add_article_body(self, text: str):
        """Add article body text."""
        p, run = self._add_text(
            text, size=Pt(8.5), color=RGBColor(0x3A, 0x4E, 0x6E), space_after=Pt(6)
        )
        return p

    def _add_article_body_multi(self, parts):
        """Add article body with bold/normal alternating parts.
        parts: list of (text, bold) tuples.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(2)
        for text, bold in parts:
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x3A, 0x4E, 0x6E)
            run.bold = bold
        return p

    @property
    def _ht(self) -> float:
        try:
            return float(self.c.montant_ht or 0)
        except (TypeError, ValueError):
            return 0.0

    @property
    def _tva_pct(self) -> float:
        try:
            return float(self.c.tva if self.c.tva is not None else 20)
        except (TypeError, ValueError):
            return 20.0

    @property
    def _tva_amt(self) -> float:
        return self._ht * self._tva_pct / 100

    @property
    def _ttc(self) -> float:
        return self._ht + self._tva_amt

    @property
    def _acompte_pct(self) -> float:
        return float(self.c.acompte if self.c.acompte is not None else 30)

    @property
    def _tranche2_pct(self) -> float:
        return float(self.c.tranche2 or 0)

    @property
    def _solde_pct_val(self) -> float:
        return _solde_pct(self.c)

    def _dev(self) -> str:
        return self.c.devise or "MAD"

    def _amt(self, n) -> str:
        return _fmt_amt(n, self._dev())

    def _date_fin(self) -> str:
        if self.c.date_debut and self.c.duree_estimee:
            try:
                d = self.c.date_debut + timedelta(days=int(self.c.duree_estimee))
                return _fmt_date(d)
            except (TypeError, ValueError):
                pass
        return "—"

    def _build_topstrip(self):
        """Add colored top strip via a thin table."""
        table = self.doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)
        colors = [NAVY_HEX, BLUE_HEX, GOLD_HEX]
        for i, color in enumerate(colors):
            cell = table.cell(0, i)
            _cell_bg(cell, color)
            cell.text = ""
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
        _set_row_height(table.rows[0], 3)
        self._add_empty(8)

    def _build_header(self):
        co = BL_COMPANY
        tagline = co["tagline_fr"] if self.fr else co["tagline_en"]

        # Header table: logo left | ref right
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)
        _set_cell_width(table.cell(0, 0), 10.5)
        _set_cell_width(table.cell(0, 1), 7)

        # Left cell
        cell_left = table.cell(0, 0)
        p = cell_left.paragraphs[0]
        run = p.add_run(co["name"])
        run.font.size = Pt(20)
        run.bold = True
        run.font.color.rgb = NAVY
        run.font.name = "Playfair Display"

        p2 = cell_left.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(tagline)
        run2.font.size = Pt(7)
        run2.font.color.rgb = GOLD
        run2.font.all_caps = True

        p3 = cell_left.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        info_lines = [
            f"📞 {co['phone']}",
            f"✉ {co['email']}",
            f"📍 {co['address']}",
            f"🏛 ICE : {co['ice']}",
        ]
        for line in info_lines:
            run3 = p3.add_run(line + "\n")
            run3.font.size = Pt(7)
            run3.font.color.rgb = GRAY

        # Right cell
        cell_right = table.cell(0, 1)
        pr = cell_right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = pr.add_run(f"  N° {self.c.numero_contrat or ''}  ")
        ref_run.font.size = Pt(9)
        ref_run.bold = True
        ref_run.font.color.rgb = GOLD
        _run_shd(ref_run, NAVY_HEX)

        pd = cell_right.add_paragraph()
        pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pd.paragraph_format.space_after = Pt(2)
        drun = pd.add_run(f"{self._t('date_sig')} : {_fmt_date(self.c.date_contrat)}")
        drun.font.size = Pt(8)
        drun.font.color.rgb = GRAY

        pb = cell_right.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        brun = pb.add_run(f"  {self._t('badge')}  ")
        brun.font.size = Pt(7.5)
        brun.bold = True
        brun.font.color.rgb = BLUE
        _run_shd(brun, "EBF2FF")

        # Separator
        sep = self.doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(4)
        sep.paragraph_format.space_after = Pt(8)
        _para_border_bottom(sep, BORDER_HEX, "4")

    def _build_title(self):
        p, run = self._add_text(
            self._t("contrat_title"),
            size=Pt(16),
            bold=True,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(2),
        )
        run.font.name = "Playfair Display"

        self._add_text(
            self._t("contrat_subtitle"),
            size=Pt(8),
            color=GRAY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(8),
        )

    def _build_parties(self):
        self._add_section_title(self._t("parties_title"))
        co = BL_COMPANY
        c = self.c

        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        # Prestataire cell (navy background)
        cell_p = table.cell(0, 0)
        _cell_bg(cell_p, NAVY_HEX)
        _set_cell_width(cell_p, 8.7)

        pp = cell_p.paragraphs[0]
        pp.paragraph_format.space_after = Pt(3)
        lr = pp.add_run(self._t("prestataire_label"))
        lr.font.size = Pt(7)
        lr.bold = True
        lr.font.color.rgb = GOLD
        lr.font.all_caps = True

        pn = cell_p.add_paragraph()
        pn.paragraph_format.space_after = Pt(2)
        nr = pn.add_run(co["name"])
        nr.font.size = Pt(12)
        nr.bold = True
        nr.font.color.rgb = WHITE

        pi = cell_p.add_paragraph()
        pi.paragraph_format.space_after = Pt(0)
        desc = self._t("prestataire_desc")
        info = f"{desc}\n{self._t('tel')} : {co['phone']}\n{self._t('email')} : {co['email']}\n{self._t('ice')} : {co['ice']}\n{self._t('adresse')} : {co['address']}"
        ir = pi.add_run(info)
        ir.font.size = Pt(8)
        ir.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Client cell (light blue border)
        cell_c = table.cell(0, 1)
        _set_cell_width(cell_c, 8.7)
        _cell_borders(
            cell_c,
            top=(BLUE_HEX, 6),
            bottom=(BLUE_HEX, 6),
            left=(BLUE_HEX, 6),
            right=(BLUE_HEX, 6),
        )

        cp = cell_c.paragraphs[0]
        cp.paragraph_format.space_after = Pt(3)
        clr = cp.add_run(self._t("client_label"))
        clr.font.size = Pt(7)
        clr.bold = True
        clr.font.color.rgb = BLUE
        clr.font.all_caps = True

        cn = cell_c.add_paragraph()
        cn.paragraph_format.space_after = Pt(2)
        cnr = cn.add_run(c.client_nom or "")
        cnr.font.size = Pt(12)
        cnr.bold = True
        cnr.font.color.rgb = NAVY

        ci = cell_c.add_paragraph()
        ci.paragraph_format.space_after = Pt(0)
        lines = []
        if c.client_adresse:
            addr = c.client_adresse
            if c.client_ville:
                addr += f" — {c.client_ville}"
            if c.client_cp:
                addr += f" {c.client_cp}"
            lines.append(f"{self._t('adresse')} : {addr}")
        if c.client_tel:
            lines.append(f"{self._t('tel')} : {c.client_tel}")
        if c.client_email:
            lines.append(f"{self._t('email')} : {c.client_email}")
        if c.client_cin:
            lines.append(f"{self._t('cin_ice')} : {c.client_cin}")
        qualite_display = QUALITE_LABELS[self.lang].get(
            c.client_qualite or "",
            c.client_qualite or ("Personne Physique" if self.fr else "Individual"),
        )
        lines.append(f"{self._t('qualite')} : {qualite_display}")
        cir = ci.add_run("\n".join(lines))
        cir.font.size = Pt(8)
        cir.font.color.rgb = TEXT

        self._add_empty(4)

    def _build_chantier(self):
        self._add_section_title(self._t("chantier_title"))
        c = self.c

        _date_fin_val = self._date_fin()
        info_items = [
            (
                self._t("adresse_chantier"),
                (c.adresse_travaux or "")
                + (f" — {c.chantier_etage}" if c.chantier_etage else ""),
            ),
            (self._t("ville"), c.chantier_ville or ""),
            (self._t("type_bien"), c.get_type_bien_display() if c.type_bien else ""),
            (self._t("surface_totale"), f"{c.surface} m²" if c.surface else ""),
            (self._t("date_debut"), _fmt_date(c.date_debut) if c.date_debut else ""),
            (
                self._t("fin_estimee"),
                (
                    (
                        _date_fin_val
                        + (
                            f" ({c.duree_estimee} {self._t('jours_ouvrables')})"
                            if c.duree_estimee
                            else ""
                        )
                    )
                    if _date_fin_val != "—"
                    else ""
                ),
            ),
        ]
        # Filter out entries with no value
        info_items = [(l, v) for l, v in info_items if v and v != "—"]

        num_rows = max(1, (len(info_items) + 2) // 3)
        table = self.doc.add_table(rows=num_rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        for idx, (label, value) in enumerate(info_items):
            row_idx = idx // 3
            col_idx = idx % 3
            cell = table.cell(row_idx, col_idx)
            _cell_bg(cell, "F8FAFE")
            _cell_borders(
                cell,
                top=(BORDER_HEX, 2),
                bottom=(BORDER_HEX, 2),
                left=(BORDER_HEX, 2),
                right=(BORDER_HEX, 2),
            )

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            lr = p.add_run(label)
            lr.font.size = Pt(7)
            lr.bold = True
            lr.font.color.rgb = GRAY
            lr.font.all_caps = True

            pv = cell.add_paragraph()
            pv.paragraph_format.space_before = Pt(0)
            vr = pv.add_run(str(value))
            vr.font.size = Pt(8.5)
            vr.bold = True
            vr.font.color.rgb = NAVY

        # Description des travaux
        if c.description_travaux:
            p_desc = self.doc.add_paragraph()
            p_desc.paragraph_format.space_before = Pt(4)
            p_desc.paragraph_format.space_after = Pt(2)
            _para_bg(p_desc, "F0F6FF")
            lbl_run = p_desc.add_run(self._t("description_travaux_label") + "\n")
            lbl_run.font.size = Pt(8.5)
            lbl_run.bold = True
            lbl_run.font.color.rgb = NAVY
            val_run = p_desc.add_run(c.description_travaux)
            val_run.font.size = Pt(8.5)
            val_run.font.color.rgb = RGBColor(0x3A, 0x4E, 0x6E)

        self._add_empty(4)

    def _build_prestations(self):
        self._add_section_title(self._t("prestations_title"))
        c = self.c
        prestations = c.prestations or []

        headers = [
            self._t("th_designation"),
            self._t("th_qty"),
            self._t("th_pu"),
            self._t("th_tva"),
            self._t("th_montant"),
        ]
        num_rows = max(len(prestations), 1) + 1  # +1 for header
        table = self.doc.add_table(rows=num_rows, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        # Header row
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            _cell_bg(cell, NAVY_HEX)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(header)
            r.font.size = Pt(7)
            r.bold = True
            r.font.color.rgb = WHITE
            if i == 4:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if not prestations:
            cell = table.cell(1, 0)
            p = cell.paragraphs[0]
            r = p.add_run(self._t("aucune_prestation"))
            r.font.size = Pt(8.5)
            r.font.color.rgb = GRAY
            r.italic = True
        else:
            for row_idx, prest in enumerate(prestations, 1):
                nom = prest.get("nom", "")
                desc = prest.get("description", prest.get("desc", ""))
                qte = prest.get("quantite", prest.get("qte", 0))
                unite = prest.get("unite", "")
                pu = float(prest.get("prix_unitaire", prest.get("pu", 0)))
                total = float(qte) * pu

                # Even row background
                if row_idx % 2 == 0:
                    for ci in range(5):
                        _cell_bg(table.cell(row_idx, ci), "FAFBFE")

                # Designation
                cell0 = table.cell(row_idx, 0)
                p0 = cell0.paragraphs[0]
                r0 = p0.add_run(f"{row_idx}. {nom}")
                r0.font.size = Pt(8.5)
                r0.bold = True
                r0.font.color.rgb = NAVY
                if desc:
                    pd = cell0.add_paragraph()
                    pd.paragraph_format.space_before = Pt(0)
                    rd = pd.add_run(desc)
                    rd.font.size = Pt(7)
                    rd.font.color.rgb = GRAY

                # Qty
                cell1 = table.cell(row_idx, 1)
                p1 = cell1.paragraphs[0]
                r1 = p1.add_run(f"{qte} {unite}")
                r1.font.size = Pt(8.5)

                # PU
                cell2 = table.cell(row_idx, 2)
                p2 = cell2.paragraphs[0]
                r2 = p2.add_run(self._amt(pu))
                r2.font.size = Pt(8.5)

                # TVA
                cell3 = table.cell(row_idx, 3)
                p3 = cell3.paragraphs[0]
                r3 = p3.add_run(f"{int(self._tva_pct)}%")
                r3.font.size = Pt(8.5)

                # Total
                cell4 = table.cell(row_idx, 4)
                p4 = cell4.paragraphs[0]
                p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r4 = p4.add_run(self._amt(total))
                r4.font.size = Pt(8.5)
                r4.bold = True
                r4.font.color.rgb = NAVY

                # Bottom border
                for ci in range(5):
                    _cell_borders(table.cell(row_idx, ci), bottom=(BORDER_HEX, 2))

        # Totals
        self._add_empty(4)
        for label, value, is_grand in [
            (self._t("subtotal_ht"), self._amt(self._ht), False),
            (
                f"{'TVA' if self.fr else 'VAT'} ({self._tva_pct:g}%)",
                self._amt(self._tva_amt),
                False,
            ),
            (self._t("total_ttc"), self._amt(self._ttc), True),
        ]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(2)
            lr = p.add_run(f"{label}   ")
            lr.font.size = Pt(9 if is_grand else 8.5)
            lr.font.color.rgb = NAVY if is_grand else TEXT
            lr.bold = is_grand
            vr = p.add_run(value)
            vr.font.size = Pt(11 if is_grand else 8.5)
            vr.bold = True
            vr.font.color.rgb = BLUE if is_grand else NAVY
            if is_grand:
                _para_border_bottom(p, NAVY_HEX, "6")

    def _build_payment(self):
        self._add_section_title(self._t("paiement_title"))
        ttc = self._ttc
        ap, t2p, sp = self._acompte_pct, self._tranche2_pct, self._solde_pct_val

        schedule = [
            (
                self._t("acompte_label"),
                ap,
                ttc * ap / 100,
                self._t("acompte_detail"),
                GOLD_HEX,
            ),
        ]
        if t2p > 0:
            schedule.append(
                (
                    self._t("tranche2_label"),
                    t2p,
                    ttc * t2p / 100,
                    self._t("tranche2_detail"),
                    BLUE_HEX,
                )
            )
        schedule.append(
            (
                self._t("solde_label"),
                sp,
                ttc * sp / 100,
                self._t("solde_detail"),
                "22C55E",
            )
        )

        table = self.doc.add_table(rows=1, cols=len(schedule))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        for i, (label, pct, amount, detail, accent) in enumerate(schedule):
            cell = table.cell(0, i)
            _cell_borders(
                cell,
                top=(accent, 6),
                bottom=(accent, 6),
                left=(accent, 6),
                right=(accent, 6),
            )

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            lr = p.add_run(f"{label} ({pct:.0f}%)")
            lr.font.size = Pt(7)
            lr.bold = True
            lr.font.color.rgb = GRAY
            lr.font.all_caps = True

            pa = cell.add_paragraph()
            pa.paragraph_format.space_after = Pt(2)
            ar = pa.add_run(self._amt(amount))
            ar.font.size = Pt(14)
            ar.bold = True
            ar.font.color.rgb = NAVY

            pd = cell.add_paragraph()
            pd.paragraph_format.space_before = Pt(0)
            dr = pd.add_run(detail)
            dr.font.size = Pt(7)
            dr.font.color.rgb = GRAY

        self._add_empty(4)

        # Mode de paiement
        mode = self.c.mode_paiement_texte or ""
        if mode:
            parts = [(f"{self._t('mode_paiement')} : ", True), (mode, False)]
            if self.c.rib:
                parts.append((f"\n{self._t('coord_bancaires')} : ", True))
                parts.append((self.c.rib, False))
            self._add_article_body_multi(parts)

    def _build_articles(self):
        self._add_section_title(self._t("articles_title"))
        c = self.c
        co_name = BL_COMPANY["name"]
        client = c.client_nom or ("le Client" if self.fr else "the Client")

        g_nb = int(c.garantie_nb or 0)
        g_text = _garantie_text(c, self.lang)
        g_type_val = c.garantie_type or "defauts"
        g_type_label = GARANTIE_TYPE_LABELS.get(
            self.lang, GARANTIE_TYPE_LABELS["fr"]
        ).get(g_type_val, GARANTIE_TYPE_LABELS["fr"].get("defauts", g_type_val))
        sp = self._solde_pct_val
        mont_solde = self._ttc * sp / 100

        fournitures_val = c.fournitures or "non_incluses"
        fournitures_text = FOURNITURES_LABELS.get(
            self.lang, FOURNITURES_LABELS["fr"]
        ).get(fournitures_val, FOURNITURES_LABELS["fr"]["non_incluses"])
        eau_val = c.eau_electricite or "client"
        eau_text = EAU_ELEC_LABELS.get(self.lang, EAU_ELEC_LABELS["fr"]).get(
            eau_val, EAU_ELEC_LABELS["fr"]["client"]
        )
        if c.penalite_retard:
            _pv = f"{c.penalite_retard:g}"
            if self.fr:
                _pv = _pv.replace(".", ",")
            penalites = f"{_pv}% par mois" if self.fr else f"{_pv}% per month"
        else:
            penalites = "1,5% par mois" if self.fr else "1.5% per month"
        resil_val = c.clause_resiliation or "30j"
        resil_text_raw = CLAUSE_RESILIATION_LABELS.get(
            self.lang, CLAUSE_RESILIATION_LABELS["fr"]
        ).get(resil_val, CLAUSE_RESILIATION_LABELS["fr"]["30j"])
        tribunal = c.tribunal or "Tribunal de Commerce de Casablanca"
        excl_garantie = c.exclusions_garantie or (
            "Usure normale, mauvaise utilisation, modifications par des tiers."
            if self.fr
            else "Normal wear, misuse, modifications by third parties."
        )
        date_debut_str = _fmt_date(c.date_debut)
        date_fin_str = self._date_fin()
        duree_parts_fr = []
        duree_parts_en = []
        if c.duree_estimee:
            duree_parts_fr = [
                (" (durée estimée : ", False),
                (f"{c.duree_estimee} {self._t('jours_ouvrables')}", True),
                (")", False),
            ]
            duree_parts_en = [
                (" (estimated duration: ", False),
                (f"{c.duree_estimee} {self._t('jours_ouvrables')}", True),
                (")", False),
            ]

        # Art 1
        self._add_article_title(self._t("art1_title"))
        if self.fr:
            self._add_article_body_multi(
                [
                    ("Le présent contrat a pour objet la réalisation par ", False),
                    (co_name, True),
                    (
                        ' (ci-après le "Prestataire") des travaux de pose de carrelage, marbre, '
                        "finitions et travaux connexes décrits dans le tableau des prestations, "
                        "au bénéfice de ",
                        False,
                    ),
                    (client, True),
                    (
                        ' (ci-après le "Client"), sur le chantier situé à l\'adresse indiquée. '
                        "Le Prestataire s'engage à exécuter ces travaux dans les règles de l'art, "
                        "avec tout le soin et le professionnalisme requis.",
                        False,
                    ),
                ]
            )
        else:
            self._add_article_body_multi(
                [
                    ("The purpose of this contract is the execution by ", False),
                    (co_name, True),
                    (
                        ' (hereinafter "The Service Provider") of tiling, marble, finishing '
                        "and related works described in the services table, for the benefit of ",
                        False,
                    ),
                    (client, True),
                    (
                        ' (hereinafter "The Client"), at the site located at the specified '
                        "address. The Service Provider undertakes to execute these works "
                        "in compliance with professional standards, with all due care and professionalism.",
                        False,
                    ),
                ]
            )

        # Art 2
        self._add_article_title(self._t("art2_title"))
        parts2 = [
            ("Fournitures : " if self.fr else "Supplies: ", True),
            (f"{fournitures_text}.\n", False),
        ]
        if c.materiaux_detail:
            mat_label = (
                "Matériaux à fournir par le client : "
                if self.fr
                else "Materials to be supplied by the client: "
            )
            parts2.append((mat_label, True))
            parts2.append((f"{c.materiaux_detail}.\n", False))
        parts2.append(
            ("Eau & Électricité : " if self.fr else "Water & Electricity: ", True)
        )
        parts2.append((f"{eau_text}\n", False))
        if self.fr:
            parts2.append(
                (
                    "Les matériaux fournis par le Client doivent être disponibles sur chantier au "
                    "minimum ",
                    False,
                )
            )
            parts2.append(("24 heures avant", True))
            parts2.append(
                (
                    " le début des travaux concernés. Tout retard de "
                    "livraison des matériaux imputable au Client pourra entraîner un décalage du "
                    "planning et/ou une facturation complémentaire.",
                    False,
                )
            )
        else:
            parts2.append(
                (
                    "Materials supplied by the Client must be available on site at least ",
                    False,
                )
            )
            parts2.append(("24 hours before", True))
            parts2.append(
                (
                    " the start of the relevant works. Any delay in material delivery "
                    "attributable to the Client may result in schedule adjustments and/or "
                    "additional charges.",
                    False,
                )
            )
        self._add_article_body_multi(parts2)

        # Art 3
        self._add_article_title(self._t("art3_title"))
        if self.fr:
            art3_parts = (
                [
                    ("Les travaux débuteront le ", False),
                    (date_debut_str, True),
                    (" et se termineront le ", False),
                    (date_fin_str, True),
                ]
                + duree_parts_fr
                + [
                    (
                        ". Ces délais sont donnés de bonne foi à titre indicatif et "
                        "peuvent être prolongés en cas de : (i) force majeure, (ii) retard de "
                        "livraison des matériaux non imputable au Prestataire, (iii) mauvaises "
                        "conditions climatiques, (iv) modifications demandées par le Client en "
                        "cours de chantier, ou (v) arrêt de chantier causé par le Client. Tout "
                        "arrêt injustifié imputé au Client fera l'objet d'une facturation "
                        "complémentaire couvrant les frais de mobilisation et d'immobilisation "
                        "des équipes.",
                        False,
                    ),
                ]
            )
            self._add_article_body_multi(art3_parts)
        else:
            art3_parts = (
                [
                    ("Works will begin on ", False),
                    (date_debut_str, True),
                    (" and end on ", False),
                    (date_fin_str, True),
                ]
                + duree_parts_en
                + [
                    (
                        ". These deadlines are given in good faith as estimates and "
                        "may be extended in the event of: (i) force majeure, (ii) material delivery "
                        "delays not attributable to the Service Provider, (iii) adverse weather "
                        "conditions, (iv) modifications requested by the Client during works, or "
                        "(v) site shutdown caused by the Client. Any unjustified shutdown "
                        "attributable to the Client will be subject to additional charges covering "
                        "mobilization and immobilization costs.",
                        False,
                    ),
                ]
            )
            self._add_article_body_multi(art3_parts)

        # Art 4
        self._add_article_title(self._t("art4_title"))
        if self.fr:
            if g_nb == 0:
                self._add_article_body_multi(
                    [
                        (
                            f"Les travaux objets du présent contrat sont réalisés ",
                            False,
                        ),
                        (g_text, True),
                        (
                            ". Le Client reconnaît avoir été expressément informé de cette absence "
                            "de garantie contractuelle et l'accepte. Cette disposition ne remet pas "
                            "en cause les obligations légales du Prestataire telles que définies par "
                            "la législation marocaine en vigueur.\n\n",
                            False,
                        ),
                        ("Exclusions de garantie : ", True),
                        (excl_garantie, False),
                    ]
                )
            else:
                self._add_article_body_multi(
                    [
                        ("Les travaux sont couverts par une garantie ", False),
                        (g_text, True),
                        (f" de type ", False),
                        (g_type_label, True),
                        (
                            " à compter de la date de réception et de signature du procès-verbal "
                            "de réception, contre tout défaut d'exécution directement imputable "
                            "au Prestataire.\n\n",
                            False,
                        ),
                        ("Exclusions de garantie : ", True),
                        (excl_garantie, False),
                    ]
                )
        else:
            if g_nb == 0:
                self._add_article_body_multi(
                    [
                        ("The works subject to this contract are performed ", False),
                        (g_text, True),
                        (
                            ". The Client acknowledges having been expressly informed of this "
                            "absence of contractual guarantee and accepts it. This provision does "
                            "not affect the Service Provider's legal obligations as defined by "
                            "current Moroccan legislation.\n\n",
                            False,
                        ),
                        ("Warranty exclusions: ", True),
                        (excl_garantie, False),
                    ]
                )
            else:
                self._add_article_body_multi(
                    [
                        ("The works are covered by a ", False),
                        (g_text, True),
                        (" guarantee (", False),
                        (g_type_label, True),
                        (
                            ") from the date of reception and signing of the acceptance "
                            "report, against any execution defect directly attributable to the "
                            "Service Provider.\n\n",
                            False,
                        ),
                        ("Warranty exclusions: ", True),
                        (excl_garantie, False),
                    ]
                )

        # Art 5
        self._add_article_title(self._t("art5_title"))
        if self.fr:
            self._add_article_body_multi(
                [
                    (
                        "La réception des travaux est effectuée contradictoirement à l'issue des "
                        "travaux. Un ",
                        False,
                    ),
                    ("procès-verbal de réception", True),
                    (
                        " sera établi et signé par les deux parties. Le solde du contrat (",
                        False,
                    ),
                    (f"{sp:.0f}%", True),
                    (" soit ", False),
                    (self._amt(mont_solde), True),
                    (
                        ") est exigible immédiatement "
                        "à la signature de ce procès-verbal. Toute réserve devra être notifiée "
                        "par écrit dans les ",
                        False,
                    ),
                    ("48 heures", True),
                    (
                        " suivant la fin des travaux ; l'absence de réserve dans ce délai vaut "
                        "acceptation définitive. Les travaux réservés feront l'objet d'une reprise "
                        "dans un délai convenu entre les parties.",
                        False,
                    ),
                ]
            )
        else:
            self._add_article_body_multi(
                [
                    (
                        "Works reception is carried out jointly at the completion of works. An ",
                        False,
                    ),
                    ("acceptance report", True),
                    (
                        " will be drawn up and signed by both parties. The contract balance (",
                        False,
                    ),
                    (f"{sp:.0f}%", True),
                    (", i.e. ", False),
                    (self._amt(mont_solde), True),
                    (
                        ") is payable immediately "
                        "upon signing this report. Any reservations must be notified in writing "
                        "within ",
                        False,
                    ),
                    ("48 hours", True),
                    (
                        " of the end of works; absence of reservations within this period "
                        "constitutes final acceptance. Reserved works will be subject to "
                        "rework within a timeframe agreed between the parties.",
                        False,
                    ),
                ]
            )

        # Art 6
        self._add_article_title(self._t("art6_title"))
        if self.fr:
            self._add_article_body_multi(
                [
                    (
                        "Tout retard de paiement entraînera l'application de pénalités de retard "
                        "au taux de ",
                        False,
                    ),
                    (penalites, True),
                    (
                        ", calculées à partir du lendemain de la date d'échéance, sans qu'une "
                        "mise en demeure préalable soit nécessaire. Le Prestataire se réserve "
                        "le droit de ",
                        False,
                    ),
                    ("suspendre les travaux", True),
                    (
                        " en cas de non-paiement après mise en demeure restée sans effet pendant "
                        "72 heures. Les frais de recouvrement éventuels seront à la charge du "
                        "Client.",
                        False,
                    ),
                ]
            )
        else:
            self._add_article_body_multi(
                [
                    (
                        "Any late payment will result in the application of late payment penalties "
                        "at the rate of ",
                        False,
                    ),
                    (penalites, True),
                    (
                        ", calculated from the day after the due date, without prior notice being "
                        "necessary. The Service Provider reserves the right to ",
                        False,
                    ),
                    ("suspend works", True),
                    (
                        " in case of non-payment after a formal notice that remains unanswered "
                        "for 72 hours. Any recovery costs will be borne by the Client.",
                        False,
                    ),
                ]
            )

        # Art 7
        self._add_article_title(self._t("art7_title"))
        if self.fr:
            resil_parts = _parse_strong(resil_text_raw)
            resil_parts.append(
                (
                    " En cas de résiliation anticipée à l'initiative du Client, "
                    "les travaux déjà réalisés seront facturés au prorata de leur avancement "
                    "constaté contradictoirement, et l'acompte versé restera acquis au "
                    "Prestataire à titre d'indemnité forfaitaire de dédit couvrant les préjudices "
                    "subis (mobilisation du matériel, main-d'œuvre réservée, sous-traitants "
                    "engagés, etc.).",
                    False,
                )
            )
            self._add_article_body_multi(resil_parts)
        else:
            resil_en = CLAUSE_RESILIATION_LABELS.get("en", {}).get(resil_val, "")
            resil_en_parts = _parse_strong(resil_en)
            resil_en_parts.append(
                (
                    " In case of early termination at the Client's initiative, "
                    "the works already completed will be invoiced pro rata based on jointly "
                    "verified progress, and the deposit paid will remain retained by the "
                    "Service Provider as a fixed penalty covering damages suffered (equipment "
                    "mobilization, reserved workforce, engaged subcontractors, etc.).",
                    False,
                )
            )
            self._add_article_body_multi(resil_en_parts)

        # Art 8
        self._add_article_title(self._t("art8_title"))
        if self.fr:
            self._add_article_body(
                "Le Prestataire s'engage à : (i) respecter les règles de sécurité en vigueur, "
                "(ii) maintenir le chantier dans un état de propreté acceptable, (iii) protéger "
                "les zones adjacentes non concernées par les travaux. Le Client s'engage à : "
                "(i) faciliter l'accès au chantier aux horaires convenus (généralement 8h00–18h00, "
                "du lundi au samedi), (ii) évacuer les meubles et objets gênants avant le début "
                "des travaux, (iii) informer le Prestataire de toute contrainte particulière "
                "(copropriété, voisinage, etc.)."
            )
        else:
            self._add_article_body(
                "The Service Provider undertakes to: (i) respect safety regulations in force, "
                "(ii) maintain the site in an acceptable state of cleanliness, (iii) protect "
                "adjacent areas unaffected by the works. The Client undertakes to: (i) facilitate "
                "access to the site during agreed hours (generally 8:00 AM - 6:00 PM, Monday "
                "to Saturday), (ii) remove furniture and obstructing items before works begin, "
                "(iii) inform the Service Provider of any particular constraints "
                "(condominium, neighbors, etc.)."
            )

        # Art 9
        self._add_article_title(self._t("art9_title"))
        if self.fr:
            self._add_article_body(
                "Le Prestataire déclare être couvert par une assurance responsabilité civile "
                "professionnelle couvrant les travaux objets du présent contrat. Sa responsabilité "
                "ne peut être engagée pour : les dommages résultant d'une mauvaise utilisation "
                "des ouvrages réalisés, des modifications effectuées par des tiers non mandatés, "
                "des vices cachés dans les matériaux fournis par le Client, ou des sinistres "
                "survenus après l'expiration de la garantie contractuelle."
            )
        else:
            self._add_article_body(
                "The Service Provider declares to be covered by professional liability insurance "
                "covering the works subject to this contract. Their liability cannot be engaged "
                "for: damages resulting from misuse of completed works, modifications made by "
                "unauthorized third parties, hidden defects in materials supplied by the Client, "
                "or incidents occurring after the expiration of the contractual guarantee."
            )

        # Art 10
        self._add_article_title(self._t("art10_title"))
        if self.fr:
            self._add_article_body(
                f"Les données personnelles collectées dans le cadre de ce contrat (nom, adresse, "
                f"coordonnées) sont utilisées exclusivement pour la gestion de la relation "
                f"contractuelle entre {co_name} et le Client. Elles ne seront en aucun cas "
                f"cédées à des tiers. Le Client dispose d'un droit d'accès, de rectification "
                f"et de suppression de ses données conformément à la législation en vigueur."
            )
        else:
            self._add_article_body(
                f"Personal data collected under this contract (name, address, contact details) "
                f"is used exclusively for managing the contractual relationship between "
                f"{co_name} and the Client. It will under no circumstances be transferred to "
                f"third parties. The Client has the right to access, rectify and delete their "
                f"data in compliance with current legislation."
            )

        # Art 11
        self._add_article_title(self._t("art11_title"))
        if self.fr:
            self._add_article_body_multi(
                [
                    (
                        "En cas de différend relatif à l'exécution ou à l'interprétation du présent "
                        "contrat, les parties s'engagent à rechercher une ",
                        False,
                    ),
                    ("solution amiable", True),
                    (
                        " dans un délai de 30 jours à compter de la notification du différend. "
                        "À défaut d'accord amiable, le litige sera porté exclusivement devant le ",
                        False,
                    ),
                    (tribunal, True),
                    (
                        ", nonobstant pluralité de défendeurs ou appel en garantie, même pour "
                        "les procédures d'urgence ou conservatoires.",
                        False,
                    ),
                ]
            )
        else:
            self._add_article_body_multi(
                [
                    (
                        "In the event of a dispute relating to the execution or interpretation of "
                        "this contract, the parties undertake to seek an ",
                        False,
                    ),
                    ("amicable solution", True),
                    (
                        " within 30 days of notification of the dispute. Failing amicable "
                        "agreement, the dispute will be brought exclusively before the ",
                        False,
                    ),
                    (tribunal, True),
                    (
                        ", notwithstanding plurality of defendants or warranty appeals, including "
                        "for emergency or conservatory proceedings.",
                        False,
                    ),
                ]
            )

        # Notes
        if c.notes:
            self._add_article_title(self._t("notes_title"))
            self._add_article_body(c.notes)

    def _build_signatures(self):
        self._add_empty(6)
        # Separator
        sep = self.doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(0)
        sep.paragraph_format.space_after = Pt(8)
        _para_border_bottom(sep, BORDER_HEX, "6")

        # "Fait à" info
        date_str = _fmt_date(self.c.date_contrat)
        ville = self.c.ville_signature if self.c.ville_signature else "Tanger"
        self._add_article_body_multi(
            [
                (f"{self._t('fait_a')} {ville}", True),
                (f", {self._t('le')} {date_str}, {self._t('en_deux_ex')}", False),
            ]
        )

        # Signature table
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        for i, (label, note, name) in enumerate(
            [
                (
                    self._t("prestataire_label"),
                    self._t("sig_cachet"),
                    BL_COMPANY["name"],
                ),
                (
                    self._t("client_label"),
                    self._t("lu_approuve"),
                    self.c.client_nom or "Client",
                ),
            ]
        ):
            cell = table.cell(0, i)
            _cell_borders(
                cell,
                top=(BORDER_HEX, 4),
                bottom=(BORDER_HEX, 4),
                left=(BORDER_HEX, 4),
                right=(BORDER_HEX, 4),
            )

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            lr = p.add_run(label)
            lr.font.size = Pt(7)
            lr.bold = True
            lr.font.color.rgb = BLUE
            lr.font.all_caps = True

            pn = cell.add_paragraph()
            pn.paragraph_format.space_after = Pt(1)
            nr = pn.add_run(note)
            nr.font.size = Pt(8)
            nr.font.color.rgb = GRAY
            nr.italic = True

            # Signature line space
            for _ in range(3):
                cell.add_paragraph().paragraph_format.space_after = Pt(4)

            ps = cell.add_paragraph()
            ps.paragraph_format.space_after = Pt(1)
            sr = ps.add_run(name)
            sr.font.size = Pt(8.5)
            sr.bold = True
            sr.font.color.rgb = NAVY

            pd = cell.add_paragraph()
            dr = pd.add_run(f"{self._t('date_sig')} : _______________")
            dr.font.size = Pt(7.5)
            dr.font.color.rgb = GRAY

    def _build_footer(self):
        co = BL_COMPANY
        self._add_empty(6)
        sep = self.doc.add_paragraph()
        _para_border_bottom(sep, BORDER_HEX, "2")
        sep.paragraph_format.space_after = Pt(4)

        now = datetime.now()
        gen_date = now.strftime("%d/%m/%Y")
        gen_time = now.strftime("%H:%M")
        p, r = self._add_text(
            f"{co['name']} — {self._t('footer_specialist')}\n"
            f"{co['phone']} | {co['email']} | {self._t('ice')} : {co['ice']}\n"
            f"{self._t('doc_genere')} {gen_date} {self._t('a')} {gen_time} — "
            f"{self._t('valeur_juridique')}",
            size=Pt(7),
            color=GRAY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    def _build(self):
        self._build_topstrip()
        self._build_header()
        self._build_title()
        self._build_parties()
        self._build_chantier()
        self._build_prestations()
        self._build_payment()
        self._build_articles()
        self._build_signatures()
        self._build_footer()

    def generate_response(self) -> HttpResponse:
        self._build()
        buf = BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        ref = (self.c.numero_contrat or "contract").replace("/", "-")
        filename = f"contrat_{self.c.id}_{ref}.docx"
        response = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response
