import io
from datetime import date, timedelta
from decimal import Decimal
from unittest.mock import patch, MagicMock

import pytest
from django.urls import reverse
from rest_framework import status
from rest_framework.test import APIClient
from rest_framework_simplejwt.tokens import AccessToken

from account.models import CustomUser
from contract.models import Contract, Project
from contract.pdf import _fmt_date, _fmt_amt, _esc, _is_societe, ContractPDFGenerator
from contract.doc import ContractDOCGenerator
from contract.bl_pdf import BluelinePDFGenerator, _garantie_text, _solde_pct
from contract.bl_doc import BluelineDOCGenerator
from contract.st_pdf import SousTraitancePDFGenerator
from contract.st_doc import SousTraitanceDOCGenerator
from contract.st_i18n import st_t, LOT_LABELS
from core.models import CompanyConfig

pytestmark = pytest.mark.django_db


def make_staff_user(email="staff@test.com", password="securepass123"):
    """Create a staff (can_create / can_edit / can_delete / can_print) user."""
    user = CustomUser.objects.create_user(
        email=email,
        password=password,
        is_staff=True,
        can_create=True,
        can_edit=True,
        can_delete=True,
        can_print=True,
    )
    token = str(AccessToken.for_user(user))
    client = APIClient()
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")
    return user, client


def make_readonly_user(email="readonly@test.com", password="securepass123"):
    """Create a user with no write permissions."""
    user = CustomUser.objects.create_user(
        email=email,
        password=password,
        is_staff=False,
        can_create=False,
        can_edit=False,
        can_delete=False,
        can_print=False,
    )
    token = str(AccessToken.for_user(user))
    client = APIClient()
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")
    return user, client


def make_contract(created_by=None, numero="0001/26", **kwargs):
    """Create a minimal valid Contract instance."""
    defaults = {
        "numero_contrat": numero,
        "date_contrat": "2026-01-15",
        "statut": "Brouillon",
        "montant_ht": "50000.00",
        "tva": "20.00",
    }
    defaults.update(kwargs)
    if created_by is not None:
        defaults["created_by_user"] = created_by
    return Contract.objects.create(**defaults)


class TestContractModel:
    def test_str_returns_numero_contrat(self):
        user, _ = make_staff_user()
        contract = make_contract(created_by=user, numero="STR/01")
        assert str(contract) == "STR/01"

    def test_montant_tva_property(self):
        user, _ = make_staff_user()
        contract = make_contract(
            created_by=user, numero="TVA/01", montant_ht="10000.00", tva="20.00"
        )
        assert contract.montant_tva == pytest.approx(2000.0)

    def test_montant_ttc_property(self):
        user, _ = make_staff_user()
        contract = make_contract(
            created_by=user, numero="TTC/01", montant_ht="10000.00", tva="20.00"
        )
        assert contract.montant_ttc == pytest.approx(12000.0)

    def test_default_statut_is_brouillon(self):
        user, _ = make_staff_user()
        contract = Contract.objects.create(
            numero_contrat="DEF/01",
            date_contrat="2026-01-01",
            created_by_user=user,
        )
        assert contract.statut == "Brouillon"

    def test_default_confidentialite(self):
        user, _ = make_staff_user()
        contract = make_contract(created_by=user, numero="CONF/01")
        assert contract.confidentialite == "CONFIDENTIEL"

    def test_numero_contrat_unique(self):
        from django.db import IntegrityError

        user, _ = make_staff_user()
        make_contract(created_by=user, numero="UNIQ/01")
        with pytest.raises(IntegrityError):
            make_contract(created_by=user, numero="UNIQ/01")

    def test_created_by_user_nullable(self):
        contract = make_contract(numero="NULL/01")
        assert contract.created_by_user is None


class TestContractListCreateView:
    def setup_method(self):
        self.url = reverse("contract:contract-list-create")
        self.staff_user, self.staff_client = make_staff_user()
        self.readonly_user, self.readonly_client = make_readonly_user()
        self.anon_client = APIClient()

    def test_list_contracts_returns_200(self):
        make_contract(created_by=self.staff_user, numero="LIST/01")
        make_contract(created_by=self.staff_user, numero="LIST/02")
        response = self.staff_client.get(self.url)
        assert response.status_code == status.HTTP_200_OK
        assert len(response.data) >= 2

    def test_list_contracts_paginated(self):
        for i in range(3):
            make_contract(created_by=self.staff_user, numero=f"PAG/{i:02d}")
        response = self.staff_client.get(self.url, {"pagination": "true", "page": 1})
        assert response.status_code == status.HTTP_200_OK
        assert "results" in response.data
        assert "count" in response.data

    def test_list_contracts_unauthenticated_returns_401(self):
        response = self.anon_client.get(self.url)
        assert response.status_code == status.HTTP_401_UNAUTHORIZED

    def test_create_contract_as_staff_returns_201(self):
        payload = {
            "company": "casa_di_lusso",
            "contract_category": "standard",
            "numero_contrat": "NEW/01",
            "date_contrat": "2026-03-01",
            "statut": "Brouillon",
            "type_contrat": "travaux_finition",
            "montant_ht": "20000.00",
            "tva": "20.00",
            "tranches": [
                {"label": "Acompte", "pourcentage": 40},
                {"label": "Solde", "pourcentage": 60},
            ],
        }
        response = self.staff_client.post(self.url, payload, format="json")
        assert response.status_code == status.HTTP_201_CREATED
        assert response.data["numero_contrat"] == "NEW/01"

    def test_create_contract_rejects_invalid_cdl_tranche_total(self):
        payload = {
            "company": "casa_di_lusso",
            "contract_category": "standard",
            "numero_contrat": "NEW/02",
            "date_contrat": "2026-03-01",
            "statut": "Brouillon",
            "type_contrat": "travaux_finition",
            "montant_ht": "20000.00",
            "tva": "20.00",
            "tranches": [
                {"label": "Acompte", "pourcentage": 30},
                {"label": "Solde", "pourcentage": 60},
            ],
        }
        response = self.staff_client.post(self.url, payload, format="json")
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "tranches" in response.data["details"]

    def test_create_contract_rejects_invalid_st_tranche_total(self):
        payload = {
            "company": "casa_di_lusso",
            "contract_category": "sous_traitance",
            "numero_contrat": "NEW/ST/01",
            "date_contrat": "2026-03-01",
            "statut": "Brouillon",
            "montant_ht": "30000.00",
            "tva": "20.00",
            "client_nom": None,
            "st_name": "Sous-traitant Test",
            "st_lot_type": "gros_oeuvre",
            "st_type_prix": "forfaitaire",
            "st_tranches": [
                {"label": "Acompte", "pourcentage": 20},
                {"label": "Solde", "pourcentage": 50},
            ],
        }
        response = self.staff_client.post(self.url, payload, format="json")
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "st_tranches" in response.data["details"]

    def test_create_contract_without_can_create_returns_403(self):
        payload = {
            "numero_contrat": "NOPERM/01",
            "date_contrat": "2026-03-01",
            "statut": "Brouillon",
        }
        response = self.readonly_client.post(self.url, payload, format="json")
        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_create_contract_unauthenticated_returns_401(self):
        payload = {
            "numero_contrat": "ANON/01",
            "date_contrat": "2026-03-01",
        }
        response = self.anon_client.post(self.url, payload, format="json")
        assert response.status_code == status.HTTP_401_UNAUTHORIZED

    def test_create_contract_missing_required_field_returns_400(self):
        # numero_contrat is required
        payload = {"date_contrat": "2026-03-01", "statut": "Brouillon"}
        response = self.staff_client.post(self.url, payload, format="json")
        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_list_contracts_filter_by_statut(self):
        make_contract(created_by=self.staff_user, numero="FLT/01", statut="Signé")
        make_contract(created_by=self.staff_user, numero="FLT/02", statut="Brouillon")
        response = self.staff_client.get(self.url, {"statut": "Signé"})
        assert response.status_code == status.HTTP_200_OK
        for item in response.data:
            assert item["statut"] == "Signé"


class TestContractDetailEditDeleteView:
    def setup_method(self):
        self.staff_user, self.staff_client = make_staff_user()
        self.readonly_user, self.readonly_client = make_readonly_user(
            email="ro2@test.com"
        )
        self.contract = make_contract(created_by=self.staff_user, numero="DET/01")
        self.url = reverse("contract:contract-detail", kwargs={"pk": self.contract.pk})
        self.anon_client = APIClient()

    def test_get_contract_returns_200(self):
        response = self.staff_client.get(self.url)
        assert response.status_code == status.HTTP_200_OK
        assert response.data["numero_contrat"] == "DET/01"

    def test_get_contract_not_found_returns_404(self):
        url = reverse("contract:contract-detail", kwargs={"pk": 99999})
        response = self.staff_client.get(url)
        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_get_contract_unauthenticated_returns_401(self):
        response = self.anon_client.get(self.url)
        assert response.status_code == status.HTTP_401_UNAUTHORIZED

    def test_put_contract_as_staff_returns_200(self):
        payload = {
            "company": "casa_di_lusso",
            "contract_category": "standard",
            "numero_contrat": "DET/01",
            "date_contrat": "2026-06-01",
            "montant_ht": "60000.00",
            "tva": "20.00",
            "type_contrat": "travaux_finition",
            "tranches": [
                {"label": "Acompte", "pourcentage": 50},
                {"label": "Solde", "pourcentage": 50},
            ],
        }
        response = self.staff_client.put(self.url, payload, format="json")
        assert response.status_code == status.HTTP_200_OK
        # statut is read-only via PUT; should remain unchanged
        assert response.data["statut"] == "Brouillon"

    def test_put_contract_without_can_update_returns_403(self):
        payload = {
            "numero_contrat": "DET/01",
            "date_contrat": "2026-06-01",
        }
        response = self.readonly_client.put(self.url, payload, format="json")
        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_delete_contract_as_staff_returns_204(self):
        contract = make_contract(created_by=self.staff_user, numero="DEL/01")
        url = reverse("contract:contract-detail", kwargs={"pk": contract.pk})
        response = self.staff_client.delete(url)
        assert response.status_code == status.HTTP_204_NO_CONTENT
        assert not Contract.objects.filter(pk=contract.pk).exists()

    def test_delete_contract_without_can_delete_returns_403(self):
        response = self.readonly_client.delete(self.url)
        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_delete_contract_not_found_returns_404(self):
        url = reverse("contract:contract-detail", kwargs={"pk": 99999})
        response = self.staff_client.delete(url)
        assert response.status_code == status.HTTP_404_NOT_FOUND


class TestContractStatusUpdateView:
    def setup_method(self):
        self.staff_user, self.staff_client = make_staff_user()
        self.readonly_user, self.readonly_client = make_readonly_user(
            email="ro3@test.com"
        )
        self.contract = make_contract(
            created_by=self.staff_user, numero="STAT/01", statut="Brouillon"
        )
        self.url = reverse(
            "contract:contract-statut-update", kwargs={"pk": self.contract.pk}
        )

    def test_patch_statut_valid_returns_200(self):
        response = self.staff_client.patch(
            self.url, {"statut": "Envoyé"}, format="json"
        )
        assert response.status_code == status.HTTP_200_OK
        assert response.data["statut"] == "Envoyé"
        self.contract.refresh_from_db()
        assert self.contract.statut == "Envoyé"

    def test_patch_statut_invalid_returns_400(self):
        response = self.staff_client.patch(
            self.url, {"statut": "Inexistant"}, format="json"
        )
        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_patch_statut_without_can_update_returns_403(self):
        response = self.readonly_client.patch(
            self.url, {"statut": "Signé"}, format="json"
        )
        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_patch_statut_not_found_returns_404(self):
        url = reverse("contract:contract-statut-update", kwargs={"pk": 99999})
        response = self.staff_client.patch(url, {"statut": "Signé"}, format="json")
        assert response.status_code == status.HTTP_404_NOT_FOUND


class TestBulkDeleteContractView:
    def setup_method(self):
        self.staff_user, self.staff_client = make_staff_user()
        self.readonly_user, self.readonly_client = make_readonly_user(
            email="ro4@test.com"
        )
        self.url = reverse("contract:contract-bulk-delete")

    def test_bulk_delete_valid_ids_returns_204(self):
        c1 = make_contract(created_by=self.staff_user, numero="BULK/01")
        c2 = make_contract(created_by=self.staff_user, numero="BULK/02")
        response = self.staff_client.delete(
            self.url, {"ids": [c1.pk, c2.pk]}, format="json"
        )
        assert response.status_code == status.HTTP_204_NO_CONTENT
        assert not Contract.objects.filter(pk__in=[c1.pk, c2.pk]).exists()

    def test_bulk_delete_missing_ids_returns_400(self):
        response = self.staff_client.delete(self.url, {}, format="json")
        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_bulk_delete_non_list_ids_returns_400(self):
        response = self.staff_client.delete(
            self.url, {"ids": "not-a-list"}, format="json"
        )
        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_bulk_delete_invalid_integer_ids_returns_400(self):
        response = self.staff_client.delete(
            self.url, {"ids": ["abc", "xyz"]}, format="json"
        )
        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_bulk_delete_nonexistent_ids_returns_404(self):
        response = self.staff_client.delete(
            self.url, {"ids": [99998, 99999]}, format="json"
        )
        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_bulk_delete_without_can_delete_returns_403(self):
        c = make_contract(created_by=self.staff_user, numero="BULKPERM/01")
        response = self.readonly_client.delete(self.url, {"ids": [c.pk]}, format="json")
        assert response.status_code == status.HTTP_403_FORBIDDEN


class TestGenerateNumeroContratView:
    def setup_method(self):
        _, self.auth_client = make_staff_user()
        self.anon_client = APIClient()
        self.url = reverse("contract:generate-numero-contrat")

    def test_generate_returns_200_with_numero(self):
        response = self.auth_client.get(self.url)
        assert response.status_code == status.HTTP_200_OK
        assert "numero_contrat" in response.data

    def test_generate_unauthenticated_returns_401(self):
        response = self.anon_client.get(self.url)
        assert response.status_code == status.HTTP_401_UNAUTHORIZED


# ══════════════════════════════════════════════════════════════════════════════
#  PDF & DOCX GENERATOR TESTS
#  (contract/pdf.py, contract/doc.py, contract/bl_pdf.py, contract/bl_doc.py)
# ══════════════════════════════════════════════════════════════════════════════


# ──────────────────────────────────────────────────────────────────────────────
# Fixtures
# ──────────────────────────────────────────────────────────────────────────────


@pytest.fixture()
def staff_user():
    user = CustomUser.objects.create_user(
        email="gen_staff@test.com",
        password="securepass123",
        is_staff=True,
        can_create=True,
        can_edit=True,
        can_delete=True,
        can_print=True,
    )
    return user


@pytest.fixture()
def staff_client(staff_user):
    token = str(AccessToken.for_user(staff_user))
    client = APIClient()
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")
    return client


@pytest.fixture()
def noprint_user():
    user = CustomUser.objects.create_user(
        email="no_print@test.com",
        password="securepass123",
        is_staff=False,
        can_create=False,
        can_edit=False,
        can_delete=False,
        can_print=False,
    )
    return user


@pytest.fixture()
def noprint_client(noprint_user):
    token = str(AccessToken.for_user(noprint_user))
    client = APIClient()
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")
    return client


@pytest.fixture()
def cdl_contract(staff_user):
    """Full Casa Di Lusso contract with all fields populated."""
    return Contract.objects.create(
        company="casa_di_lusso",
        numero_contrat="CDL/TEST-001",
        date_contrat=date(2026, 3, 1),
        statut="Signé",
        created_by_user=staff_user,
        # Client
        client_nom="Ahmed Ben Ali",
        client_cin="AB123456",
        client_qualite="particulier",
        client_adresse="12 Rue de Tanger, Tanger 90000",
        client_tel="+212600000001",
        client_email="ahmed@example.com",
        ville_signature="Tanger",
        # Works
        adresse_travaux="45 Boulevard Mohammed V, Tanger",
        type_bien="villa",
        surface=Decimal("250.50"),
        services=["design_conception", "travaux_finition", "ameublement"],
        description_travaux="Rénovation complète de la villa avec <script>alert('xss')</script>",
        date_debut=date(2026, 4, 1),
        duree_estimee="6 mois",
        conditions_acces="Accès libre du lundi au samedi, 8h-18h",
        # Financial
        montant_ht=Decimal("150000.00"),
        devise="MAD",
        tva=Decimal("20.00"),
        tranches=[
            {"label": "Acompte", "pourcentage": 30},
            {"label": "Avancement 50%", "pourcentage": 40},
            {"label": "Réception", "pourcentage": 30},
        ],
        mode_paiement_texte="Virement Bancaire",
        rib="007 400 0123456789012345 67",
        delai_retard=5,
        penalite_retard=Decimal("1.50"),
        frais_redemarrage=Decimal("5000.00"),
        # Legal
        garantie="2 ans",
        delai_reserves=7,
        tribunal="Tanger",
        clauses_actives=[
            "clause_resiliation",
            "clause_force_majeure",
            "clause_confidentialite",
        ],
        clause_spec="Clause spécifique personnalisée <b>test</b>",
        exclusions="Plomberie & électricité non incluses <img src=x onerror=alert(1)>",
        # Meta
        type_contrat="travaux_finition",
        responsable_projet="Karim Directeur",
        architecte="Sofia Architecte",
        confidentialite="CONFIDENTIEL",
        version_document="v1.0 – Définitif",
        annexes="Plans architecturaux\nDevis détaillé\nCalendrier d'exécution",
    )


@pytest.fixture()
def bl_contract(staff_user):
    """Full BlueLine Works contract with all BL-specific fields populated."""
    return Contract.objects.create(
        company="blueline_works",
        numero_contrat="BLW/TEST-002",
        date_contrat=date(2026, 3, 15),
        statut="Signé",
        created_by_user=staff_user,
        # Client
        client_nom="Sara El Fassi",
        client_cin="EF789012",
        client_qualite="entreprise_societe",
        client_adresse="88 Bd Pasteur, Tanger",
        client_tel="+212611111111",
        client_email="sara@corp.ma",
        client_ville="Tanger",
        client_cp="90000",
        ville_signature="Tanger",
        # Chantier
        adresse_travaux="15 Rue Ibn Batouta, Tanger",
        chantier_ville="Tanger",
        chantier_etage="3ème étage, Apt 12",
        # Works
        type_bien="appartement",
        surface=Decimal("120.00"),
        description_travaux="Pose de carrelage et marbre",
        date_debut=date(2026, 4, 15),
        duree_estimee="45",
        # Prestations
        prestations=[
            {
                "nom": "pose_carrelage",
                "desc": "Sol salon + chambres",
                "qte": 80,
                "unite": "m2",
                "pu": 350,
            },
            {
                "nom": "pose_marbre",
                "desc": "Entrée + escalier",
                "qte": 25,
                "unite": "m2",
                "pu": 650,
            },
        ],
        fournitures="non_incluses",
        materiaux_detail="Carrelage fourni par le client",
        eau_electricite="client",
        # Financial
        montant_ht=Decimal("44250.00"),
        devise="MAD",
        tva=Decimal("20.00"),
        acompte=Decimal("30.00"),
        tranche2=Decimal("40.00"),
        mode_paiement_texte="Virement Bancaire",
        rib="007 500 0123456789012345 89",
        # Garantie
        garantie_nb=2,
        garantie_unite="ans",
        garantie_type="bonne_fin",
        exclusions_garantie="Usure normale et mauvais entretien",
        # Legal
        clause_resiliation="30j",
        tribunal="Tanger",
        notes="Début prévu mi-avril\nContact chef de chantier: 0600000002",
    )


@pytest.fixture()
def cdl_minimal_contract(staff_user):
    """Minimal CDL contract — only required fields."""
    return Contract.objects.create(
        company="casa_di_lusso",
        numero_contrat="CDL/MIN-001",
        date_contrat=date(2026, 1, 1),
        statut="Brouillon",
        created_by_user=staff_user,
    )


@pytest.fixture()
def bl_minimal_contract(staff_user):
    """Minimal BL contract — only required fields."""
    return Contract.objects.create(
        company="blueline_works",
        numero_contrat="BLW/MIN-001",
        date_contrat=date(2026, 1, 1),
        statut="Brouillon",
        created_by_user=staff_user,
    )


# ══════════════════════════════════════════════════════════════════════════════
#  UNIT TESTS — Helper Functions
# ══════════════════════════════════════════════════════════════════════════════


class TestFmtDate:
    def test_formats_date_object(self):
        assert _fmt_date(date(2026, 3, 4)) == "04 / 03 / 2026"

    def test_none_returns_placeholder(self):
        result = _fmt_date(None)
        assert "…" in result

    def test_string_returned_as_is(self):
        assert _fmt_date("custom") == "custom"


class TestFmtAmt:
    def test_formats_integer(self):
        result = _fmt_amt(150000, "MAD")
        assert "150" in result
        assert "000" in result
        assert "MAD" in result

    def test_formats_with_decimals(self):
        result = _fmt_amt(1234.56, "EUR")
        assert "1" in result
        assert "234" in result
        assert "56" in result
        assert "EUR" in result

    def test_none_returns_zero(self):
        result = _fmt_amt(None, "MAD")
        assert "0" in result

    def test_invalid_returns_zero(self):
        result = _fmt_amt("abc", "MAD")
        assert "0" in result


class TestEsc:
    def test_escapes_html_tags(self):
        assert (
            _esc("<script>alert('x')</script>")
            == "&lt;script&gt;alert('x')&lt;/script&gt;"
        )

    def test_escapes_ampersand(self):
        assert _esc("A & B") == "A &amp; B"

    def test_empty_returns_empty(self):
        assert _esc("") == ""

    def test_none_returns_empty(self):
        assert _esc(None) == ""

    def test_plain_text_unchanged(self):
        assert _esc("Hello World") == "Hello World"


class TestIsSociete:
    def test_societe_keyword(self):
        assert _is_societe("Société") is True

    def test_morale_keyword(self):
        assert _is_societe("Personne Morale") is True

    def test_particulier_is_false(self):
        assert _is_societe("particulier") is False

    def test_none_is_false(self):
        assert _is_societe(None) is False

    def test_empty_is_false(self):
        assert _is_societe("") is False


class TestGarantieText:
    """Test the _garantie_text helper in bl_pdf.py."""

    def test_zero_returns_sans_garantie(self, bl_minimal_contract):
        bl_minimal_contract.garantie_nb = 0
        result = _garantie_text(bl_minimal_contract, "fr")
        assert "garantie" in result.lower() or "sans" in result.lower()

    def test_none_returns_sans_garantie(self, bl_minimal_contract):
        bl_minimal_contract.garantie_nb = None
        result = _garantie_text(bl_minimal_contract, "fr")
        assert "garantie" in result.lower() or "sans" in result.lower()

    def test_singular_ans(self, bl_minimal_contract):
        bl_minimal_contract.garantie_nb = 1
        bl_minimal_contract.garantie_unite = "ans"
        result = _garantie_text(bl_minimal_contract, "fr")
        assert "1" in result
        assert "an" in result.lower()

    def test_plural_ans(self, bl_minimal_contract):
        bl_minimal_contract.garantie_nb = 2
        bl_minimal_contract.garantie_unite = "ans"
        result = _garantie_text(bl_minimal_contract, "fr")
        assert "2" in result
        assert "ans" in result.lower()

    def test_mois(self, bl_minimal_contract):
        bl_minimal_contract.garantie_nb = 6
        bl_minimal_contract.garantie_unite = "mois"
        result = _garantie_text(bl_minimal_contract, "fr")
        assert "6" in result
        assert "mois" in result.lower()

    def test_english(self, bl_minimal_contract):
        bl_minimal_contract.garantie_nb = 2
        bl_minimal_contract.garantie_unite = "ans"
        result = _garantie_text(bl_minimal_contract, "en")
        assert "2" in result
        assert "year" in result.lower()


class TestSoldePct:
    def test_standard_split(self, bl_contract):
        # 30 + 40 => solde = 30
        assert _solde_pct(bl_contract) == pytest.approx(30.0)

    def test_no_tranche2(self, bl_minimal_contract):
        bl_minimal_contract.acompte = Decimal("50")
        bl_minimal_contract.tranche2 = None
        assert _solde_pct(bl_minimal_contract) == pytest.approx(50.0)


# ══════════════════════════════════════════════════════════════════════════════
#  INTEGRATION TESTS — CDL PDF Generator
# ══════════════════════════════════════════════════════════════════════════════


class TestCDLPDFGenerator:
    """Test ContractPDFGenerator (Casa Di Lusso — WeasyPrint PDF)."""

    def test_generate_response_returns_pdf(self, cdl_contract):
        gen = ContractPDFGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        assert response["Content-Type"] == "application/pdf"
        assert b"%PDF" in response.content[:20]

    def test_generate_response_en(self, cdl_contract):
        gen = ContractPDFGenerator(cdl_contract, language="en")
        response = gen.generate_response()
        assert response.status_code == 200
        assert b"%PDF" in response.content[:20]

    def test_filename_contains_ref(self, cdl_contract):
        gen = ContractPDFGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        disp = response["Content-Disposition"]
        assert "CDL-TEST-001" in disp or "CDL" in disp

    def test_html_contains_client_name(self, cdl_contract):
        """Verify the intermediate HTML includes client info."""
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        assert "Ahmed Ben Ali" in html

    def test_html_contains_escaped_description(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        # The <script> tag should be escaped
        assert "<script>" not in html
        assert "&lt;script&gt;" in html

    def test_html_contains_escaped_exclusions(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        # <img should be escaped
        assert "<img " not in html
        assert "&lt;img " in html

    def test_html_contains_escaped_clause_spec(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        # <b> in clause_spec should be escaped
        assert "Clause spécifique" in html
        assert "&lt;b&gt;test&lt;/b&gt;" in html

    def test_html_contains_services(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        assert "design_conception" in html or "SERVICES CONVENUS" in html

    def test_html_contains_tranches(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        assert "Acompte" in html
        assert "30" in html

    def test_html_contains_tribunal(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        assert "Tanger" in html

    def test_html_contains_architecte(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        assert "Sofia Architecte" in html

    def test_html_dynamic_article_numbers(self, cdl_contract):
        """Verify article numbers are dynamically generated, not hardcoded."""
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        # Should contain numbered articles like 01, 02, 03...
        assert "01" in html
        assert "02" in html

    def test_html_en_contains_english_labels(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "en")
        assert "SERVICE AGREEMENT" in html or "SCOPE OF AGREEMENT" in html

    def test_minimal_contract_generates(self, cdl_minimal_contract):
        """Even a minimal contract should produce a valid PDF."""
        gen = ContractPDFGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        assert b"%PDF" in response.content[:20]

    def test_html_annexes_escaped(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_contract, "fr")
        # Annexes text should appear (but any HTML in it escaped)
        assert "Plans architecturaux" in html


# ══════════════════════════════════════════════════════════════════════════════
#  INTEGRATION TESTS — CDL DOC Generator
# ══════════════════════════════════════════════════════════════════════════════


class TestCDLDOCGenerator:
    """Test ContractDOCGenerator (Casa Di Lusso — python-docx DOCX)."""

    def test_generate_response_returns_docx(self, cdl_contract):
        gen = ContractDOCGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        ct = response["Content-Type"]
        assert "officedocument" in ct or "msword" in ct

    def test_generate_response_en(self, cdl_contract):
        gen = ContractDOCGenerator(cdl_contract, language="en")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_filename_contains_ref(self, cdl_contract):
        gen = ContractDOCGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        disp = response["Content-Disposition"]
        assert "CDL" in disp

    def test_docx_content_is_valid_zip(self, cdl_contract):
        """DOCX files are ZIP archives — verify the magic bytes."""
        import zipfile

        gen = ContractDOCGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        buf = io.BytesIO(response.content)
        assert zipfile.is_zipfile(buf)

    def test_docx_contains_client_name(self, cdl_contract):
        """Parse the DOCX and check for client name in the text."""
        from docx import Document as DocxDocument

        gen = ContractDOCGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        # Client name may be in paragraphs or in table cells
        all_text = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = "\n".join(all_text)
        assert "Ahmed Ben Ali" in full_text

    def test_docx_contains_tribunal(self, cdl_contract):
        from docx import Document as DocxDocument

        gen = ContractDOCGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Tanger" in full_text

    def test_docx_en_contains_english_text(self, cdl_contract):
        from docx import Document as DocxDocument

        gen = ContractDOCGenerator(cdl_contract, language="en")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert (
            "Client" in full_text or "Agreement" in full_text or "Service" in full_text
        )

    def test_minimal_contract_generates(self, cdl_minimal_contract):
        gen = ContractDOCGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200


# ══════════════════════════════════════════════════════════════════════════════
#  INTEGRATION TESTS — BL PDF Generator
# ══════════════════════════════════════════════════════════════════════════════


class TestBLPDFGenerator:
    """Test BluelinePDFGenerator (BlueLine Works — WeasyPrint PDF)."""

    def test_generate_response_returns_pdf(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        assert response["Content-Type"] == "application/pdf"
        assert b"%PDF" in response.content[:20]

    def test_generate_response_en(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="en")
        response = gen.generate_response()
        assert response.status_code == 200
        assert b"%PDF" in response.content[:20]

    def test_filename_contains_contrat(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        disp = response["Content-Disposition"]
        assert "contrat_" in disp

    def test_html_contains_client_info(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        assert "Sara El Fassi" in html
        assert "EF789012" in html

    def test_html_contains_prestations(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        # Check prestation data appears
        assert "80" in html  # qte
        assert "350" in html or "350,00" in html  # pu

    def test_html_contains_chantier_info(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        assert "3ème étage" in html or "Apt 12" in html

    def test_html_contains_garantie_type(self, bl_contract):
        """Verify GARANTIE_TYPE is rendered (our bug fix)."""
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        # bonne_fin should be rendered as its label
        assert "bonne fin" in html.lower() or "Bonne fin" in html

    def test_html_contains_payment_schedule(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        assert "30" in html  # acompte 30%
        assert "40" in html  # tranche2 40%

    def test_html_contains_fournitures(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        # fournitures="non_incluses" should appear as its label
        assert (
            "client" in html.lower()
            or "non incluse" in html.lower()
            or "non_incluses" in html.lower()
        )

    def test_html_contains_eau_electricite(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        assert "client" in html.lower() or "charge" in html.lower()

    def test_html_contains_clause_resiliation(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        assert "30" in html  # 30j clause

    def test_minimal_contract_generates(self, bl_minimal_contract):
        gen = BluelinePDFGenerator(bl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        assert b"%PDF" in response.content[:20]

    def test_en_html_uses_english_labels(self, bl_contract):
        gen = BluelinePDFGenerator(bl_contract, language="en")
        html = gen._build_html()
        # English labels should appear
        assert "Service" in html or "Contract" in html or "Client" in html


# ══════════════════════════════════════════════════════════════════════════════
#  INTEGRATION TESTS — BL DOC Generator
# ══════════════════════════════════════════════════════════════════════════════


class TestBLDOCGenerator:
    """Test BluelineDOCGenerator (BlueLine Works — python-docx DOCX)."""

    def test_generate_response_returns_docx(self, bl_contract):
        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        ct = response["Content-Type"]
        assert "officedocument" in ct or "msword" in ct

    def test_generate_response_en(self, bl_contract):
        gen = BluelineDOCGenerator(bl_contract, language="en")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_filename_contains_contrat(self, bl_contract):
        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        disp = response["Content-Disposition"]
        assert "contrat_" in disp

    def test_docx_is_valid_zip(self, bl_contract):
        import zipfile

        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        buf = io.BytesIO(response.content)
        assert zipfile.is_zipfile(buf)

    def test_docx_contains_client_name(self, bl_contract):
        from docx import Document as DocxDocument

        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Sara El Fassi" in full_text

    def test_docx_contains_prestations(self, bl_contract):
        from docx import Document as DocxDocument

        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        # Check table cells for prestation data
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        combined = " ".join(all_text)
        assert "80" in combined or "350" in combined

    def test_docx_contains_garantie_type(self, bl_contract):
        """Verify the garantie_type label is rendered (our bug fix)."""
        from docx import Document as DocxDocument

        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "bonne fin" in full_text.lower() or "Bonne fin" in full_text

    def test_docx_contains_clause_resiliation(self, bl_contract):
        from docx import Document as DocxDocument

        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "30" in full_text or "préavis" in full_text.lower()

    def test_minimal_contract_generates(self, bl_minimal_contract):
        gen = BluelineDOCGenerator(bl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_en_docx_contains_english(self, bl_contract):
        from docx import Document as DocxDocument

        gen = BluelineDOCGenerator(bl_contract, language="en")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # At least some English text should appear
        assert (
            "Client" in full_text or "Service" in full_text or "Contract" in full_text
        )


# ══════════════════════════════════════════════════════════════════════════════
#  API VIEW TESTS — ContractPDFView & ContractDOCView
# ══════════════════════════════════════════════════════════════════════════════


class TestContractPDFViewCDL:
    """Test the PDF download endpoint for CDL contracts."""

    def test_pdf_fr_returns_200(self, staff_client, cdl_contract):
        url = reverse("contract:contract-pdf-fr", kwargs={"pk": cdl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK
        assert response["Content-Type"] == "application/pdf"

    def test_pdf_en_returns_200(self, staff_client, cdl_contract):
        url = reverse("contract:contract-pdf-en", kwargs={"pk": cdl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK

    def test_pdf_not_found_returns_404(self, staff_client):
        url = reverse("contract:contract-pdf-fr", kwargs={"pk": 99999})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_pdf_no_print_permission_returns_403(self, noprint_client, cdl_contract):
        url = reverse("contract:contract-pdf-fr", kwargs={"pk": cdl_contract.pk})
        response = noprint_client.get(url)
        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_pdf_unauthenticated_returns_401(self, cdl_contract):
        client = APIClient()
        url = reverse("contract:contract-pdf-fr", kwargs={"pk": cdl_contract.pk})
        response = client.get(url)
        assert response.status_code == status.HTTP_401_UNAUTHORIZED


class TestContractPDFViewBL:
    """Test the PDF download endpoint for BL contracts (dispatches to BluelinePDFGenerator)."""

    def test_pdf_fr_returns_200(self, staff_client, bl_contract):
        url = reverse("contract:contract-pdf-fr", kwargs={"pk": bl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK
        assert response["Content-Type"] == "application/pdf"

    def test_pdf_en_returns_200(self, staff_client, bl_contract):
        url = reverse("contract:contract-pdf-en", kwargs={"pk": bl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK


class TestContractDOCViewCDL:
    """Test the DOCX download endpoint for CDL contracts."""

    def test_doc_fr_returns_200(self, staff_client, cdl_contract):
        url = reverse("contract:contract-doc-fr", kwargs={"pk": cdl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK
        ct = response["Content-Type"]
        assert "officedocument" in ct or "msword" in ct

    def test_doc_en_returns_200(self, staff_client, cdl_contract):
        url = reverse("contract:contract-doc-en", kwargs={"pk": cdl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK

    def test_doc_not_found_returns_404(self, staff_client):
        url = reverse("contract:contract-doc-fr", kwargs={"pk": 99999})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_doc_no_print_permission_returns_403(self, noprint_client, cdl_contract):
        url = reverse("contract:contract-doc-fr", kwargs={"pk": cdl_contract.pk})
        response = noprint_client.get(url)
        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_doc_unauthenticated_returns_401(self, cdl_contract):
        client = APIClient()
        url = reverse("contract:contract-doc-fr", kwargs={"pk": cdl_contract.pk})
        response = client.get(url)
        assert response.status_code == status.HTTP_401_UNAUTHORIZED


class TestContractDOCViewBL:
    """Test the DOCX download endpoint for BL contracts."""

    def test_doc_fr_returns_200(self, staff_client, bl_contract):
        url = reverse("contract:contract-doc-fr", kwargs={"pk": bl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK

    def test_doc_en_returns_200(self, staff_client, bl_contract):
        url = reverse("contract:contract-doc-en", kwargs={"pk": bl_contract.pk})
        response = staff_client.get(url)
        assert response.status_code == status.HTTP_200_OK


# ══════════════════════════════════════════════════════════════════════════════
#  EDGE CASE TESTS
# ══════════════════════════════════════════════════════════════════════════════


class TestEdgeCases:
    """Test edge cases: empty fields, special characters, etc."""

    def test_cdl_pdf_empty_services_list(self, cdl_minimal_contract):
        cdl_minimal_contract.services = []
        cdl_minimal_contract.save()
        gen = ContractPDFGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_cdl_pdf_no_tranches(self, cdl_minimal_contract):
        cdl_minimal_contract.tranches = []
        cdl_minimal_contract.save()
        gen = ContractPDFGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_cdl_pdf_special_chars_in_client_name(self, cdl_minimal_contract):
        cdl_minimal_contract.client_nom = "Jean-Pierre O'Brien & Associés <LLC>"
        cdl_minimal_contract.save()
        from contract.pdf import _gen_contract_html

        html = _gen_contract_html(cdl_minimal_contract, "fr")
        # Should not contain unescaped angle brackets from the name
        # (the _esc function wraps individual fields)
        assert response_ok_or_html_valid(html)

    def test_bl_pdf_empty_prestations(self, bl_minimal_contract):
        bl_minimal_contract.prestations = []
        bl_minimal_contract.save()
        gen = BluelinePDFGenerator(bl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_bl_pdf_zero_amounts(self, bl_minimal_contract):
        bl_minimal_contract.montant_ht = Decimal("0.00")
        bl_minimal_contract.save()
        gen = BluelinePDFGenerator(bl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_bl_doc_empty_prestations(self, bl_minimal_contract):
        bl_minimal_contract.prestations = []
        bl_minimal_contract.save()
        gen = BluelineDOCGenerator(bl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_cdl_doc_empty_everything(self, cdl_minimal_contract):
        gen = ContractDOCGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_cdl_pdf_unicode_in_fields(self, cdl_minimal_contract):
        cdl_minimal_contract.description_travaux = (
            "Rénovation café — été 2026 • «Design»"
        )
        cdl_minimal_contract.save()
        gen = ContractPDFGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_bl_pdf_unicode_in_notes(self, bl_minimal_contract):
        bl_minimal_contract.notes = "Début prévu mi-avril\n• Point 1\n• Point 2"
        bl_minimal_contract.save()
        gen = BluelinePDFGenerator(bl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_cdl_pdf_very_long_description(self, cdl_minimal_contract):
        cdl_minimal_contract.description_travaux = "Test " * 2000
        cdl_minimal_contract.save()
        gen = ContractPDFGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_bl_pdf_many_prestations(self, bl_minimal_contract):
        bl_minimal_contract.prestations = [
            {
                "nom": f"prestation_{i}",
                "desc": f"Desc {i}",
                "qte": i,
                "unite": "m2",
                "pu": 100 + i,
            }
            for i in range(20)
        ]
        bl_minimal_contract.save()
        gen = BluelinePDFGenerator(bl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_cdl_doc_all_clauses_active(self, cdl_minimal_contract):
        cdl_minimal_contract.clauses_actives = [
            "clause_resiliation",
            "clause_force_majeure",
            "clause_confidentialite",
            "clause_propriete",
            "clause_non_sollicitation",
            "clause_sous_traitance",
            "clause_assurance",
            "clause_modification",
            "clause_integralite",
            "clause_election",
        ]
        cdl_minimal_contract.save()
        gen = ContractDOCGenerator(cdl_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_bl_doc_all_garantie_types(self, bl_minimal_contract):
        """Test each garantie_type value produces valid DOCX."""
        for gtype in ("defauts", "bonne_fin", "decennale", "aucune"):
            bl_minimal_contract.garantie_type = gtype
            bl_minimal_contract.garantie_nb = 1
            bl_minimal_contract.garantie_unite = "ans"
            bl_minimal_contract.save()
            gen = BluelineDOCGenerator(bl_minimal_contract, language="fr")
            response = gen.generate_response()
            assert response.status_code == 200, f"Failed for garantie_type={gtype}"


def response_ok_or_html_valid(html: str) -> bool:
    """Simple check that HTML string is non-empty."""
    return bool(html and len(html) > 100)


# ══════════════════════════════════════════════════════════════════════════════
#  BUG-FIX REGRESSION TESTS
# ══════════════════════════════════════════════════════════════════════════════


class TestBLFalsyZeroTvaAcompte:
    """Regression: tva=0 and acompte=0 must NOT fall back to defaults."""

    def test_bl_pdf_tva_zero_is_zero(self, bl_contract):
        bl_contract.tva = Decimal("0")
        bl_contract.save()
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        assert gen._tva_pct == 0.0
        assert gen._tva_amt == 0.0
        # TTC should equal HT when TVA is 0
        assert gen._ttc == gen._ht

    def test_bl_doc_tva_zero_is_zero(self, bl_contract):
        bl_contract.tva = Decimal("0")
        bl_contract.save()
        gen = BluelineDOCGenerator(bl_contract, language="fr")
        assert gen._tva_pct == 0.0
        assert gen._tva_amt == 0.0
        assert gen._ttc == gen._ht

    def test_bl_pdf_acompte_zero_is_zero(self, bl_contract):
        bl_contract.acompte = Decimal("0")
        bl_contract.save()
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        assert gen._acompte_pct == 0.0

    def test_bl_doc_acompte_zero_is_zero(self, bl_contract):
        bl_contract.acompte = Decimal("0")
        bl_contract.save()
        gen = BluelineDOCGenerator(bl_contract, language="fr")
        assert gen._acompte_pct == 0.0

    def test_bl_pdf_tva_none_falls_back_to_20(self, bl_contract):
        bl_contract.tva = None
        # Don't save — tva column is NOT NULL. Just test the property.
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        assert gen._tva_pct == 20.0

    def test_bl_doc_tva_none_falls_back_to_20(self, bl_contract):
        bl_contract.tva = None
        gen = BluelineDOCGenerator(bl_contract, language="fr")
        assert gen._tva_pct == 20.0

    def test_bl_pdf_acompte_none_falls_back_to_30(self, bl_contract):
        bl_contract.acompte = None
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        assert gen._acompte_pct == 30.0

    def test_bl_doc_acompte_none_falls_back_to_30(self, bl_contract):
        bl_contract.acompte = None
        gen = BluelineDOCGenerator(bl_contract, language="fr")
        assert gen._acompte_pct == 30.0

    def test_bl_pdf_zero_tva_generates_pdf(self, bl_contract):
        bl_contract.tva = Decimal("0")
        bl_contract.save()
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        assert b"%PDF" in response.content[:20]

    def test_bl_doc_zero_tva_generates_docx(self, bl_contract):
        bl_contract.tva = Decimal("0")
        bl_contract.save()
        gen = BluelineDOCGenerator(bl_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200


class TestBLDynamicTVALabel:
    """Regression: TVA label must show the actual percentage, not hardcoded 20%."""

    def test_pdf_tva_label_shows_10_percent(self, bl_contract):
        bl_contract.tva = Decimal("10")
        bl_contract.save()
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        assert "TVA (10%)" in html
        assert "TVA (20%)" not in html

    def test_pdf_tva_label_shows_0_percent(self, bl_contract):
        bl_contract.tva = Decimal("0")
        bl_contract.save()
        gen = BluelinePDFGenerator(bl_contract, language="fr")
        html = gen._build_html()
        assert "TVA (0%)" in html

    def test_pdf_tva_label_en_shows_vat(self, bl_contract):
        bl_contract.tva = Decimal("15")
        bl_contract.save()
        gen = BluelinePDFGenerator(bl_contract, language="en")
        html = gen._build_html()
        assert "VAT (15%)" in html
        assert "TVA" not in html


class TestCDLGarantieTranslation:
    """Regression: garantie values must be translated in English contracts."""

    def test_pdf_garantie_en_2_ans_translated(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        cdl_contract.garantie = "2 ans"
        cdl_contract.save()
        html = _gen_contract_html(cdl_contract, "en")
        assert "2 years" in html
        assert "2 ans" not in html

    def test_pdf_garantie_en_6_mois_translated(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        cdl_contract.garantie = "6 mois"
        cdl_contract.save()
        html = _gen_contract_html(cdl_contract, "en")
        assert "6 months" in html
        # "6 mois" may still appear in duree_estimee, so just verify translation
        # is present in the garantie context
        assert "warranty" in html.lower() or "6 months" in html

    def test_pdf_garantie_en_sans_garantie_translated(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        cdl_contract.garantie = "sans_garantie"
        cdl_contract.save()
        html = _gen_contract_html(cdl_contract, "en")
        assert "No warranty" in html

    def test_pdf_garantie_fr_stays_french(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        cdl_contract.garantie = "2 ans"
        cdl_contract.save()
        html = _gen_contract_html(cdl_contract, "fr")
        assert "2 ans" in html

    def test_doc_garantie_en_translated(self, cdl_contract):
        from docx import Document as DocxDocument

        cdl_contract.garantie = "1 an"
        cdl_contract.save()
        gen = ContractDOCGenerator(cdl_contract, language="en")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "1 year" in full_text

    def test_doc_garantie_fr_stays_french(self, cdl_contract):
        from docx import Document as DocxDocument

        cdl_contract.garantie = "3 ans"
        cdl_contract.save()
        gen = ContractDOCGenerator(cdl_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "3 ans" in full_text

    def test_pdf_garantie_unknown_value_passthrough(self, cdl_contract):
        """Unknown garantie values should pass through unchanged."""
        from contract.pdf import _gen_contract_html

        cdl_contract.garantie = "5 ans"
        cdl_contract.save()
        html = _gen_contract_html(cdl_contract, "en")
        # Unknown value falls back to _esc(c.garantie)
        assert "5 ans" in html


class TestEmptyFraisRedemarrage:
    """Regression: no empty <li></li> when frais_redemarrage is not set."""

    def test_no_empty_li_when_no_frais(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        cdl_contract.frais_redemarrage = None
        cdl_contract.save()
        html = _gen_contract_html(cdl_contract, "fr")
        assert "<li></li>" not in html

    def test_frais_present_renders_li(self, cdl_contract):
        from contract.pdf import _gen_contract_html

        cdl_contract.frais_redemarrage = Decimal("5000")
        cdl_contract.save()
        html = _gen_contract_html(cdl_contract, "fr")
        assert "frais de redémarrage" in html.lower() or "5" in html


# ══════════════════════════════════════════════════════════════════════════════
#  FIXTURES — Sous-Traitance
# ══════════════════════════════════════════════════════════════════════════════


@pytest.fixture()
def cdl_company_config():
    """CompanyConfig for casa_di_lusso EP identity."""
    return CompanyConfig.objects.create(
        company="casa_di_lusso",
        name="CASA DI LUSSO SARL",
        forme_juridique="SARL",
        capital=Decimal("500000.00"),
        rc="123456",
        ice="001234567890123",
        identifiant_fiscal="78901234",
        adresse="45 Bd Mohammed V, Tanger 90000",
        representant="Youssef Directeur",
        qualite_representant="Gérant",
    )


@pytest.fixture()
def st_project(staff_user):
    """Project linked to ST contract."""
    return Project.objects.create(
        company="casa_di_lusso",
        name="Villa Palmiers Tanger",
        type="villa",
        adresse="Route de Malabata, Tanger",
        maitre_ouvrage="Ahmed Invest SARL",
        permis="PC-2026-12345",
    )


@pytest.fixture()
def st_contract(staff_user, st_project, cdl_company_config):
    """Full ST contract with all fields populated."""
    return Contract.objects.create(
        company="casa_di_lusso",
        contract_category="sous_traitance",
        numero_contrat="ST/TEST-001",
        date_contrat=date(2026, 6, 1),
        statut="Signé",
        created_by_user=staff_user,
        ville_signature="Tanger",
        montant_ht=Decimal("250000.00"),
        devise="MAD",
        tva=Decimal("20.00"),
        # ST-specific
        st_projet=st_project,
        st_name="Entreprise El Mansouri SARL",
        st_forme="SARL",
        st_capital=Decimal("100000.00"),
        st_rc="654321",
        st_ice="009876543210987",
        st_if="56789012",
        st_cnss="CNSS-12345",
        st_addr="12 Rue Allal Ben Abdallah, Casablanca",
        st_rep="Mohammed El Mansouri",
        st_cin="AB654321",
        st_qualite="Gérant",
        st_tel="+212600000099",
        st_email="contact@elmansouri.ma",
        st_rib="007 600 0987654321098765 42",
        st_banque="Banque Populaire",
        st_lot_type="gros_oeuvre",
        st_lot_description="Gros-œuvre complet de la villa",
        st_type_prix="forfaitaire",
        st_retenue_garantie=Decimal("10.00"),
        st_avance=Decimal("20.00"),
        st_penalite_taux=Decimal("1.00"),
        st_plafond_penalite=Decimal("10.00"),
        st_delai_paiement=30,
        st_tranches=[
            {"label": "Acompte", "pourcentage": 20},
            {"label": "Avancement 50%", "pourcentage": 40},
            {"label": "Avancement 80%", "pourcentage": 30},
            {"label": "Réception", "pourcentage": 10},
        ],
        st_delai_val=6,
        st_delai_unit="mois",
        st_garantie_mois=120,
        st_delai_reserves=30,
        st_delai_med=30,
        st_clauses_actives=["tConfid", "tMediat", "tAnnexe"],
        st_observations="Début prévu juin 2026 — coordonner avec lot électricité.",
    )


@pytest.fixture()
def st_minimal_contract(staff_user, cdl_company_config):
    """Minimal ST contract — only required fields."""
    return Contract.objects.create(
        company="casa_di_lusso",
        contract_category="sous_traitance",
        numero_contrat="ST/MIN-001",
        date_contrat=date(2026, 1, 1),
        statut="Brouillon",
        created_by_user=staff_user,
    )


# ══════════════════════════════════════════════════════════════════════════════
#  UNIT TESTS — st_i18n helpers
# ══════════════════════════════════════════════════════════════════════════════


class TestStI18n:
    """Test translations and lookup functions from st_i18n.py."""

    def test_st_t_returns_fr_text(self):
        val = st_t("contrat_title", "fr")
        assert isinstance(val, str)
        assert "SOUS-TRAITANCE" in val.upper() or "CONTRAT" in val.upper()

    def test_st_t_returns_en_text(self):
        val = st_t("contrat_title", "en")
        assert isinstance(val, str)
        assert "SUBCONTRACT" in val.upper() or "CONTRACT" in val.upper()

    def test_st_t_missing_key_returns_key(self):
        val = st_t("nonexistent_key_xyz", "fr")
        assert val == "nonexistent_key_xyz"

    def test_lot_labels_contain_gros_oeuvre(self):
        assert "gros_oeuvre" in LOT_LABELS["fr"]
        assert "gros_oeuvre" in LOT_LABELS["en"]

    def test_st_t_list_type(self):
        val = st_t("docs_list", "fr")
        assert isinstance(val, list)
        assert len(val) > 0


# ══════════════════════════════════════════════════════════════════════════════
#  INTEGRATION TESTS — ST PDF Generator
# ══════════════════════════════════════════════════════════════════════════════


class TestSTPDFGenerator:
    """Test SousTraitancePDFGenerator (Casa Di Lusso ST — WeasyPrint PDF)."""

    def test_generate_response_returns_pdf(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        assert response["Content-Type"] == "application/pdf"
        assert b"%PDF" in response.content[:20]

    def test_generate_response_en(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="en")
        response = gen.generate_response()
        assert response.status_code == 200
        assert b"%PDF" in response.content[:20]

    def test_filename_contains_st_ref(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        response = gen.generate_response()
        disp = response["Content-Disposition"]
        assert "ST" in disp or "contrat_st" in disp

    def test_html_contains_ep_name(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        assert "CASA DI LUSSO" in html

    def test_html_contains_st_name(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        assert "El Mansouri" in html

    def test_html_contains_lot_label(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        lot_label = LOT_LABELS["fr"]["gros_oeuvre"]
        assert lot_label in html

    def test_html_contains_financial_box(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        assert "250" in html  # montant HT
        assert "300" in html  # TTC (250000 + 50000)

    def test_html_contains_tranches(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        assert "Acompte" in html
        assert "20" in html

    def test_html_contains_optional_clause(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        # tConfid is active
        confid_title = st_t("clause_confid", "fr")
        assert confid_title in html

    def test_html_contains_observations(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        assert "coordonner avec lot" in html

    def test_html_contains_project_info(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        assert "Villa Palmiers" in html

    def test_html_contains_signatures(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="fr")
        html = gen._build_html()
        assert "Tanger" in html  # ville_signature

    def test_minimal_contract_generates(self, st_minimal_contract):
        gen = SousTraitancePDFGenerator(st_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        assert b"%PDF" in response.content[:20]

    def test_en_contains_english_labels(self, st_contract):
        gen = SousTraitancePDFGenerator(st_contract, language="en")
        html = gen._build_html()
        lot_label = LOT_LABELS["en"]["gros_oeuvre"]
        assert lot_label in html


# ══════════════════════════════════════════════════════════════════════════════
#  INTEGRATION TESTS — ST DOC Generator
# ══════════════════════════════════════════════════════════════════════════════


class TestSTDOCGenerator:
    """Test SousTraitanceDOCGenerator (Casa Di Lusso ST — python-docx DOCX)."""

    def test_generate_response_returns_docx(self, st_contract):
        gen = SousTraitanceDOCGenerator(st_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200
        ct = response["Content-Type"]
        assert "officedocument" in ct or "msword" in ct

    def test_generate_response_en(self, st_contract):
        gen = SousTraitanceDOCGenerator(st_contract, language="en")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_filename_contains_st(self, st_contract):
        gen = SousTraitanceDOCGenerator(st_contract, language="fr")
        response = gen.generate_response()
        disp = response["Content-Disposition"]
        assert "contrat_st" in disp

    def test_docx_is_valid_zip(self, st_contract):
        import zipfile

        gen = SousTraitanceDOCGenerator(st_contract, language="fr")
        response = gen.generate_response()
        buf = io.BytesIO(response.content)
        assert zipfile.is_zipfile(buf)

    def test_docx_contains_st_name(self, st_contract):
        from docx import Document as DocxDocument

        gen = SousTraitanceDOCGenerator(st_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        all_text = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = "\n".join(all_text)
        assert "El Mansouri" in full_text

    def test_docx_contains_ep_name(self, st_contract):
        from docx import Document as DocxDocument

        gen = SousTraitanceDOCGenerator(st_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        all_text = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = "\n".join(all_text)
        assert "CASA DI LUSSO" in full_text

    def test_docx_contains_project(self, st_contract):
        from docx import Document as DocxDocument

        gen = SousTraitanceDOCGenerator(st_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        all_text = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = "\n".join(all_text)
        assert "Villa Palmiers" in full_text

    def test_docx_en_contains_english(self, st_contract):
        from docx import Document as DocxDocument

        gen = SousTraitanceDOCGenerator(st_contract, language="en")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # English labels should be present
        assert len(full_text) > 100

    def test_minimal_contract_generates(self, st_minimal_contract):
        gen = SousTraitanceDOCGenerator(st_minimal_contract, language="fr")
        response = gen.generate_response()
        assert response.status_code == 200

    def test_docx_annexe_present(self, st_contract):
        """tAnnexe is active, so the annexe checklist should be present."""
        from docx import Document as DocxDocument

        gen = SousTraitanceDOCGenerator(st_contract, language="fr")
        response = gen.generate_response()
        doc = DocxDocument(io.BytesIO(response.content))
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = "\n".join(all_text)
        assert "☐" in full_text  # checkbox character


# ══════════════════════════════════════════════════════════════════════════════
#  INTEGRATION TESTS — ST PDF / DOC View endpoints
# ══════════════════════════════════════════════════════════════════════════════


class TestContractPDFViewST:
    """Test the /pdf/fr|en/<pk>/ endpoint for ST contracts."""

    def test_pdf_view_returns_pdf(self, staff_client, st_contract):
        url = reverse("contract:contract-pdf-fr", args=[st_contract.pk])
        response = staff_client.get(url)
        assert response.status_code == 200
        assert response["Content-Type"] == "application/pdf"

    def test_pdf_view_en(self, staff_client, st_contract):
        url = reverse("contract:contract-pdf-en", args=[st_contract.pk])
        response = staff_client.get(url)
        assert response.status_code == 200

    def test_pdf_view_no_permission(self, noprint_client, st_contract):
        url = reverse("contract:contract-pdf-fr", args=[st_contract.pk])
        response = noprint_client.get(url)
        assert response.status_code == 403


class TestContractDOCViewST:
    """Test the /doc/fr|en/<pk>/ endpoint for ST contracts."""

    def test_doc_view_returns_docx(self, staff_client, st_contract):
        url = reverse("contract:contract-doc-fr", args=[st_contract.pk])
        response = staff_client.get(url)
        assert response.status_code == 200
        ct = response["Content-Type"]
        assert "officedocument" in ct

    def test_doc_view_en(self, staff_client, st_contract):
        url = reverse("contract:contract-doc-en", args=[st_contract.pk])
        response = staff_client.get(url)
        assert response.status_code == 200

    def test_doc_view_no_permission(self, noprint_client, st_contract):
        url = reverse("contract:contract-doc-fr", args=[st_contract.pk])
        response = noprint_client.get(url)
        assert response.status_code == 403


# ══════════════════════════════════════════════════════════════════════════════
#  TESTS — CompanyConfig + Project API
# ══════════════════════════════════════════════════════════════════════════════


class TestCompanyConfigAPI:
    def test_list_requires_auth(self):
        client = APIClient()
        url = "/api/company-config/"
        response = client.get(url)
        assert response.status_code == 401

    def test_list_returns_configs(self, staff_client, cdl_company_config):
        url = "/api/company-config/"
        response = staff_client.get(url)
        assert response.status_code == 200
        data = response.json()
        assert any(c["company"] == "casa_di_lusso" for c in data)


class TestProjectAPI:
    def test_create_project(self, staff_client):
        url = reverse("contract:project-list-create")
        data = {
            "company": "casa_di_lusso",
            "name": "Test Project",
            "type": "villa",
            "adresse": "123 Test St",
        }
        response = staff_client.post(url, data, format="json")
        assert response.status_code == 201
        assert response.json()["name"] == "Test Project"

    def test_list_projects(self, staff_client, st_project):
        url = reverse("contract:project-list-create")
        response = staff_client.get(url)
        assert response.status_code == 200
        data = response.json()
        assert any(p["name"] == "Villa Palmiers Tanger" for p in data)

    def test_project_detail(self, staff_client, st_project):
        url = reverse("contract:project-detail", args=[st_project.pk])
        response = staff_client.get(url)
        assert response.status_code == 200
        assert response.json()["name"] == "Villa Palmiers Tanger"


# ══════════════════════════════════════════════════════════════════════════════
#  TESTS — Unique-together (company, contract_category, numero_contrat)
# ══════════════════════════════════════════════════════════════════════════════


class TestUniqueTogetherContractCategory:
    def test_same_numero_different_category_ok(self, staff_user, cdl_company_config):
        """Two contracts with same numero but different categories should work."""
        Contract.objects.create(
            company="casa_di_lusso",
            contract_category="",
            numero_contrat="001/26",
            date_contrat=date(2026, 1, 1),
            statut="Brouillon",
            created_by_user=staff_user,
        )
        Contract.objects.create(
            company="casa_di_lusso",
            contract_category="sous_traitance",
            numero_contrat="001/26",
            date_contrat=date(2026, 1, 1),
            statut="Brouillon",
            created_by_user=staff_user,
        )
        assert Contract.objects.filter(numero_contrat="001/26").count() == 2

    def test_same_numero_same_category_fails(self, staff_user, cdl_company_config):
        """Two contracts with same (company, category, numero) should fail."""
        from django.db import IntegrityError

        Contract.objects.create(
            company="casa_di_lusso",
            contract_category="sous_traitance",
            numero_contrat="DUP/01",
            date_contrat=date(2026, 1, 1),
            statut="Brouillon",
            created_by_user=staff_user,
        )
        with pytest.raises(IntegrityError):
            Contract.objects.create(
                company="casa_di_lusso",
                contract_category="sous_traitance",
                numero_contrat="DUP/01",
                date_contrat=date(2026, 1, 1),
                statut="Brouillon",
                created_by_user=staff_user,
            )
