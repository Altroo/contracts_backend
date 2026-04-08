# i18n: skip-file — bilingual document generator; FR+EN content is intentional
from io import BytesIO

from django.http import HttpResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from .pdf import _fmt_date, _fmt_amt, _format_penalite_retard, _is_societe
from .i18n import (
    TYPELABEL,
    TYPE_BIEN_LABELS,
    MODE_NAMES,
    CTYPES_DISPLAY,
    CONFID_LABELS,
    QUALITE_LABELS,
    GARANTIE_LABELS,
)

DARK = RGBColor(0x0F, 0x0F, 0x1A)
GOLD = RGBColor(0xB8, 0x97, 0x3A)
GOLD_PALE_HEX = "F7F0E0"
CREAM_HEX = "FAF6EC"
CREAM2_HEX = "F9F5EC"
INK = RGBColor(0x1C, 0x1C, 0x2E)
RED = RGBColor(0xB5, 0x34, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x6B, 0x6B, 0x80)
BORDER_HEX = "E2D9C8"
DARK_HEX = "0F0F1A"


def _cell_bg(cell, hex_color: str):
    """Set table cell background fill."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _para_bg(para, hex_color: str):
    """Set paragraph shading (full-width highlight bar)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _para_border_left(para, color: str = "B8973A", size: str = "12"):
    """Add a left border to a paragraph (for highlight / art-title style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _para_border_bottom(para, color: str = "E2D9C8", size: str = "4"):
    """Add a bottom border to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)


def _para_border_top(para, color: str = "EEEEEE", size: str = "4"):
    """Add a top border to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    pBdr.append(top)


def _para_box_borders(para, color: str = "E2D9C8", size: str = "4"):
    """Wrap paragraph in a full box border (all 4 sides)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def _cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders.  Each arg is a (color, size) tuple or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(val[1]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val[0])
            tcBorders.append(el)
        else:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_width(cell, cm_val):
    """Set explicit cell width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm_val * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _set_row_height(row, pt_val):
    """Set minimum row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(pt_val * 20)))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


def _run_shd(run, hex_color: str):
    """Apply a character-level shading (badge background) to a run."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


class ContractDOCGenerator:
    """Generate a python-docx DOCX for a Contract instance, matching the PDF output."""

    def __init__(self, contract, language: str = "fr"):
        self.contract = contract
        self.language = language
        self.fr = language == "fr"
        self.doc = Document()
        self._configure_page()

    def _configure_page(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

        # Explicit Normal-style defaults
        normal = self.doc.styles["Normal"]
        normal.font.name = "Inter"
        normal.font.size = Pt(9)
        normal.font.color.rgb = INK
        normal.paragraph_format.space_after = Pt(2)
        normal.paragraph_format.line_spacing = 1.75

    def _add_empty(self, pt=2):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(pt)
        return p

    def _add_text(
        self, text, size=Pt(9), bold=False, color=INK, align=None, space_after=Pt(3)
    ):
        p = self.doc.add_paragraph()
        if align:
            p.alignment = align
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.font.size = size
        run.bold = bold
        run.font.color.rgb = color
        return p, run

    def _runs(self, p, *segs, size=Pt(9), color=INK):
        """Add formatted text runs to a paragraph.

        Each *seg* is either a plain ``str`` (normal text) or a tuple:
        ``(text, bold)`` or ``(text, bold, italic)``.
        """
        for seg in segs:
            if isinstance(seg, str):
                text, bold, italic = seg, False, False
            else:
                text = seg[0]
                bold = seg[1] if len(seg) > 1 else False
                italic = seg[2] if len(seg) > 2 else False
            r = p.add_run(text)
            r.font.size = size
            r.font.color.rgb = color
            if bold:
                r.bold = True
            if italic:
                r.italic = True

    def _p(self, *segs, space_after=Pt(3)):
        """Body paragraph with mixed bold/plain segments."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = space_after
        self._runs(p, *segs)
        return p

    def _bullet(self, *segs, indent=Pt(18)):
        """Unordered list item (bullet)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = indent
        r = p.add_run("\u2022 ")
        r.font.size = Pt(9)
        r.font.color.rgb = INK
        self._runs(p, *segs)
        return p

    def _subbullet(self, *segs):
        """Nested bullet (deeper indent)."""
        return self._bullet(*segs, indent=Pt(36))

    def _nbullet(self, idx, *segs):
        """Ordered (numbered) list item."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(f"{idx}. ")
        r.font.size = Pt(9)
        r.font.color.rgb = INK
        self._runs(p, *segs)
        return p

    def _highlight(self, *segs):
        """Gold-pale highlight box with left gold border."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _para_bg(p, GOLD_PALE_HEX)
        _para_border_left(p, "B8973A", "12")
        self._runs(p, *segs)
        return p

    def _warning(self, *segs):
        """Red warning box with pink background and border."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _para_bg(p, "FDF0EC")
        _para_box_borders(p, "f0c4bc", "4")
        self._runs(p, *segs, color=RED)
        return p

    def _sub_title(self, text):
        """Sub-section title (e.g. 12.1 – …)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = INK
        return p

    def _divider(self):
        """Horizontal separator line."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _para_border_bottom(p, "EDE8DA", "4")

    def _services_box(self, services):
        """Service tags box (cream background, gold tag badges)."""
        if not services:
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        _para_bg(p, CREAM2_HEX)
        _para_box_borders(p, "EDE8DA", "4")
        title_r = p.add_run(
            ("SERVICES CONVENUS" if self.fr else "AGREED SERVICES") + "\n"
        )
        title_r.font.size = Pt(7)
        title_r.bold = True
        title_r.font.color.rgb = GOLD
        title_r.font.all_caps = True
        for i, s in enumerate(services):
            if i > 0:
                sp = p.add_run("  ")
                sp.font.size = Pt(8)
            sr = p.add_run(f" {s} ")
            sr.font.size = Pt(7.5)
            sr.bold = True
            sr.font.color.rgb = GOLD
            _run_shd(sr, DARK_HEX)

    def _plan_grid(self, cards):
        """Planning info cards rendered as a mini-table.

        *cards*: list of ``{"label": ..., "val": ...}`` dicts.
        """
        if not cards:
            return
        table = self.doc.add_table(rows=2, cols=len(cards))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_table_borders(table)
        for j, card in enumerate(cards):
            # Label row
            lbl_cell = table.rows[0].cells[j]
            _cell_bg(lbl_cell, CREAM2_HEX)
            _cell_borders(
                lbl_cell,
                top=("EDE8DA", 4),
                left=("EDE8DA", 4),
                right=("EDE8DA", 4),
            )
            lp = lbl_cell.paragraphs[0]
            lp.paragraph_format.space_before = Pt(3)
            lp.paragraph_format.space_after = Pt(0)
            lr = lp.add_run(card["label"].upper())
            lr.font.size = Pt(7)
            lr.bold = True
            lr.font.color.rgb = MUTED
            lr.font.all_caps = True
            # Value row
            val_cell = table.rows[1].cells[j]
            _cell_bg(val_cell, CREAM2_HEX)
            _cell_borders(
                val_cell,
                bottom=("EDE8DA", 4),
                left=("EDE8DA", 4),
                right=("EDE8DA", 4),
            )
            vp = val_cell.paragraphs[0]
            vp.paragraph_format.space_before = Pt(0)
            vp.paragraph_format.space_after = Pt(3)
            vr = vp.add_run(card["val"])
            vr.font.size = Pt(9)
            vr.bold = True
            vr.font.color.rgb = INK
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _pay_table(self, tranches, montant_ttc, tot_pct, devise):
        """Payment schedule table (dark header, striped rows, gold footer)."""
        headers = ["#", "Description", "%", "Montant" if self.fr else "Amount"]
        data_rows = []
        for i, tr in enumerate(tranches):
            pct = float(tr.get("pourcentage", 0))
            amt = montant_ttc * pct / 100
            lbl = tr.get("label", f"Tranche {i + 1}")
            data_rows.append([str(i + 1), lbl, f"{int(pct)}%", _fmt_amt(amt, devise)])

        ncols = 4
        table = self.doc.add_table(rows=0, cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_table_borders(table)

        # Header row
        hrow = table.add_row()
        for j, h in enumerate(headers):
            cell = hrow.cells[j]
            _cell_bg(cell, DARK_HEX)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(h)
            r.font.size = Pt(7.5)
            r.bold = True
            r.font.color.rgb = GOLD
            r.font.all_caps = True
            if j == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for i, row_data in enumerate(data_rows):
            row = table.add_row()
            is_even = i % 2 == 1
            for j, val in enumerate(row_data):
                cell = row.cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(val)
                r.font.size = Pt(8.5)
                r.font.color.rgb = INK
                _cell_borders(cell, bottom=(BORDER_HEX, 2))
                if is_even:
                    _cell_bg(cell, CREAM2_HEX)
                if j == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer (TOTAL) row
        frow = table.add_row()
        # Merge first two cells for "TOTAL" label
        frow.cells[0].merge(frow.cells[1])
        for j in range(ncols):
            cell = frow.cells[j]
            _cell_bg(cell, GOLD_PALE_HEX)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if j == 0:
                r = p.add_run("TOTAL")
                r.font.size = Pt(8.5)
                r.bold = True
                r.font.color.rgb = INK
            elif j == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"{int(tot_pct)}%")
                r.font.size = Pt(8.5)
                r.bold = True
                r.font.color.rgb = INK
            elif j == 3:
                r = p.add_run(_fmt_amt(montant_ttc, devise))
                r.font.size = Pt(8.5)
                r.bold = True
                r.font.color.rgb = INK

        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _art(self, num, title):
        """Render the gold article header bar (ART. XX  TITLE)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        _para_bg(p, GOLD_PALE_HEX)
        _para_border_left(p, "B8973A", "12")
        badge = p.add_run(f" ART. {num} ")
        badge.font.size = Pt(7.5)
        badge.bold = True
        badge.font.color.rgb = GOLD
        _run_shd(badge, DARK_HEX)
        tr = p.add_run(f"  {title}")
        tr.font.size = Pt(8)
        tr.bold = True
        tr.font.color.rgb = DARK
        tr.font.all_caps = True

    def _next_art(self, title):
        """Increment article counter and render the title bar."""
        self._n += 1
        self._art(str(self._n).zfill(2), title)

    def _add_top_strip(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        _para_border_bottom(p, "B8973A", "12")

    def _add_header(self, c, ref, date_str, version_str, confid_label, ville):
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_table_borders(table)

        # Left: logo + info
        left = table.rows[0].cells[0]
        _set_cell_width(left, 10)
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        r1 = lp.add_run("CASA DI ")
        r1.font.size = Pt(18)
        r1.bold = True
        r1.font.color.rgb = DARK
        r2 = lp.add_run("LUSSO")
        r2.font.size = Pt(18)
        r2.bold = True
        r2.font.color.rgb = GOLD

        tag = left.add_paragraph()
        tag.paragraph_format.space_before = Pt(2)
        tag.paragraph_format.space_after = Pt(2)
        tr = tag.add_run(
            "Design \u00b7 "
            + ("Travaux" if self.fr else "Works")
            + " \u00b7 "
            + ("Ameublement" if self.fr else "Furnishing")
            + " \u00b7 "
            + ("Maroc" if self.fr else "Morocco")
        )
        tr.font.size = Pt(7.5)
        tr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        tr.font.all_caps = True

        info = left.add_paragraph()
        info.paragraph_format.space_before = Pt(0)
        info.paragraph_format.space_after = Pt(0)
        for line in [
            "SARL \u00b7 RC 143377 \u00b7 ICE 003389356000001",
            "IF 60116256 \u00b7 CNSS 5001474",
            "Route N1, Al Moustakbal Roundabout \u2013 Tanger",
        ]:
            ir = info.add_run(line + "\n")
            ir.font.size = Pt(7)
            ir.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

        if c.responsable_projet:
            rp_label = "Chef de projet" if self.fr else "Project Manager"
            rp = info.add_run(f"{rp_label} : {c.responsable_projet}")
            rp.font.size = Pt(7)
            rp.font.color.rgb = MUTED
            rp.bold = True

        # Right: ref badge + dates
        right = table.rows[0].cells[1]
        _set_cell_width(right, 7)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        ref_p = right.paragraphs[0]
        ref_p.paragraph_format.space_after = Pt(4)
        ref_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_p.add_run(f"  {ref}  ")
        ref_run.font.size = Pt(9)
        ref_run.bold = True
        ref_run.font.color.rgb = GOLD
        _run_shd(ref_run, DARK_HEX)

        date_p = right.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(0)
        for label, val in [
            ("Date :", date_str),
            ("Version :", version_str),
            ("Classe :" if self.fr else "Class :", confid_label),
        ]:
            lr = date_p.add_run(label + " ")
            lr.font.size = Pt(8)
            lr.bold = True
            lr.font.color.rgb = INK
            vr = date_p.add_run(val + "\n")
            vr.font.size = Pt(8)
            vr.font.color.rgb = MUTED

        vr2 = date_p.add_run(f"{ville}, {'Maroc' if self.fr else 'Morocco'}")
        vr2.font.size = Pt(8)
        vr2.font.color.rgb = MUTED

        # Separator line under header
        sep = self.doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(8)
        sep.paragraph_format.space_after = Pt(10)
        _para_border_bottom(sep, BORDER_HEX, "2")

    def _add_confidential(self):
        txt = (
            "DOCUMENT CONFIDENTIEL \u2013 USAGE EXCLUSIF DES PARTIES SIGNATAIRES"
            if self.fr
            else "CONFIDENTIAL DOCUMENT \u2013 FOR SIGNATORY PARTIES ONLY"
        )
        p, r = self._add_text(
            txt,
            size=Pt(7),
            color=RGBColor(0xCC, 0xCC, 0xCC),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(10),
        )
        r.font.all_caps = True

    def _add_title_block(self, ctype_display):
        # Gold mini-bar
        bar = self.doc.add_paragraph()
        bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bar.paragraph_format.space_after = Pt(0)
        bar.paragraph_format.space_before = Pt(6)
        br = bar.add_run("\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501")
        br.font.size = Pt(5)
        br.font.color.rgb = GOLD

        # Main title
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(8)
        title_p.paragraph_format.space_after = Pt(2)
        _para_border_top(title_p, "EEEEEE", "4")
        tr = title_p.add_run(
            "CONTRAT DE PRESTATIONS DE SERVICES" if self.fr else "SERVICE AGREEMENT"
        )
        tr.font.size = Pt(14)
        tr.bold = True
        tr.font.color.rgb = DARK

        # Subtitle
        sub, sr = self._add_text(
            "CASA DI LUSSO SARL \u00b7 RC 143377 \u00b7 Tanger, "
            + ("Maroc" if self.fr else "Morocco"),
            size=Pt(8),
            color=RGBColor(0xAA, 0xAA, 0xAA),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(2),
        )
        sr.font.all_caps = True

        # Type badge
        badge, br2 = self._add_text(
            f"  {ctype_display}  ",
            size=Pt(7.5),
            color=GOLD,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(4),
        )
        br2.font.all_caps = True

        # Bottom border
        bot = self.doc.add_paragraph()
        bot.paragraph_format.space_before = Pt(2)
        bot.paragraph_format.space_after = Pt(10)
        _para_border_bottom(bot, "EEEEEE", "4")

    def _add_parties(
        self,
        c,
        client_nom,
        client_cin,
        client_qualite,
        client_adresse,
        client_tel,
        client_email,
    ):
        bp, br = self._add_text(
            "ENTRE LES SOUSSIGN\u00c9S :" if self.fr else "BETWEEN THE UNDERSIGNED:",
            size=Pt(9),
            bold=True,
            color=RGBColor(0x66, 0x66, 0x66),
            space_after=Pt(6),
        )
        br.italic = True

        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_table_borders(table)

        prest = table.rows[0].cells[0]
        _cell_bg(prest, DARK_HEX)
        _set_cell_width(prest, 8.5)

        lbl_p = prest.paragraphs[0]
        lbl_p.paragraph_format.space_after = Pt(4)
        lr = lbl_p.add_run("LE PRESTATAIRE" if self.fr else "THE SERVICE PROVIDER")
        lr.font.size = Pt(7)
        lr.bold = True
        lr.font.color.rgb = GOLD
        lr.font.all_caps = True

        prest_lines = [
            ("CASA DI LUSSO SARL", True),
            ("RC : 143377 \u00b7 ICE : 003389356000001", False),
            ("IF : 60116256 \u00b7 CNSS : 5001474", False),
            ("Route N1, Al Moustakbal Roundabout", False),
            ("Tanger \u2013 " + ("Maroc" if self.fr else "Morocco"), False),
            ("", False),
            (
                (
                    "Ci-apr\u00e8s d\u00e9nomm\u00e9e \u00ab Le Prestataire \u00bb"
                    if self.fr
                    else "Hereinafter referred to as \u00abThe Service Provider\u00bb"
                ),
                False,
            ),
        ]
        for text, bold in prest_lines:
            pp = prest.add_paragraph()
            pp.paragraph_format.space_before = Pt(0)
            pp.paragraph_format.space_after = Pt(0)
            rr = pp.add_run(text)
            rr.font.size = Pt(8.5)
            rr.font.color.rgb = WHITE if bold else RGBColor(0xCC, 0xCC, 0xCC)
            rr.bold = bold
            if text.startswith("Ci-apr\u00e8s") or text.startswith("Hereinafter"):
                rr.italic = True
                rr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        # ── Client (cream bg with gold border) ──
        client = table.rows[0].cells[1]
        _cell_bg(client, CREAM_HEX)
        _cell_borders(
            client,
            top=("B8973A", 4),
            bottom=("B8973A", 4),
            left=("B8973A", 4),
            right=("B8973A", 4),
        )
        _set_cell_width(client, 8.5)

        clbl_p = client.paragraphs[0]
        clbl_p.paragraph_format.space_after = Pt(4)
        clr = clbl_p.add_run("LE CLIENT" if self.fr else "THE CLIENT")
        clr.font.size = Pt(7)
        clr.bold = True
        clr.font.color.rgb = GOLD
        clr.font.all_caps = True

        cin_lbl = "CIN / ICE" if self.fr else "ID / Passport"
        qual_lbl = "Qualit\u00e9" if self.fr else "Status"
        client_lines = [
            (client_nom, True),
            (f"{cin_lbl} : {client_cin}", False),
            (f"{qual_lbl} : {client_qualite}", False),
            (client_adresse, False),
            (client_tel, False),
        ]
        if client_email:
            client_lines.append((client_email, False))
        client_lines.append(("", False))
        ci_client = (
            "Ci-apr\u00e8s d\u00e9nomm\u00e9(e) \u00ab Le Client \u00bb"
            if self.fr
            else "Hereinafter referred to as \u00abThe Client\u00bb"
        )
        client_lines.append((ci_client, False))

        for text, bold in client_lines:
            cp = client.add_paragraph()
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            cr = cp.add_run(text)
            cr.font.size = Pt(8.5)
            cr.font.color.rgb = DARK if bold else INK
            cr.bold = bold
            if text.startswith("Ci-apr\u00e8s") or text.startswith("Hereinafter"):
                cr.italic = True
                cr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        self._add_empty(8)

    def _add_signatures(self, ville, date_str, sig_name, sig_role):
        sep = self.doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(14)
        sep.paragraph_format.space_after = Pt(8)
        _para_border_bottom(sep, BORDER_HEX, "4")

        # Header row
        sig_header = self.doc.add_table(rows=1, cols=2)
        _remove_table_borders(sig_header)
        lp = sig_header.rows[0].cells[0].paragraphs[0]
        lr = lp.add_run("\u270d SIGNATURES")
        lr.font.size = Pt(8)
        lr.bold = True
        lr.font.color.rgb = DARK
        lr.font.all_caps = True

        rp = sig_header.rows[0].cells[1].paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_date_txt = (
            f"Fait \u00e0 {ville}, le {date_str}"
            if self.fr
            else f"Done in {ville}, on {date_str}"
        )
        rr = rp.add_run(sig_date_txt)
        rr.font.size = Pt(8)
        rr.font.color.rgb = MUTED

        self._add_empty(4)

        # Signature boxes
        sig_table = self.doc.add_table(rows=4, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_table_borders(sig_table)

        lbl_client = "LE CLIENT" if self.fr else "THE CLIENT"
        note_client = (
            "Lu et approuv\u00e9 \u2013 Bon pour accord"
            if self.fr
            else "Read and approved \u2013 Agreement"
        )
        note_prest = "Signature & Cachet" if self.fr else "Signature & Stamp"
        dir_gen = "Direction G\u00e9n\u00e9rale" if self.fr else "General Management"
        for j, (label, note, name, role) in enumerate(
            [
                (lbl_client, note_client, sig_name, sig_role),
                (
                    "CASA DI LUSSO SARL",
                    note_prest,
                    dir_gen,
                    "CASA DI LUSSO SARL",
                ),
            ]
        ):
            cell0 = sig_table.rows[0].cells[j]
            _cell_borders(
                cell0, top=(BORDER_HEX, 4), left=(BORDER_HEX, 4), right=(BORDER_HEX, 4)
            )
            p0 = cell0.paragraphs[0]
            r0 = p0.add_run(label)
            r0.font.size = Pt(7)
            r0.bold = True
            r0.font.color.rgb = GOLD
            r0.font.all_caps = True

            cell1 = sig_table.rows[1].cells[j]
            _cell_borders(cell1, left=(BORDER_HEX, 4), right=(BORDER_HEX, 4))
            p1 = cell1.paragraphs[0]
            r1 = p1.add_run(note)
            r1.font.size = Pt(8)
            r1.italic = True
            r1.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

            cell2 = sig_table.rows[2].cells[j]
            _cell_borders(cell2, left=(BORDER_HEX, 4), right=(BORDER_HEX, 4))
            _set_row_height(sig_table.rows[2], 50)

            cell3 = sig_table.rows[3].cells[j]
            _cell_borders(
                cell3,
                bottom=(BORDER_HEX, 4),
                left=(BORDER_HEX, 4),
                right=(BORDER_HEX, 4),
            )
            p3 = cell3.paragraphs[0]
            rn = p3.add_run(name)
            rn.font.size = Pt(8.5)
            rn.bold = True
            rn.font.color.rgb = INK
            rr2 = p3.add_run(f"\n{role}")
            rr2.font.size = Pt(7.5)
            rr2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        # Initials row
        self._add_empty(6)
        init_p = self.doc.add_paragraph()
        init_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        init_p.paragraph_format.space_before = Pt(6)
        init_p.paragraph_format.space_after = Pt(6)
        _para_bg(init_p, CREAM2_HEX)
        _para_box_borders(init_p, "EDE8DA", "4")
        lbl_provider = "Prestataire" if self.fr else "Service Provider"
        initials_txt = (
            "Paraphes des parties (chaque page)"
            if self.fr
            else "Initials of parties (each page)"
        )
        ir = init_p.add_run(
            "Client  ________________    "
            f"{initials_txt}    "
            f"________________  {lbl_provider}"
        )
        ir.font.size = Pt(8)
        ir.font.color.rgb = MUTED

    def _add_footer(self, ref, version_str, confid_label):
        sep = self.doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(10)
        sep.paragraph_format.space_after = Pt(4)
        _para_border_bottom(sep, BORDER_HEX, "2")

        ft = self.doc.add_table(rows=1, cols=3)
        _remove_table_borders(ft)

        fp = ft.rows[0].cells[0].paragraphs[0]
        fr1 = fp.add_run("CASA DI LUSSO")
        fr1.font.size = Pt(10)
        fr1.bold = True
        fr1.font.color.rgb = GOLD

        cp = ft.rows[0].cells[1].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(
            f"{ref} \u00b7 RC 143377 \u00b7 Tanger, {'Maroc' if self.fr else 'Morocco'}"
        )
        cr.font.size = Pt(7)
        cr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

        rp2 = ft.rows[0].cells[2].paragraphs[0]
        rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp2.add_run(f"{version_str} \u00b7 {confid_label}")
        rr.font.size = Pt(7)
        rr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    def _build_articles(self):
        """Build every article body using direct python-docx calls.

        The numbering follows pdf.py: counter starts at 1 and is
        pre-incremented, so the first article rendered is ART. 01.
        Conditional articles only increment the counter when present,
        keeping numbering contiguous.
        """
        c = self.contract
        clauses = c.clauses_actives if isinstance(c.clauses_actives, list) else []
        self._n = 0
        fr = self.fr
        lang = self.language

        type_label = TYPELABEL[lang].get(c.type_contrat or "", c.type_contrat or "")
        type_bien_lbl = TYPE_BIEN_LABELS[lang].get(
            c.type_bien or "autre", c.type_bien or ""
        )
        surface_str = f" \u2013 {c.surface}\u00a0m\u00b2" if c.surface else ""
        services = c.services if isinstance(c.services, list) else []
        adresse = c.adresse_travaux or (
            "\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026"
            "\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026"
        )

        self._next_art("OBJET DU CONTRAT" if fr else "SCOPE OF AGREEMENT")
        self._p(
            (
                "Le pr\u00e9sent contrat a pour objet la r\u00e9alisation par "
                if fr
                else "This contract covers the execution by "
            ),
            ("CASA DI LUSSO SARL", True),
            (
                " (ci-apr\u00e8s \u00ab\u202fLe Prestataire\u202f\u00bb) de travaux "
                "et prestations de type "
                if fr
                else " (hereinafter \u00abThe Service Provider\u00bb) of works and services of type "
            ),
            (type_label, True),
            (
                " dans le bien immobilier du Client."
                if fr
                else " on the Client\u2019s property."
            ),
        )
        self._p(
            ("Type de bien\u202f: " if fr else "Property type: ", True),
            f"{type_bien_lbl}{surface_str}",
        )
        self._p(
            ("Adresse du chantier\u202f: " if fr else "Site address: ", True), adresse
        )
        self._services_box(services)
        if c.description_travaux:
            self._highlight(
                (
                    (
                        "D\u00e9tail des travaux\u202f:\n"
                        if fr
                        else "Works description:\n"
                    ),
                    True,
                ),
                c.description_travaux,
            )
        self._p(
            (
                "Tout travail ou prestation non mentionn\u00e9(e) express\u00e9ment au "
                "pr\u00e9sent article est "
                if fr
                else "Any work or service not expressly mentioned in "
                "this article is "
            ),
            ("formellement exclu(e)" if fr else "expressly excluded", True),
            (
                " du pr\u00e9sent contrat et ne pourra \u00eatre r\u00e9clam\u00e9(e) "
                "sans nouvel accord contractuel \u00e9crit."
                if fr
                else " from this contract and cannot be claimed "
                "without a new written contractual agreement."
            ),
        )

        if c.date_debut or c.duree_estimee:
            d_str = _fmt_date(c.date_debut) if c.date_debut else ""
            dur = c.duree_estimee or ""
            cards = []
            if d_str:
                cards.append(
                    {
                        "label": ("Date de d\u00e9but" if fr else "Start Date"),
                        "val": d_str,
                    }
                )
            if dur:
                cards.append(
                    {
                        "label": (
                            "Dur\u00e9e estim\u00e9e" if fr else "Estimated Duration"
                        ),
                        "val": dur,
                    }
                )
            if d_str and dur:
                cards.append(
                    {
                        "label": (
                            "Livraison pr\u00e9visionnelle"
                            if fr
                            else "Expected Delivery"
                        ),
                        "val": ("Selon planning" if fr else "Per schedule"),
                    }
                )
            self._next_art(
                "PLANNING & D\u00c9LAIS D\u2019EX\u00c9CUTION"
                if fr
                else "SCHEDULE & EXECUTION TIMELINE"
            )
            self._plan_grid(cards)
            self._p(
                (
                    "Le planning communiqu\u00e9 est "
                    if fr
                    else "The communicated schedule is "
                ),
                ("estimatif" if fr else "indicative", True),
                (
                    " et non contractuellement opposable. Le Prestataire mettra tout "
                    "en \u0153uvre pour respecter les d\u00e9lais convenus, sous r\u00e9serve "
                    "de la bonne ex\u00e9cution des obligations du Client (paiements, "
                    "validations, acc\u00e8s chantier). Tout retard imputable au Client "
                    "entra\u00eene automatiquement un d\u00e9calage du planning sans "
                    "possibilit\u00e9 d\u2019invoquer une quelconque p\u00e9nalit\u00e9 "
                    "\u00e0 l\u2019encontre du Prestataire."
                    if fr
                    else " and not contractually binding. The Service Provider will make every effort "
                    "to meet agreed timelines, subject to the Client\u2019s timely fulfillment "
                    "of obligations (payments, approvals, site access). Any delay attributable "
                    "to the Client automatically pushes the schedule back with no right "
                    "to invoke penalties against the Service Provider."
                ),
            )

        self._next_art(
            "FORCE CONTRACTUELLE ET PREUVE" if fr else "CONTRACTUAL FORCE AND EVIDENCE"
        )
        self._p(
            (
                "Le pr\u00e9sent contrat, le devis sign\u00e9 par les deux parties, les "
                "plans et visuels approuv\u00e9s par \u00e9crit, ainsi que toutes les "
                "confirmations officielles \u00e9crites, constituent "
                if fr
                else "This contract, the quotation signed by both parties, the plans "
                "and visuals approved in writing, and all official written "
                "confirmations constitute "
            ),
            (
                (
                    "la base l\u00e9gale exclusive et exhaustive"
                    if fr
                    else "the sole and exhaustive legal basis"
                ),
                True,
            ),
            (
                " des engagements r\u00e9ciproques des parties."
                if fr
                else " of the parties\u2019 mutual commitments."
            ),
        )
        self._bullet(
            (
                "Tout \u00e9change oral, r\u00e9union informelle, discussion "
                "t\u00e9l\u00e9phonique ou message non valid\u00e9 officiellement par "
                "voie \u00e9crite ou \u00e9lectronique "
                if fr
                else "Any oral exchange, informal meeting, telephone discussion or message not "
                "officially validated in writing or electronically "
            ),
            (
                (
                    "n\u2019a aucune valeur contractuelle"
                    if fr
                    else "has no contractual value"
                ),
                True,
            ),
            ".",
        )
        self._bullet(
            (
                "Les messages WhatsApp, SMS, e-mails et tout document sign\u00e9 ou "
                "valid\u00e9 par le Client "
                if fr
                else "WhatsApp messages, SMS, emails and any document signed or "
                "validated by the Client "
            ),
            ("font foi" if fr else "shall serve as evidence", True),
            (
                " et peuvent \u00eatre produits comme preuves en cas de litige."
                if fr
                else " and may be produced in any dispute."
            ),
        )
        self._bullet(
            (
                "Toute modification au pr\u00e9sent contrat doit faire l\u2019objet "
                "d\u2019un "
                if fr
                else "Any modification to this contract must be formalized by a "
            ),
            (
                "avenant \u00e9crit sign\u00e9" if fr else "written amendment signed",
                True,
            ),
            (
                " par les deux parties pour \u00eatre valide."
                if fr
                else " by both parties to be valid."
            ),
        )
        self._highlight(
            (
                "En cas de contradiction entre le pr\u00e9sent contrat et tout autre "
                "document, "
                if fr
                else "In case of contradiction between this contract and any other document, "
            ),
            (
                "le pr\u00e9sent contrat prime" if fr else "this contract prevails",
                True,
            ),
            (
                ", sauf avenant sign\u00e9 post\u00e9rieurement."
                if fr
                else ", unless a later signed amendment exists."
            ),
        )

        self._next_art(
            "OBLIGATIONS DU PRESTATAIRE" if fr else "SERVICE PROVIDER OBLIGATIONS"
        )
        self._p(
            "Le Prestataire s\u2019engage \u00e0\u202f:"
            if fr
            else "The Service Provider undertakes to:"
        )
        self._bullet(
            (
                "Ex\u00e9cuter les travaux conform\u00e9ment aux "
                if fr
                else "Execute works in compliance with "
            ),
            ("r\u00e8gles de l\u2019art" if fr else "professional standards", True),
            (
                " et aux normes techniques marocaines en vigueur."
                if fr
                else " and applicable Moroccan technical regulations."
            ),
        )
        self._bullet(
            ("Mettre \u00e0 disposition une " if fr else "Provide a "),
            (
                (
                    "\u00e9quipe qualifi\u00e9e et exp\u00e9riment\u00e9e"
                    if fr
                    else "qualified and experienced team"
                ),
                True,
            ),
            (
                ", sous la supervision d\u2019un responsable de chantier d\u00e9sign\u00e9."
                if fr
                else ", under the supervision of a designated site manager."
            ),
        )
        self._bullet(
            (
                "Respecter les choix esth\u00e9tiques et techniques "
                if fr
                else "Respect aesthetic and technical choices "
            ),
            ("valid\u00e9s par \u00e9crit" if fr else "validated in writing", True),
            (" par le Client." if fr else " by the Client."),
        )
        self._bullet(
            (
                "Informer le Client de tout al\u00e9a ou difficult\u00e9 susceptible "
                "d\u2019affecter le planning ou le budget."
                if fr
                else "Inform the Client of any hazard or difficulty likely to affect "
                "the schedule or budget."
            ),
        )
        self._bullet(
            ("Veiller au maintien d\u2019un chantier " if fr else "Maintain a "),
            ("propre et s\u00e9curis\u00e9" if fr else "clean and safe worksite", True),
            (
                " tout au long de l\u2019ex\u00e9cution."
                if fr
                else " throughout execution."
            ),
        )
        self._bullet(
            (
                "Assurer la coordination entre les diff\u00e9rents corps de m\u00e9tier "
                "intervenants."
                if fr
                else "Coordinate between the various trades involved."
            ),
        )
        if c.responsable_projet:
            self._p(
                (
                    (
                        "Responsable de projet d\u00e9sign\u00e9\u202f: "
                        if fr
                        else "Designated project manager: "
                    ),
                    True,
                ),
                c.responsable_projet,
            )
        if c.architecte:
            self._p(
                (
                    (
                        "Architecte / Designer associ\u00e9\u202f: "
                        if fr
                        else "Associated Architect / Designer: "
                    ),
                    True,
                ),
                c.architecte,
            )
        self._highlight(
            (
                "Le Prestataire est tenu \u00e0 une "
                if fr
                else "The Service Provider is bound by a "
            ),
            ("obligation de moyen" if fr else "best-efforts obligation", True),
            (
                " et non de r\u00e9sultat pour les travaux dont les mat\u00e9riaux, "
                "choix esth\u00e9tiques et contraintes sp\u00e9cifiques sont impos\u00e9s "
                "par le Client."
                if fr
                else " (not a results obligation) for works where materials, "
                "aesthetic choices and specific constraints are imposed by the Client."
            ),
        )

        acces_extra = f" ({c.conditions_acces})" if c.conditions_acces else ""
        self._next_art(
            "OBLIGATIONS RENFORC\u00c9ES DU CLIENT"
            if fr
            else "REINFORCED CLIENT OBLIGATIONS"
        )
        self._p(
            (
                "Le Client s\u2019engage contractuellement et irr\u00e9vocablement "
                "\u00e0\u202f:"
                if fr
                else "The Client contractually and irrevocably undertakes to:"
            ),
        )
        self._bullet(
            ("Respecter" if fr else "Respect", True),
            (
                " l\u2019ensemble des membres de l\u2019\u00e9quipe du Prestataire et "
                "leurs d\u00e9cisions techniques professionnelles."
                if fr
                else " all members of the Service Provider\u2019s team and their professional technical decisions."
            ),
        )
        self._bullet(
            ("Fournir toutes les " if fr else "Provide all "),
            ("validations \u00e9crites" if fr else "written validations", True),
            (
                " (plans, choix de mat\u00e9riaux, modifications) dans des d\u00e9lais "
                "raisonnables n\u2019exc\u00e9dant pas 72h, sauf accord contraire."
                if fr
                else " (plans, material choices, modifications) within reasonable timeframes "
                "not exceeding 72 hours, unless otherwise agreed."
            ),
        )
        self._bullet(
            ("R\u00e9gler les " if fr else "Make "),
            (
                (
                    "paiements aux dates convenues"
                    if fr
                    else "payments on the agreed dates"
                ),
                True,
            ),
            (
                " dans l\u2019\u00e9ch\u00e9ancier contractuel."
                if fr
                else " in the contractual schedule."
            ),
        )
        self._bullet(
            (
                "N\u2019exercer aucune pression, menace, intimidation ni contrainte "
                "sur les \u00e9quipes, les sous-traitants ou les fournisseurs du "
                "Prestataire."
                if fr
                else "Not exert any pressure, threats, intimidation, or constraint "
                "on the teams, subcontractors, or suppliers."
            ),
        )
        self._bullet(
            ("Garantir un " if fr else "Guarantee "),
            (
                (
                    "acc\u00e8s libre, s\u00e9curis\u00e9 et continu"
                    if fr
                    else "free, secure, and continuous site access"
                ),
                True,
            ),
            (
                f" au chantier selon les conditions convenues.{acces_extra}"
                if fr
                else f" under the agreed conditions.{acces_extra}"
            ),
        )
        self._bullet(
            *(
                [
                    ("Assumer l\u2019enti\u00e8re responsabilit\u00e9", True),
                    " de toutes les d\u00e9cisions esth\u00e9tiques et techniques qu\u2019il "
                    "a valid\u00e9es par \u00e9crit.",
                ]
                if fr
                else [
                    "Take ",
                    ("full responsibility", True),
                    " for all aesthetic and technical decisions validated in writing.",
                ]
            ),
        )
        self._bullet(
            ("S\u2019abstenir de faire intervenir " if fr else "Refrain from having "),
            (
                (
                    "toute autre entreprise ou artisan"
                    if fr
                    else "any other company or tradesperson"
                ),
                True,
            ),
            (
                " sur le chantier sans accord pr\u00e9alable \u00e9crit du Prestataire."
                if fr
                else " on site without prior written consent from the Service Provider."
            ),
        )
        self._bullet(
            (
                "Assurer la pr\u00e9sence d\u2019un repr\u00e9sentant habilit\u00e9 \u00e0 "
                if fr
                else "Ensure the presence of an authorized representative empowered to "
            ),
            ("prendre des d\u00e9cisions" if fr else "make decisions", fr),
            (
                " lors des r\u00e9unions de chantier planifi\u00e9es."
                if fr
                else " at scheduled site meetings."
            ),
        )
        self._warning(
            (
                "\u26a0\ufe0f En cas de non-respect de l\u2019une de ces obligations, le "
                "Prestataire se r\u00e9serve le droit de "
                if fr
                else "\u26a0\ufe0f In case of non-compliance, the "
                "Service Provider reserves the right to "
            ),
            (
                (
                    "suspendre imm\u00e9diatement les travaux"
                    if fr
                    else "immediately suspend works"
                ),
                True,
            ),
            (
                " sans pr\u00e9avis ni indemnit\u00e9, et de facturer les frais de "
                "red\u00e9marrage pr\u00e9vus au pr\u00e9sent contrat."
                if fr
                else " without notice or compensation, and to invoice the restart fees "
                "provided in this contract."
            ),
        )

        self._next_art(
            "D\u00c9LAIS, RETARDS ET AL\u00c9AS"
            if fr
            else "TIMELINES, DELAYS AND CONTINGENCIES"
        )
        self._p(
            (
                "Le Prestataire ne saurait \u00eatre tenu responsable de tout retard "
                "r\u00e9sultant des causes ci-apr\u00e8s list\u00e9es, qui constituent "
                "autant de "
                if fr
                else "The Service Provider cannot be held liable for any delay resulting "
                "from the following causes, which constitute automatic "
            ),
            (
                (
                    "cas de suspension automatique"
                    if fr
                    else "suspension of the contractual timeline"
                ),
                True,
            ),
            (" du d\u00e9lai contractuel\u202f:" if fr else ":"),
        )
        self._bullet(
            (
                "Retard ou d\u00e9faut de paiement de la part du Client au-del\u00e0 "
                "du d\u00e9lai tol\u00e9r\u00e9\u202f;"
                if fr
                else "Late or missing payment from the Client beyond the tolerated delay;"
            ),
        )
        self._bullet(
            (
                "Retard du Client dans la validation des plans, choix de mat\u00e9riaux "
                "ou modifications\u202f;"
                if fr
                else "Client delay in validating plans, material choices, or modifications;"
            ),
        )
        self._bullet(
            (
                "Changement de design ou de prestation d\u00e9cid\u00e9 apr\u00e8s "
                "validation initiale\u202f;"
                if fr
                else "Design or scope changes decided after initial validation;"
            ),
        )
        self._bullet(
            (
                "Retard d\u2019importation, rupture de stock chez les fournisseurs ou "
                "d\u00e9lais douaniers\u202f;"
                if fr
                else "Import delays, supplier stock shortages, or customs issues;"
            ),
        )
        self._bullet(
            (
                "Cas de force majeure (pand\u00e9mie, catastrophe naturelle, conflit, "
                "d\u00e9cision administrative)\u202f;"
                if fr
                else "Force majeure (pandemic, natural disaster, conflict, administrative order);"
            ),
        )
        self._bullet(
            (
                "Probl\u00e8me structurel impr\u00e9vu d\u00e9couvert en cours de "
                "chantier\u202f;"
                if fr
                else "Unforeseen structural issues discovered during works;"
            ),
        )
        self._bullet(
            (
                "Intervention de tiers non mandat\u00e9s par le Prestataire\u202f;"
                if fr
                else "Unauthorized third-party interventions;"
            ),
        )
        self._bullet(
            (
                "D\u00e9faut d\u2019acc\u00e8s au chantier."
                if fr
                else "Denied or blocked site access."
            ),
        )
        self._highlight(
            (
                "Tout retard imputable au Client "
                if fr
                else "Any delay attributable to the Client "
            ),
            ("repousse automatiquement" if fr else "automatically extends", True),
            (
                " le planning d\u2019un nombre de jours \u00e9quivalent, sans "
                "possibilit\u00e9 pour le Client d\u2019invoquer une quelconque "
                "p\u00e9nalit\u00e9 \u00e0 l\u2019encontre du Prestataire."
                if fr
                else " the timeline by the equivalent number of days, with no right for "
                "the Client to invoke penalties against the Service Provider."
            ),
        )

        self._next_art(
            "MODIFICATIONS, AVENANTS ET TRAVAUX SUPPL\u00c9MENTAIRES"
            if fr
            else "MODIFICATIONS, AMENDMENTS AND ADDITIONAL WORKS"
        )
        self._p(
            (
                "Le p\u00e9rim\u00e8tre des travaux d\u00e9fini \u00e0 "
                "l\u2019Article\u00a01 est ferme et d\u00e9finitif \u00e0 compter de la "
                "date de signature. Toute modification ult\u00e9rieure est soumise aux "
                "r\u00e8gles suivantes\u202f:"
                if fr
                else "The works scope defined in Article 1 is firm and final from the "
                "signing date. Any subsequent modification is subject to the following rules:"
            ),
        )
        self._nbullet(
            1,
            (
                "La demande de modification doit \u00eatre formul\u00e9e "
                if fr
                else "Any modification request must be made "
            ),
            ("par \u00e9crit" if fr else "in writing", True),
            (
                " (email ou document sign\u00e9)."
                if fr
                else " (email or signed document)."
            ),
        )
        self._nbullet(
            2,
            (
                "Le Prestataire \u00e9tablira un "
                if fr
                else "The Service Provider will issue a "
            ),
            ("devis compl\u00e9mentaire" if fr else "supplementary quotation", True),
            (
                " dans un d\u00e9lai raisonnable."
                if fr
                else " within a reasonable timeframe."
            ),
        )
        self._nbullet(
            3,
            (
                "La modification ne sera ex\u00e9cut\u00e9e qu\u2019apr\u00e8s "
                if fr
                else "The modification will only be executed after "
            ),
            (
                (
                    "signature du devis additionnel"
                    if fr
                    else "signing the additional quotation"
                ),
                True,
            ),
            (
                " et versement de l\u2019acompte correspondant."
                if fr
                else " and payment of the corresponding deposit."
            ),
        )
        self._nbullet(
            4,
            ("Tout changement entra\u00eene une " if fr else "Any change entails a "),
            ("r\u00e9vision du planning" if fr else "revision of the timeline", True),
            (
                " \u00e0 la hausse ou \u00e0 la baisse selon la nature des travaux."
                if fr
                else ", up or down, depending on the nature of works."
            ),
        )
        self._nbullet(
            5,
            ("Les " if fr else ""),
            ("moins-values" if fr else "Reductions", True),
            (
                " (r\u00e9ductions de prestations) n\u2019ouvrent pas droit au "
                "remboursement des sommes d\u00e9j\u00e0 vers\u00e9es."
                if fr
                else " in scope do not entitle the Client to a refund of amounts already paid."
            ),
        )
        self._warning(
            "\u26a0\ufe0f ",
            ("Aucune modification verbale" if fr else "No verbal modification", True),
            (
                " ne sera prise en compte. Tout travail suppl\u00e9mentaire "
                "ex\u00e9cut\u00e9 sans avenant sign\u00e9 sera factur\u00e9 sur la base "
                "du tarif horaire en vigueur du Prestataire."
                if fr
                else " will be considered. Any additional work performed without a signed "
                "amendment will be invoiced at the Service Provider\u2019s current hourly rate."
            ),
        )

        montant_ht = float(c.montant_ht or 0)
        tva_pct = float(c.tva or 0)
        montant_tva = montant_ht * tva_pct / 100
        montant_ttc = montant_ht + montant_tva
        devise = c.devise or "MAD"
        mode_label = MODE_NAMES[lang].get(
            c.mode_paiement_texte or "virement", c.mode_paiement_texte or ""
        )
        tranches = c.tranches if isinstance(c.tranches, list) else []
        tot_pct = sum(float(t.get("pourcentage", 0)) for t in tranches)
        rib_str = f" \u2014 {c.rib}" if c.rib else ""
        delai_ret = c.delai_retard if c.delai_retard is not None else 5
        penalite = _format_penalite_retard(c, lang)

        self._next_art(
            "CONDITIONS FINANCI\u00c8RES ET IRR\u00c9VOCABILIT\u00c9 DES PAIEMENTS"
            if fr
            else "FINANCIAL CONDITIONS AND PAYMENT IRREVOCABILITY"
        )

        # Montant HT paragraph (with TVA info)
        if tva_pct == 0:
            self._p(
                (
                    (
                        "Montant total HT\u202f: "
                        if fr
                        else "Total amount (excl. tax): "
                    ),
                    True,
                ),
                f"{_fmt_amt(montant_ht, devise)} "
                + ("(non assujetti TVA)" if fr else "(VAT-exempt)"),
            )
        else:
            self._p(
                (
                    (
                        "Montant total HT\u202f: "
                        if fr
                        else "Total amount (excl. tax): "
                    ),
                    True,
                ),
                f"{_fmt_amt(montant_ht, devise)} \u2014 "
                f"{'TVA' if fr else 'VAT'}\u00a0{int(tva_pct)}%\u202f: "
                f"{_fmt_amt(montant_tva, devise)} \u2014 ",
                (
                    f"{'TTC' if fr else 'Incl. tax'}\u202f: {_fmt_amt(montant_ttc, devise)}",
                    True,
                ),
            )

        self._p(
            (
                "Les paiements sont \u00e9chelonn\u00e9s selon le tableau ci-apr\u00e8s. "
                "Chaque versement est une condition suspensive \u00e0 la poursuite des "
                "travaux\u202f:"
                if fr
                else "Payments are scheduled as follows. Each payment is a condition for "
                "the continuation of works:"
            ),
        )
        self._pay_table(tranches, montant_ttc, tot_pct, devise)
        self._p(
            (("Mode de r\u00e8glement\u202f: " if fr else "Payment method: "), True),
            f"{mode_label}{rib_str}",
        )

        # Payment terms bullets
        self._bullet(
            ("Tout acompte vers\u00e9 est " if fr else "Any deposit paid is "),
            (
                (
                    "d\u00e9finitivement acquis au Prestataire"
                    if fr
                    else "definitively retained by the Service Provider"
                ),
                True,
            ),
            (
                " et non remboursable, sauf faute lourde d\u00fbment constat\u00e9e par "
                "d\u00e9cision judiciaire d\u00e9finitive."
                if fr
                else " and non-refundable, except in case of serious fault established by "
                "final court decision."
            ),
        )
        self._bullet(
            (
                "En cas de retard de paiement sup\u00e9rieur \u00e0 "
                if fr
                else "In case of payment delay exceeding "
            ),
            (f"{delai_ret} " + ("jours" if fr else "calendar days"), True),
            (" calendaires\u202f:" if fr else ":"),
        )
        self._subbullet(
            ("Les travaux sont " if fr else "Works are "),
            ("imm\u00e9diatement suspendus" if fr else "immediately suspended", True),
            (
                " et les \u00e9quipes retir\u00e9es du chantier\u202f;"
                if fr
                else " and teams removed from site;"
            ),
        )
        self._subbullet(
            ("Des " if fr else ""),
            ("p\u00e9nalit\u00e9s de retard" if fr else "Late payment penalties", True),
            (" de " if fr else " of "),
            (penalite, True),
            (
                " de retard sont automatiquement applicables sur les sommes "
                "dues\u202f;"
                if fr
                else " automatically apply on overdue amounts;"
            ),
        )
        if c.frais_redemarrage:
            self._subbullet(
                ("Des " if fr else ""),
                ("frais de red\u00e9marrage" if fr else "Restart fees", True),
                (" d\u2019un montant de " if fr else " of "),
                (f"{_fmt_amt(float(c.frais_redemarrage), devise)}", True),
                (
                    " seront factur\u00e9s avant toute reprise des travaux\u202f;"
                    if fr
                    else " will be invoiced before any resumption of works;"
                ),
            )
        self._subbullet(
            (
                "Le planning est automatiquement r\u00e9vis\u00e9 sans droit \u00e0 "
                "indemnit\u00e9 pour le Client."
                if fr
                else "The schedule is automatically revised with no right to compensation "
                "for the Client."
            ),
        )

        self._next_art(
            "LIMITATION DE RESPONSABILIT\u00c9" if fr else "LIMITATION OF LIABILITY"
        )
        self._p(
            (
                "La responsabilit\u00e9 du Prestataire est express\u00e9ment "
                if fr
                else "The Service Provider\u2019s liability is expressly "
            ),
            (
                (
                    "limit\u00e9e et exclusivement engag\u00e9e"
                    if fr
                    else "limited and exclusively engaged"
                ),
                True,
            ),
            (
                " pour les travaux qu\u2019il a lui-m\u00eame r\u00e9alis\u00e9s. "
                "Le Prestataire est formellement "
                if fr
                else " for works it has itself carried out. The Service Provider is formally "
            ),
            (
                (
                    "d\u00e9gag\u00e9 de toute responsabilit\u00e9"
                    if fr
                    else "released from all liability"
                ),
                True,
            ),
            (" pour\u202f:" if fr else " for:"),
        )
        self._bullet(
            ("Les " if fr else "Pre-existing or in-progress "),
            ("d\u00e9fauts structurels" if fr else "structural defects", True),
            (
                " pr\u00e9existants ou d\u00e9couverts en cours de chantier\u202f;"
                if fr
                else ";"
            ),
        )
        self._bullet(
            ("Les " if fr else ""),
            ("infiltrations" if fr else "Infiltrations", True),
            (
                " et probl\u00e8mes provenant d\u2019\u00e9l\u00e9ments ext\u00e9rieurs "
                "au p\u00e9rim\u00e8tre du contrat\u202f;"
                if fr
                else " and issues from elements outside the contract scope;"
            ),
        )
        self._bullet(
            ("Les " if fr else ""),
            (
                (
                    "mat\u00e9riaux, produits ou \u00e9quipements"
                    if fr
                    else "Materials, products, or equipment"
                ),
                True,
            ),
            (
                " fournis par le Client ou impos\u00e9s par lui\u202f;"
                if fr
                else " supplied or imposed by the Client;"
            ),
        )
        self._bullet(
            ("Les travaux r\u00e9alis\u00e9s par " if fr else "Works carried out by "),
            (
                (
                    "d\u2019autres entreprises ou artisans"
                    if fr
                    else "other companies or tradespeople"
                ),
                True,
            ),
            (" intervenant sur le chantier\u202f;" if fr else " on site;"),
        )
        self._bullet(
            (
                "Les d\u00e9g\u00e2ts r\u00e9sultant d\u2019un "
                if fr
                else "Damage resulting from "
            ),
            ("usage non conforme" if fr else "improper use", True),
            (
                ", de modifications, ou d\u2019un d\u00e9faut d\u2019entretien "
                "apr\u00e8s livraison\u202f;"
                if fr
                else ", modifications, or lack of maintenance after delivery;"
            ),
        )
        self._bullet(
            ("Les cons\u00e9quences des " if fr else "Consequences of "),
            (
                (
                    "d\u00e9cisions esth\u00e9tiques ou techniques"
                    if fr
                    else "aesthetic or technical decisions"
                ),
                True,
            ),
            (" valid\u00e9es par le Client." if fr else " validated by the Client."),
        )
        self._warning(
            (
                "La responsabilit\u00e9 totale du Prestataire, toutes causes "
                "confondues, est "
                if fr
                else "The Service Provider\u2019s total liability, for all causes, is "
            ),
            (
                (
                    "strictement plafonn\u00e9e au montant total TTC effectivement "
                    "encaiss\u00e9"
                    if fr
                    else "strictly capped at the total amount actually received"
                ),
                True,
            ),
            (
                " dans le cadre du pr\u00e9sent contrat."
                if fr
                else " under this contract."
            ),
        )

        delai_res = c.delai_reserves if c.delai_reserves is not None else 7
        self._next_art(
            "R\u00c9CEPTION ET LIVRAISON DES TRAVAUX"
            if fr
            else "ACCEPTANCE AND DELIVERY OF WORKS"
        )
        self._p(
            (
                (
                    "R\u00e9ception provisoire\u202f: "
                    if fr
                    else "Provisional acceptance: "
                ),
                True,
            ),
            (
                "Elle est effectu\u00e9e en pr\u00e9sence physique et obligatoire du "
                "Client (ou de son repr\u00e9sentant mandat\u00e9 par \u00e9crit). Un "
                if fr
                else "Carried out with the Client (or their duly mandated representative) "
                "present. A "
            ),
            (
                (
                    "proc\u00e8s-verbal de r\u00e9ception"
                    if fr
                    else "formal acceptance report"
                ),
                True,
            ),
            (
                " sera \u00e9tabli et sign\u00e9 par les deux parties."
                if fr
                else " will be drafted and signed by both parties."
            ),
        )
        self._p(
            ("Toute " if fr else "Any "),
            ("r\u00e9serve" if fr else "reservations", True),
            (
                " (malfa\u00e7on, non-conformit\u00e9) doit \u00eatre formul\u00e9e de "
                "mani\u00e8re pr\u00e9cise et par \u00e9crit dans un d\u00e9lai de "
                if fr
                else " (defects, non-conformity) must be precisely stated in writing within "
            ),
            (f"{delai_res} " + ("jours ouvr\u00e9s" if fr else "working days"), True),
            (
                " suivant la r\u00e9ception provisoire."
                if fr
                else " of provisional acceptance."
            ),
        )
        self._bullet(
            (
                "Les r\u00e9serves mentionn\u00e9es feront l\u2019objet d\u2019une "
                if fr
                else "Stated reservations will be subject to a "
            ),
            ("reprise contradictoire" if fr else "jointly supervised correction", True),
            (
                " dans un d\u00e9lai raisonnable."
                if fr
                else " within a reasonable timeframe."
            ),
        )
        self._bullet(
            (
                f"Pass\u00e9 le d\u00e9lai de {delai_res} jours sans r\u00e9serve "
                "\u00e9crite, les travaux sont r\u00e9put\u00e9s "
                if fr
                else f"After {delai_res} days with no written reservation, works are deemed "
            ),
            (
                (
                    "r\u00e9ceptionn\u00e9s d\u00e9finitivement et sans r\u00e9serve"
                    if fr
                    else "definitively accepted without reservation"
                ),
                True,
            ),
            ".",
        )
        self._bullet(
            ("La r\u00e9ception vaut " if fr else "Acceptance marks the "),
            ("point de d\u00e9part" if fr else "start date", True),
            (" du d\u00e9lai de garantie." if fr else " of the warranty period."),
        )
        self._highlight(
            (
                "L\u2019absence du Client \u00e0 la r\u00e9ception, sans "
                "d\u00e9l\u00e9gation \u00e9crite pr\u00e9alable, vaut "
                if fr
                else "The Client\u2019s absence at acceptance, without prior written "
                "delegation, constitutes "
            ),
            ("acceptation tacite" if fr else "tacit acceptance", True),
            (" des travaux sans r\u00e9serve." if fr else " without reservation."),
        )

        lang_key = "fr" if fr else "en"
        garantie = (
            GARANTIE_LABELS.get(lang_key, {}).get(c.garantie, c.garantie)
            if c.garantie
            else ("1 an" if fr else "1 year")
        )
        self._next_art(
            "GARANTIE ET APR\u00c8S-LIVRAISON" if fr else "WARRANTY AND POST-DELIVERY"
        )
        self._p(
            ("Le Prestataire accorde une " if fr else "The Service Provider grants a "),
            (
                ("garantie contractuelle de " if fr else "contractual warranty of ")
                + f"{garantie}",
                True,
            ),
            (
                " sur les travaux r\u00e9alis\u00e9s dans le cadre du pr\u00e9sent "
                "contrat, \u00e0 compter de la date de r\u00e9ception d\u00e9finitive."
                if fr
                else " on all works performed under this contract, starting from the date "
                "of final acceptance."
            ),
        )
        self._p(
            ("Cette garantie couvre les " if fr else "This warranty covers "),
            ("d\u00e9fauts d\u2019ex\u00e9cution" if fr else "execution defects", True),
            (
                " directement imputables aux travaux du Prestataire, \u00e0 "
                "l\u2019exclusion de\u202f:"
                if fr
                else " directly attributable to the Service Provider\u2019s works, excluding:"
            ),
        )
        self._bullet(
            (
                "L\u2019usure normale li\u00e9e \u00e0 l\u2019utilisation\u202f;"
                if fr
                else "Normal wear from use;"
            ),
        )
        self._bullet(
            (
                "Les d\u00e9gradations caus\u00e9es par le Client, ses locataires ou "
                "des tiers\u202f;"
                if fr
                else "Damage caused by the Client, tenants, or third parties;"
            ),
        )
        self._bullet(
            (
                "Les interventions de tiers sur les travaux garantis\u202f;"
                if fr
                else "Third-party interventions on warranted works;"
            ),
        )
        self._bullet(
            (
                "Un d\u00e9faut d\u2019entretien ou une utilisation non conforme \u00e0 "
                "la destination des travaux\u202f;"
                if fr
                else "Lack of maintenance or improper use;"
            ),
        )
        self._bullet(
            (
                "Les probl\u00e8mes structurels non li\u00e9s aux travaux "
                "r\u00e9alis\u00e9s."
                if fr
                else "Structural issues unrelated to the works performed."
            ),
        )
        self._p(
            (
                "La mise en jeu de la garantie requiert une "
                if fr
                else "Invoking the warranty requires a "
            ),
            (
                "notification \u00e9crite" if fr else "detailed written notification",
                True,
            ),
            (
                " d\u00e9taill\u00e9e adress\u00e9e au Prestataire avec photos "
                "\u00e0 l\u2019appui."
                if fr
                else " sent to the Service Provider with supporting photos."
            ),
        )

        if "c-comportement" in clauses:
            self._next_art(
                "CLAUSE DE COMPORTEMENT ET PROTECTION DES \u00c9QUIPES"
                if fr
                else "CONDUCT AND TEAM PROTECTION CLAUSE"
            )
            self._p(
                (
                    "Le Client s\u2019engage formellement \u00e0 maintenir des "
                    "relations "
                    if fr
                    else "The Client formally undertakes to maintain "
                ),
                (
                    (
                        "respectueuses et professionnelles"
                        if fr
                        else "respectful and professional"
                    ),
                    True,
                ),
                (
                    " avec l\u2019ensemble du personnel, des sous-traitants et des "
                    "fournisseurs du Prestataire."
                    if fr
                    else " relations with all staff, subcontractors and suppliers of the Service Provider."
                ),
            )
            self._p(
                (
                    "Constituent des motifs de "
                    if fr
                    else "The following constitute grounds for "
                ),
                (
                    (
                        "r\u00e9siliation imm\u00e9diate"
                        if fr
                        else "immediate termination"
                    ),
                    True,
                ),
                (
                    " et sans indemnit\u00e9 pour le Prestataire\u202f:"
                    if fr
                    else " without compensation for the Service Provider:"
                ),
            )
            self._bullet(
                (
                    "Tout comportement abusif, humiliant, discriminatoire ou "
                    "mena\u00e7ant envers les \u00e9quipes\u202f;"
                    if fr
                    else "Any abusive, humiliating, discriminatory, or threatening behavior toward teams;"
                ),
            )
            self._bullet(
                (
                    "Toute pression ou intimidation visant \u00e0 obtenir des "
                    "prestations non contractuelles\u202f;"
                    if fr
                    else "Any pressure or intimidation to obtain non-contractual services;"
                ),
            )
            self._bullet(
                (
                    "Toute ing\u00e9rence non autoris\u00e9e dans les d\u00e9cisions "
                    "techniques du Prestataire\u202f;"
                    if fr
                    else "Any unauthorized interference in the Service Provider\u2019s technical decisions;"
                ),
            )
            self._bullet(
                (
                    "Tout d\u00e9nigrement public du Prestataire ou de ses "
                    "\u00e9quipes."
                    if fr
                    else "Any public disparagement of the Service Provider or its teams."
                ),
            )
            self._warning(
                (
                    "En cas de r\u00e9siliation pour comportement abusif, "
                    if fr
                    else "In case of termination for abusive conduct, "
                ),
                (
                    (
                        "toutes les sommes d\u00e9j\u00e0 vers\u00e9es restent "
                        "d\u00e9finitivement acquises"
                        if fr
                        else "all amounts already paid remain definitively retained"
                    ),
                    True,
                ),
                (
                    " au Prestataire, et les travaux en cours seront factur\u00e9s "
                    "au prorata."
                    if fr
                    else " by the Service Provider, and works in progress will be invoiced pro-rata."
                ),
            )

        if "c-prop-intel" in clauses:
            self._next_art(
                "PROPRI\u00c9T\u00c9 INTELLECTUELLE" if fr else "INTELLECTUAL PROPERTY"
            )
            self._p(
                (
                    "L\u2019ensemble des cr\u00e9ations r\u00e9alis\u00e9es dans le "
                    "cadre du pr\u00e9sent contrat \u2014 incluant sans limitation les "
                    if fr
                    else "All creations produced under this contract \u2014 including without limitation "
                ),
                (
                    (
                        "plans, dessins, moodboards, visuels, concepts "
                        "d\u2019am\u00e9nagement, id\u00e9es cr\u00e9atives, prototypes "
                        "et \u00e9tudes"
                        if fr
                        else "plans, drawings, moodboards, visuals, layout concepts, creative "
                        "ideas, prototypes and studies"
                    ),
                    True,
                ),
                (" \u2014 demeurent la " if fr else " \u2014 remain the "),
                (
                    (
                        "propri\u00e9t\u00e9 intellectuelle exclusive et "
                        "inali\u00e9nable"
                        if fr
                        else "exclusive and inalienable intellectual property"
                    ),
                    True,
                ),
                " de CASA DI LUSSO SARL." if fr else " of CASA DI LUSSO SARL.",
            )
            self._p(
                (
                    "Le Client b\u00e9n\u00e9ficie d\u2019une "
                    if fr
                    else "The Client is granted a "
                ),
                (
                    (
                        "licence d\u2019utilisation personnelle"
                        if fr
                        else "personal use license"
                    ),
                    True,
                ),
                (
                    " pour le bien objet du contrat uniquement. Il lui est "
                    "formellement interdit de\u202f:"
                    if fr
                    else " for the property covered by this contract only. The Client is "
                    "strictly prohibited from:"
                ),
            )
            self._bullet(
                (
                    "Reproduire, adapter ou diffuser ces cr\u00e9ations \u00e0 des "
                    "fins commerciales ou publicitaires\u202f;"
                    if fr
                    else "Reproducing, adapting, or distributing these creations for "
                    "commercial or advertising purposes;"
                ),
            )
            self._bullet(
                (
                    "Transmettre ou c\u00e9der ces documents \u00e0 des tiers\u202f;"
                    if fr
                    else "Transferring or assigning these documents to third parties;"
                ),
            )
            self._bullet(
                (
                    "Utiliser ces cr\u00e9ations pour un bien immobilier "
                    "diff\u00e9rent de celui vis\u00e9 au contrat\u202f;"
                    if fr
                    else "Using these creations for any property other than the one covered "
                    "by this contract;"
                ),
            )
            self._bullet(
                (
                    "D\u00e9poser ces cr\u00e9ations \u00e0 titre de marque, brevet ou "
                    "dessin industriel."
                    if fr
                    else "Registering these creations as a trademark, patent, or industrial design."
                ),
            )
            self._p(
                (
                    "Toute violation est susceptible d\u2019engager la "
                    if fr
                    else "Any violation may engage the Client\u2019s "
                ),
                (
                    (
                        "responsabilit\u00e9 civile et p\u00e9nale"
                        if fr
                        else "civil and criminal liability"
                    ),
                    True,
                ),
                (
                    " du Client au titre du droit marocain de la propri\u00e9t\u00e9 "
                    "intellectuelle."
                    if fr
                    else " under Moroccan intellectual property law."
                ),
            )

        if "c-image" in clauses:
            self._next_art(
                "DROIT \u00c0 L\u2019IMAGE ET COMMUNICATION"
                if fr
                else "IMAGE RIGHTS AND COMMUNICATION"
            )
            self._p(
                ("Sauf refus " if fr else "Unless the Client provides "),
                (
                    "expr\u00e8s et \u00e9crit" if fr else "express written objection",
                    True,
                ),
                (
                    " notifi\u00e9 dans les 8 jours suivant la signature du "
                    "pr\u00e9sent contrat, le Client autorise CASA DI LUSSO \u00e0\u202f:"
                    if fr
                    else " within 8 days of signing this contract, the Client authorizes "
                    "CASA DI LUSSO to:"
                ),
            )
            self._bullet(
                (
                    "Photographier et filmer le chantier et le projet "
                    "finalis\u00e9\u202f;"
                    if fr
                    else "Photograph and film the worksite and completed project;"
                ),
            )
            self._bullet(
                (
                    "Publier ces visuels sur son "
                    if fr
                    else "Publish these visuals on their "
                ),
                (
                    (
                        "site internet, r\u00e9seaux sociaux, portfolio"
                        if fr
                        else "website, social media, portfolio"
                    ),
                    True,
                ),
                (
                    " et supports de communication\u202f;"
                    if fr
                    else " and marketing materials;"
                ),
            )
            self._bullet(
                (
                    "Citer le projet \u00e0 titre de r\u00e9f\u00e9rence commerciale."
                    if fr
                    else "Reference the project as a commercial reference."
                ),
            )
            self._p(
                (
                    "Cette autorisation est conc\u00e9d\u00e9e \u00e0 titre "
                    if fr
                    else "This authorization is granted "
                ),
                (
                    (
                        "gratuit, non exclusif et pour une dur\u00e9e illimit\u00e9e"
                        if fr
                        else "free of charge, non-exclusively and for an unlimited duration"
                    ),
                    True,
                ),
                (
                    ". Le Prestataire s\u2019engage \u00e0 ne jamais mentionner les "
                    "informations personnelles du Client sans accord expr\u00e8s."
                    if fr
                    else ". The Service Provider undertakes never to mention the Client\u2019s "
                    "personal information without express consent."
                ),
            )

        if "c-confidential" in clauses:
            self._next_art("CONFIDENTIALIT\u00c9" if fr else "CONFIDENTIALITY")
            self._p(
                (
                    "Les deux parties s\u2019engagent mutuellement \u00e0 traiter "
                    "comme "
                    if fr
                    else "Both parties mutually undertake to treat as "
                ),
                (
                    "strictement confidentielles" if fr else "strictly confidential",
                    True,
                ),
                (
                    " toutes les informations \u00e9chang\u00e9es dans le cadre du "
                    "pr\u00e9sent contrat, incluant notamment\u202f:"
                    if fr
                    else " all information exchanged under this contract, including:"
                ),
            )
            self._bullet(
                (
                    "Les montants financiers et conditions tarifaires\u202f;"
                    if fr
                    else "Financial amounts and pricing conditions;"
                ),
            )
            self._bullet(
                (
                    "Les plans, documents techniques et sp\u00e9cifications du "
                    "projet\u202f;"
                    if fr
                    else "Plans, technical documents and project specifications;"
                ),
            )
            self._bullet(
                (
                    "Toute information relative aux m\u00e9thodes de travail du "
                    "Prestataire."
                    if fr
                    else "Any information relating to the Service Provider\u2019s working methods."
                ),
            )
            self._p(
                (
                    "Cette obligation de confidentialit\u00e9 survit \u00e0 la "
                    if fr
                    else "This confidentiality obligation survives "
                ),
                (
                    (
                        "r\u00e9siliation ou fin du contrat"
                        if fr
                        else "termination or expiry"
                    ),
                    True,
                ),
                (
                    " pour une dur\u00e9e de "
                    if fr
                    else " of the contract for a period of "
                ),
                ("5 ans" if fr else "5 years", True),
                ".",
            )

        if "c-sous-traiter" in clauses:
            self._next_art("DROIT DE SOUS-TRAITANCE" if fr else "SUBCONTRACTING RIGHTS")
            self._p(
                (
                    "Le Prestataire se r\u00e9serve le droit de "
                    if fr
                    else "The Service Provider reserves the right to "
                ),
                (
                    (
                        "recourir \u00e0 des sous-traitants"
                        if fr
                        else "engage qualified subcontractors"
                    ),
                    True,
                ),
                (
                    " qualifi\u00e9s pour l\u2019ex\u00e9cution de tout ou partie des "
                    "travaux pr\u00e9vus au pr\u00e9sent contrat, notamment pour les "
                    "corps de m\u00e9tier sp\u00e9cialis\u00e9s."
                    if fr
                    else " for the execution of all or part of the works under this contract, "
                    "particularly for specialized trades."
                ),
            )
            self._p(
                ("Le Prestataire demeure " if fr else "The Service Provider remains "),
                ("seul responsable" if fr else "solely responsible", True),
                (
                    " vis-\u00e0-vis du Client de la bonne ex\u00e9cution des travaux "
                    "sous-trait\u00e9s. Le Client renonce express\u00e9ment \u00e0 "
                    "invoquer la sous-traitance comme motif de contestation, "
                    "d\u00e8s lors que les travaux sont conformes au pr\u00e9sent "
                    "contrat."
                    if fr
                    else " to the Client for the proper execution of subcontracted works. "
                    "The Client expressly waives the right to invoke subcontracting as "
                    "a ground for dispute, provided works conform to this contract."
                ),
            )

        # ── ART – RÉVISION DES PRIX MATÉRIAUX (conditional) ──────────
        if "c-materiau-prix" in clauses:
            self._next_art(
                "CLAUSE DE R\u00c9VISION DES PRIX MAT\u00c9RIAUX"
                if fr
                else "MATERIAL PRICE REVISION CLAUSE"
            )
            self._p(
                (
                    "Dans l\u2019hypoth\u00e8se d\u2019une "
                    if fr
                    else "In the event of a "
                ),
                (
                    (
                        "hausse des prix des mat\u00e9riaux sup\u00e9rieure \u00e0 "
                        "10%"
                        if fr
                        else "material price increase exceeding 10%"
                    ),
                    True,
                ),
                (
                    " par rapport aux prix de r\u00e9f\u00e9rence au jour de la "
                    "signature, imputable \u00e0 des causes ext\u00e9rieures "
                    "(inflation, rupture d\u2019approvisionnement, d\u00e9cision "
                    "gouvernementale), le Prestataire se r\u00e9serve le droit "
                    "de\u202f:"
                    if fr
                    else " compared to reference prices at the date of signing, attributable "
                    "to external causes (inflation, supply shortage, government "
                    "decision), the Service Provider reserves the right to:"
                ),
            )
            self._bullet(
                (
                    "Notifier le Client par \u00e9crit des nouvelles conditions "
                    "tarifaires\u202f;"
                    if fr
                    else "Notify the Client in writing of the new pricing conditions;"
                ),
            )
            self._bullet(
                (
                    "Ajuster le montant du contrat \u00e0 hauteur de la hausse "
                    "constat\u00e9e, apr\u00e8s accord \u00e9crit du Client."
                    if fr
                    else "Adjust the contract amount proportionally to the increase, "
                    "subject to the Client\u2019s written agreement."
                ),
            )
            self._p(
                (
                    "En cas de d\u00e9saccord, les parties s\u2019engagent \u00e0 "
                    if fr
                    else "In case of disagreement, the parties commit to "
                ),
                (
                    "n\u00e9gocier de bonne foi" if fr else "negotiating in good faith",
                    True,
                ),
                (
                    " une solution dans un d\u00e9lai de 15 jours."
                    if fr
                    else " within 15 days."
                ),
            )

        if "c-force-maj" in clauses:
            self._next_art("FORCE MAJEURE")
            self._p(
                (
                    "Aucune partie ne pourra \u00eatre tenue responsable de "
                    "l\u2019inex\u00e9cution de ses obligations lorsque celle-ci "
                    "r\u00e9sulte d\u2019un "
                    if fr
                    else "Neither party shall be held liable for failure to perform its "
                    "obligations when such failure results from a "
                ),
                ("cas de force majeure" if fr else "force majeure event", True),
                (
                    ", c\u2019est-\u00e0-dire d\u2019un \u00e9v\u00e9nement "
                    "impr\u00e9visible, irr\u00e9sistible et ext\u00e9rieur aux "
                    "parties, incluant notamment\u202f:"
                    if fr
                    else ", being an unforeseeable, irresistible event beyond the parties\u2019 "
                    "control, including without limitation:"
                ),
            )
            self._bullet(
                (
                    "Catastrophes naturelles, \u00e9pid\u00e9mies, "
                    "pand\u00e9mies\u202f;"
                    if fr
                    else "Natural disasters, epidemics, pandemics;"
                ),
            )
            self._bullet(
                (
                    "Conflits arm\u00e9s, \u00e9meutes, actes terroristes\u202f;"
                    if fr
                    else "Armed conflicts, riots, terrorist acts;"
                ),
            )
            self._bullet(
                (
                    "D\u00e9cisions gouvernementales ou administratives\u202f;"
                    if fr
                    else "Government or administrative orders;"
                ),
            )
            self._bullet(
                (
                    "Gr\u00e8ves g\u00e9n\u00e9rales affectant les secteurs "
                    "d\u2019approvisionnement."
                    if fr
                    else "General strikes affecting supply chains."
                ),
            )
            self._p(
                (
                    "La partie invoquant la force majeure doit en notifier "
                    "l\u2019autre par "
                    if fr
                    else "The invoking party must notify the other "
                ),
                (
                    (
                        "\u00e9crit dans les 48 heures"
                        if fr
                        else "in writing within 48 hours"
                    ),
                    True,
                ),
                (
                    ". Si la force majeure dure plus de "
                    if fr
                    else ". If force majeure lasts more than "
                ),
                ("60 jours cons\u00e9cutifs" if fr else "60 consecutive days", True),
                (
                    ", chacune des parties pourra r\u00e9silier le contrat sans "
                    "indemnit\u00e9, sous r\u00e9serve de r\u00e8glement des "
                    "prestations d\u00e9j\u00e0 effectu\u00e9es."
                    if fr
                    else ", either party may terminate without compensation, "
                    "subject to payment for works already performed."
                ),
            )

        if "c-abandon-chant" in clauses:
            self._next_art(
                "ABANDON DE CHANTIER PAR LE CLIENT" if fr else "CLIENT SITE ABANDONMENT"
            )
            self._p(
                ("Est consid\u00e9r\u00e9 comme " if fr else ""),
                ("abandon de chantier" if fr else "Site abandonment", True),
                (
                    " le fait pour le Client de refuser ou d\u2019emp\u00eacher "
                    "l\u2019acc\u00e8s au chantier pendant plus de "
                    if fr
                    else " is defined as the Client refusing or preventing access to the "
                    "site for more than "
                ),
                ("15 jours cons\u00e9cutifs" if fr else "15 consecutive days", True),
                (
                    " sans justification valable et accept\u00e9e par le Prestataire."
                    if fr
                    else " without valid justification accepted by the Service Provider."
                ),
            )
            self._p(
                (
                    "En cas d\u2019abandon de chantier, le Prestataire est en droit "
                    "de\u202f:"
                    if fr
                    else "In case of site abandonment, the Service Provider is entitled to:"
                ),
            )
            self._bullet(
                (
                    "Facturer l\u2019int\u00e9gralit\u00e9 des travaux "
                    "r\u00e9alis\u00e9s et mat\u00e9riaux command\u00e9s\u202f;"
                    if fr
                    else "Invoice all completed works and ordered materials;"
                ),
            )
            self._bullet(
                (
                    "Retenir d\u00e9finitivement l\u2019ensemble des sommes "
                    "d\u00e9j\u00e0 vers\u00e9es\u202f;"
                    if fr
                    else "Definitively retain all amounts already paid;"
                ),
            )
            self._bullet(
                (
                    "R\u00e9cup\u00e9rer tout mat\u00e9riel et outil lui "
                    "appartenant\u202f;"
                    if fr
                    else "Recover all equipment and tools belonging to it;"
                ),
            )
            self._bullet(
                ("Facturer des " if fr else "Invoice "),
                (
                    (
                        "indemnit\u00e9s d\u2019immobilisation"
                        if fr
                        else "immobilization indemnities"
                    ),
                    True,
                ),
                (
                    " pour les \u00e9quipes mobilis\u00e9es."
                    if fr
                    else " for mobilized teams."
                ),
            )

        if "c-non-debauch" in clauses:
            self._next_art(
                "CLAUSE DE NON-D\u00c9BAUCHAGE DU PERSONNEL"
                if fr
                else "NON-POACHING OF STAFF CLAUSE"
            )
            self._p(
                (
                    "Le Client s\u2019engage, pendant la dur\u00e9e du pr\u00e9sent "
                    "contrat et pendant une p\u00e9riode de "
                    if fr
                    else "The Client undertakes, for the duration of this contract and "
                    "for a period of "
                ),
                ("24 mois" if fr else "24 months", True),
                (
                    " suivant son terme ou sa r\u00e9siliation, \u00e0 ne pas "
                    "directement ou indirectement\u202f:"
                    if fr
                    else " following its end or termination, not to directly or indirectly:"
                ),
            )
            self._bullet(
                (
                    "Recruter, solliciter ou engager tout membre du personnel ou "
                    "sous-traitant du Prestataire\u202f;"
                    if fr
                    else "Recruit, solicit, or hire any member of the Service Provider\u2019s staff "
                    "or subcontractors;"
                ),
            )
            self._bullet(
                (
                    "Inciter tout membre du personnel \u00e0 quitter le Prestataire."
                    if fr
                    else "Encourage any staff member to leave the Service Provider."
                ),
            )
            self._p(
                (
                    "En cas de violation, le Client s\u2019engage \u00e0 verser au "
                    "Prestataire une "
                    if fr
                    else "In case of violation, the Client agrees to pay the Service Provider a "
                ),
                ("indemnit\u00e9 forfaitaire" if fr else "lump-sum indemnity", True),
                (" \u00e9quivalente \u00e0 " if fr else " equivalent to "),
                (
                    "12 mois de salaire brut" if fr else "12 months\u2019 gross salary",
                    True,
                ),
                (
                    " de la personne concern\u00e9e."
                    if fr
                    else " of the person concerned."
                ),
            )

        self._next_art("R\u00c9SILIATION DU CONTRAT" if fr else "CONTRACT TERMINATION")
        _rn = self._n  # capture dynamic article number for sub-sections
        self._sub_title(
            f"{_rn}.1 \u2013 R\u00e9siliation \u00e0 l\u2019initiative du Client"
            if fr
            else f"{_rn}.1 \u2013 Termination by the Client"
        )
        self._p(
            (
                "Le Client peut mettre fin au pr\u00e9sent contrat par "
                if fr
                else "The Client may terminate this contract by "
            ),
            (
                (
                    "lettre recommand\u00e9e avec accus\u00e9 de r\u00e9ception"
                    if fr
                    else "registered mail with acknowledgment of receipt"
                ),
                True,
            ),
            (
                ". En pareil cas, le Client s\u2019engage \u00e0 r\u00e9gler\u202f:"
                if fr
                else ". In such case, the Client undertakes to pay:"
            ),
        )
        self._bullet(
            (
                "La totalit\u00e9 des travaux r\u00e9alis\u00e9s \u00e0 la date de "
                "r\u00e9siliation, au prorata\u202f;"
                if fr
                else "All works completed at the termination date, on a pro-rata basis;"
            ),
        )
        self._bullet(
            (
                "Le co\u00fbt de l\u2019ensemble des mat\u00e9riaux command\u00e9s, "
                "livr\u00e9s ou en cours de commande\u202f;"
                if fr
                else "The cost of all materials ordered, delivered, or on order;"
            ),
        )
        self._bullet(
            (
                "Les frais de d\u00e9mobilisation des \u00e9quipes\u202f;"
                if fr
                else "Team demobilization costs;"
            ),
        )
        self._bullet(
            ("Une " if fr else "A "),
            (
                "indemnit\u00e9 de r\u00e9siliation" if fr else "termination indemnity",
                True,
            ),
            (
                " de 15% du montant restant d\u00fb."
                if fr
                else " of 15% of the remaining amount due."
            ),
        )
        self._p(
            (
                "L\u2019acompte initial vers\u00e9 \u00e0 l\u2019ouverture du "
                "chantier reste "
                if fr
                else "The initial deposit paid at project start is "
            ),
            (
                (
                    "d\u00e9finitivement et irr\u00e9vocablement acquis"
                    if fr
                    else "definitively and irrevocably retained"
                ),
                True,
            ),
            (" au Prestataire." if fr else " by the Service Provider."),
        )
        self._sub_title(
            f"{_rn}.2 \u2013 R\u00e9siliation \u00e0 l\u2019initiative du Prestataire"
            if fr
            else f"{_rn}.2 \u2013 Termination by the Service Provider"
        )
        self._p(
            (
                "Le Prestataire peut r\u00e9silier le pr\u00e9sent contrat sans "
                "indemnit\u00e9 en cas de\u202f:"
                if fr
                else "The Service Provider may terminate this contract without compensation in "
                "the event of:"
            ),
        )
        self._bullet(
            (
                "Non-paiement d\u2019une ou plusieurs tranches au-del\u00e0 du "
                "d\u00e9lai tol\u00e9r\u00e9\u202f;"
                if fr
                else "Non-payment of one or more installments beyond the tolerated delay;"
            ),
        )
        self._bullet(
            (
                "Comportement abusif, mena\u00e7ant ou irrespectueux envers les "
                "\u00e9quipes\u202f;"
                if fr
                else "Abusive, threatening, or disrespectful behavior toward teams;"
            ),
        )
        self._bullet(
            (
                "Blocage r\u00e9p\u00e9t\u00e9 et injustifi\u00e9 du chantier\u202f;"
                if fr
                else "Repeated and unjustified site blockage;"
            ),
        )
        self._bullet(
            (
                "Refus r\u00e9p\u00e9t\u00e9 de fournir les validations "
                "n\u00e9cessaires \u00e0 l\u2019avancement des travaux\u202f;"
                if fr
                else "Repeated refusal to provide validations needed for works to progress;"
            ),
        )
        self._bullet(
            (
                "Insolvabilit\u00e9 ou redressement judiciaire du Client."
                if fr
                else "Client insolvency or receivership."
            ),
        )

        if "c-anti-litige" in clauses:
            tribunal = c.tribunal if c.tribunal else "Tanger"
            self._next_art(
                "M\u00c9DIATION OBLIGATOIRE ET R\u00c8GLEMENT DES LITIGES"
                if fr
                else "MANDATORY MEDIATION AND DISPUTE RESOLUTION"
            )
            self._p(
                (
                    "Avant tout recours judiciaire, les parties s\u2019engagent "
                    "express\u00e9ment \u00e0 rechercher une "
                    if fr
                    else "Before any legal action, parties expressly undertake to seek an "
                ),
                ("solution amiable" if fr else "amicable solution", True),
                (
                    " selon la proc\u00e9dure suivante\u202f:"
                    if fr
                    else " through the following procedure:"
                ),
            )
            self._nbullet(
                1,
                (
                    "La partie plaignante adresse une "
                    if fr
                    else "The complaining party sends a "
                ),
                (
                    "mise en demeure \u00e9crite" if fr else "formal written notice",
                    True,
                ),
                (
                    " \u00e0 l\u2019autre partie par lettre recommand\u00e9e\u202f;"
                    if fr
                    else " to the other party by registered mail;"
                ),
            )
            self._nbullet(
                2,
                ("L\u2019autre partie dispose de " if fr else "The other party has "),
                ("15 jours ouvr\u00e9s" if fr else "15 working days", True),
                (
                    " pour y r\u00e9pondre et proposer une solution\u202f;"
                    if fr
                    else " to respond and propose a solution;"
                ),
            )
            self._nbullet(
                3,
                (
                    "En l\u2019absence d\u2019accord dans un d\u00e9lai de "
                    if fr
                    else "Failing agreement within "
                ),
                ("30 jours" if fr else "30 days", True),
                (
                    " suivant la mise en demeure, les parties peuvent saisir un "
                    "m\u00e9diateur agr\u00e9\u00e9 d\u2019un commun accord\u202f;"
                    if fr
                    else " of the formal notice, parties may jointly appoint an accredited mediator;"
                ),
            )
            self._nbullet(
                4,
                (
                    "\u00c0 d\u00e9faut de m\u00e9diation aboutie, les parties "
                    "pourront saisir les juridictions comp\u00e9tentes."
                    if fr
                    else "Failing successful mediation, parties may refer "
                    "to the competent courts."
                ),
            )
            self._p(
                (
                    "Les parties conviennent d\u2019attribuer "
                    if fr
                    else "The parties agree to grant "
                ),
                ("comp\u00e9tence exclusive" if fr else "exclusive jurisdiction", True),
                (
                    " aux juridictions commerciales du "
                    if fr
                    else " to the commercial courts of the "
                ),
                (f"Tribunal de {tribunal}" if fr else f"Court of {tribunal}", True),
                (
                    " pour conna\u00eetre de tout litige relatif au pr\u00e9sent "
                    "contrat."
                    if fr
                    else " for any dispute arising from this contract."
                ),
            )
            self._p(
                (
                    "Le pr\u00e9sent contrat est r\u00e9gi et interpr\u00e9t\u00e9 "
                    "conform\u00e9ment au "
                    if fr
                    else "This contract is governed and interpreted in accordance with "
                ),
                ("droit marocain" if fr else "Moroccan law", True),
                (
                    ", notamment le Dahir des obligations et contrats (DOC)."
                    if fr
                    else ", in particular the Dahir of Obligations and Contracts (DOC)."
                ),
            )

        self._next_art(
            "LOI APPLICABLE ET DISPOSITION FINALE"
            if fr
            else "APPLICABLE LAW AND FINAL PROVISIONS"
        )
        self._p(
            (
                "Le pr\u00e9sent contrat est r\u00e9gi et interpr\u00e9t\u00e9 "
                "conform\u00e9ment au "
                if fr
                else "This contract is governed and interpreted in accordance with the "
            ),
            (
                "droit du Royaume du Maroc" if fr else "laws of the Kingdom of Morocco",
                True,
            ),
            (
                ". Tout sujet non express\u00e9ment trait\u00e9 par le pr\u00e9sent "
                "contrat sera r\u00e9gi par les dispositions l\u00e9gales marocaines "
                "applicables, notamment le Dahir du 12 ao\u00fbt 1913 formant Code "
                "des Obligations et des Contrats (DOC)."
                if fr
                else ". Any matter not expressly covered by this contract shall be "
                "governed by applicable Moroccan legal provisions, notably the Dahir of "
                "12 August 1913 forming the Code of Obligations and Contracts (DOC)."
            ),
        )
        self._p(
            (
                "Le pr\u00e9sent contrat constitue l\u2019int\u00e9gralit\u00e9 de "
                "l\u2019accord entre les parties et annule et remplace tous les "
                "accords, n\u00e9gociations et propositions ant\u00e9rieurs relatifs "
                "\u00e0 son objet."
                if fr
                else "This contract constitutes the entire agreement between the parties "
                "and supersedes all prior agreements, negotiations and proposals "
                "relating to its subject matter."
            ),
        )
        self._p(
            (
                "Si l\u2019une des clauses du pr\u00e9sent contrat est "
                "d\u00e9clar\u00e9e nulle ou inapplicable, les autres clauses "
                "demeurent "
                if fr
                else "If any clause of this contract is declared void or unenforceable, "
                "the remaining clauses shall remain "
            ),
            ("pleinement en vigueur" if fr else "in full force and effect", True),
            ".",
        )
        if c.exclusions:
            self._divider()
            self._p(
                (
                    (
                        "Exclusions contractuelles express\u00e9ment convenues\u202f:\n"
                        if fr
                        else "Expressly agreed exclusions:\n"
                    ),
                    True,
                ),
                c.exclusions,
            )
        if c.clause_spec:
            self._divider()
            self._p(
                (
                    (
                        "Clauses sp\u00e9cifiques additionnelles\u202f:\n"
                        if fr
                        else "Additional specific clauses:\n"
                    ),
                    True,
                ),
                c.clause_spec,
            )
        if c.annexes:
            self._highlight(
                (
                    (
                        "Annexes jointes au pr\u00e9sent contrat\u202f:\n"
                        if fr
                        else "Annexes attached to this contract:\n"
                    ),
                    True,
                ),
                c.annexes,
            )

    def _build(self):
        c = self.contract
        lang = self.language
        fr = self.fr

        confid_label = CONFID_LABELS.get(
            c.confidentialite or "confidentiel", "CONFIDENTIEL"
        )
        ctype_display = CTYPES_DISPLAY[lang].get(
            c.type_contrat or "travaux_finition", c.type_contrat or ""
        )
        date_str = _fmt_date(c.date_contrat)
        version_str = "v1.0 \u2013 D\u00e9finitif" if fr else "v1.0 \u2013 Final"
        ville = c.ville_signature if c.ville_signature else "Tanger"
        ref = c.numero_contrat or ""

        client_nom = (
            c.client_nom
            or "\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026"
        )
        client_cin = (
            c.client_cin
            or "\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026"
        )
        client_qualite = QUALITE_LABELS[lang].get(
            c.client_qualite or "",
            c.client_qualite or ("Personne Physique" if fr else "Individual"),
        )
        client_adresse = (
            c.client_adresse
            or "\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026"
        )
        client_tel = (
            c.client_tel
            or "\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026"
        )
        client_email = c.client_email or ""
        sig_role = (
            ("Repr\u00e9sentant l\u00e9gal" if fr else "Legal Representative")
            if _is_societe(c.client_qualite)
            else ("Particulier" if fr else "Individual")
        )
        sig_name = (
            c.client_nom
            or "\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026"
        )

        # ── assemble ──
        self._add_top_strip()
        self._add_header(c, ref, date_str, version_str, confid_label, ville)
        self._add_confidential()
        self._add_title_block(ctype_display)
        self._add_parties(
            c,
            client_nom,
            client_cin,
            client_qualite,
            client_adresse,
            client_tel,
            client_email,
        )

        # Articles — direct python-docx, same content as the PDF
        self._build_articles()

        # Signatures & footer
        self._add_signatures(ville, date_str, sig_name, sig_role)
        self._add_footer(ref, version_str, confid_label)

    def generate_response(self) -> HttpResponse:
        self._build()
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        filename = f"contrat_{self.contract.id}_{self.contract.numero_contrat.replace('/', '-')}.docx"
        response = HttpResponse(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response
